"""
Module de rapport d'application pour App_PADESCE.

Fournit les fonctions de construction, d'export et d'envoi du rapport
journalier de l'application.
"""

import io
import logging
import re
from datetime import date, datetime, time, timedelta
from email.mime.image import MIMEImage
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from typing import Any, Iterable

from django.conf import settings
from django.contrib.auth import get_user_model
from django.core.mail import EmailMessage, get_connection
from django.db.models import Count, Max, Min, Q
from django.http import QueryDict
from django.utils import timezone

from App_PADESCE.appels.models import (
    CALL_ANALYSIS_THRESHOLD_STATUSES,
    CALL_COMPLETED_STATUSES,
    CALL_TENTATIVE_STATUSES,
    Appel,
    AppelCGA,
    AppelFormateur,
)
from App_PADESCE.appels.cga_report import (
    build_cga_calls_report_workbook,
    get_cga_calls_report_filename,
)
from App_PADESCE.core.analysis_rules import analysis_threshold_target
from App_PADESCE.core.models import UserActivity
from App_PADESCE.formations.models import Classe
from App_PADESCE.reporting.network_excel import build_padesce_source_index, normalize_network_lookup
from App_PADESCE.reporting.padesce_calls_excel import build_padesce_calls_report

logger = logging.getLogger(__name__)
CONSOLE_EMAIL_BACKEND = "django.core.mail.backends.console.EmailBackend"
SMTP_EMAIL_BACKEND = "django.core.mail.backends.smtp.EmailBackend"

# ---------------------------------------------------------------------------
# Helpers — imports optionnels (python-docx)
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    HAS_DOCX = True
except ImportError:  # pragma: no cover
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Utilitaires dates / config
# ---------------------------------------------------------------------------


def parse_report_dates(start_str: str | None, end_str: str | None) -> tuple[date, date]:
    today = timezone.localdate()
    try:
        start = date.fromisoformat(str(start_str).strip()) if start_str else today
    except ValueError:
        start = today
    try:
        end = date.fromisoformat(str(end_str).strip()) if end_str else today
    except ValueError:
        end = today
    if end < start:
        end = start
    return start, end


def normalize_report_call_scope(value: str | None) -> str:
    return "cga" if str(value or "").strip().lower() == "cga" else "padesce"


def parse_email_recipients(raw: str | Iterable[str] | None) -> list[str]:
    if raw is None:
        return []
    if isinstance(raw, str):
        values = re.split(r"[,;\n]+", raw)
    else:
        values = []
        for item in raw:
            values.extend(re.split(r"[,;\n]+", str(item)))

    recipients: list[str] = []
    seen: set[str] = set()
    for value in values:
        email = str(value or "").strip()
        if not email:
            continue
        key = email.lower()
        if key in seen:
            continue
        seen.add(key)
        recipients.append(email)
    return recipients


def get_report_email_recipients(raw_override: str | Iterable[str] | None = None) -> list[str]:
    raw = raw_override
    if raw is None:
        raw = (getattr(settings, "REPORT_EMAIL_TO", "") or "").strip()
        if not raw:
            raw = settings.__dict__.get("REPORT_EMAIL_TO") or ""
    return parse_email_recipients(raw)


def _build_report_email_connection() -> tuple[object | None, str | None]:
    backend = str(getattr(settings, "EMAIL_BACKEND", "") or "").strip()
    if backend != CONSOLE_EMAIL_BACKEND:
        return None, None

    missing = [
        field_name
        for field_name in (
            "EMAIL_HOST",
            "EMAIL_HOST_USER",
            "EMAIL_HOST_PASSWORD",
            "DEFAULT_FROM_EMAIL",
        )
        if not getattr(settings, field_name, "")
    ]
    if missing:
        return (
            None,
            "EMAIL_BACKEND est en mode console et la configuration SMTP est incomplète: "
            + ", ".join(missing),
        )

    logger.info("Envoi rapport: backend console détecté, bascule forcée vers SMTP.")
    return (
        get_connection(
            backend=SMTP_EMAIL_BACKEND,
            host=settings.EMAIL_HOST,
            port=getattr(settings, "EMAIL_PORT", None),
            username=settings.EMAIL_HOST_USER,
            password=settings.EMAIL_HOST_PASSWORD,
            use_tls=getattr(settings, "EMAIL_USE_TLS", False),
            use_ssl=getattr(settings, "EMAIL_USE_SSL", False),
            timeout=getattr(settings, "EMAIL_TIMEOUT", None),
        ),
        None,
    )


# ---------------------------------------------------------------------------
# Section 3 — Moteur principal
# ---------------------------------------------------------------------------


def build_application_report(
    start_date: date,
    end_date: date,
    selected_class_code: str | None = None,
    call_scope: str = "padesce",
) -> dict:
    tz = timezone.get_current_timezone()
    start_dt = timezone.make_aware(datetime.combine(start_date, time.min), tz)
    end_dt = timezone.make_aware(datetime.combine(end_date, time.max), tz)
    now = timezone.localtime()
    report_day_start = timezone.make_aware(datetime.combine(end_date, time.min), tz)
    report_day_end = timezone.make_aware(datetime.combine(end_date, time.max), tz)
    daily_reference_is_today = end_date == timezone.localdate()
    daily_reference_label = (
        "aujourd'hui" if daily_reference_is_today else f"le {end_date.strftime('%d/%m/%Y')}"
    )
    selected_class_code = (selected_class_code or "").strip()
    selected_class = _get_selected_class(selected_class_code)
    call_scope = normalize_report_call_scope(call_scope)
    call_scope_label = "CGA" if call_scope == "cga" else "PADESCE / Formateurs"

    user_model = get_user_model()
    active_period_users = UserActivity.objects.filter(
        last_seen__gte=start_dt, last_seen__lte=end_dt
    )
    active_24h_users = UserActivity.objects.filter(last_seen__gte=now - timedelta(hours=24))

    padesce_qs = Appel.objects.filter(
        is_active=True, updated_at__gte=start_dt, updated_at__lte=end_dt
    )
    formateur_qs = AppelFormateur.objects.filter(
        is_active=True, updated_at__gte=start_dt, updated_at__lte=end_dt
    )
    cga_qs = AppelCGA.objects.filter(
        is_active=True, updated_at__gte=start_dt, updated_at__lte=end_dt
    )

    if call_scope == "cga":
        call_sources = [_build_call_source_summary("CGA", cga_qs)]
    else:
        call_sources = [
            _build_call_source_summary("PADESCE", padesce_qs),
            _build_call_source_summary("Formateurs", formateur_qs),
        ]
    call_totals = {
        "total": sum(source["total"] for source in call_sources),
        "completed": sum(source["completed"] for source in call_sources),
        "pending": sum(source["pending"] for source in call_sources),
        "in_progress": sum(source["in_progress"] for source in call_sources),
        "callbacks": sum(source["callbacks"] for source in call_sources),
        "with_audio": sum(source["with_audio"] for source in call_sources),
    }
    processed_total = sum(source["processed"] for source in call_sources)
    call_totals["processed"] = processed_total
    call_totals["completion_rate"] = (
        round((call_totals["completed"] / processed_total) * 100, 2) if processed_total else 0.0
    )

    hourly_rows = (
        _build_hourly_rows(_hourly_completed_counts(cga_qs))
        if call_scope == "cga"
        else _build_hourly_rows(
            _hourly_completed_counts(padesce_qs),
            _hourly_completed_counts(formateur_qs),
        )
    )
    best_hour = (
        max(hourly_rows, key=lambda row: (row["completed"], row["total"], -row["hour"]))
        if hourly_rows
        else None
    )

    bug_summary = _build_bug_summary(start_dt, end_dt)
    anomaly_summary = (
        _build_cga_anomaly_summary(cga_qs)
        if call_scope == "cga"
        else _build_anomaly_summary(selected_class, selected_class_code=selected_class_code)
    )
    analysis_summary = _build_analysis_summary(padesce_qs)
    cga_dimensions = _build_cga_dimension_summary(cga_qs)
    formateurs_summary = _build_formateurs_satisfaction_summary(start_dt, end_dt)
    user_call_rows = (
        _build_user_call_rows(cga_qs, day_start=report_day_start, day_end=report_day_end)
        if call_scope == "cga"
        else _build_user_call_rows(
            padesce_qs,
            formateur_qs,
            day_start=report_day_start,
            day_end=report_day_end,
        )
    )
    cga_outcomes = {
        "interesses": cga_qs.filter(interet="OUI").count(),
        "pas_interesses": cga_qs.filter(interet="NON").count(),
        "indisponibles": cga_qs.filter(indisponible="OUI").count(),
        "faux_numeros": cga_qs.filter(mauvais_numero="OUI").count(),
    }

    recipients = get_report_email_recipients()
    mail_status = {
        "configured_recipients": recipients,
        "missing_configuration": not bool(recipients),
        "smtp_ready": bool(getattr(settings, "EMAIL_HOST", ""))
        and bool(getattr(settings, "DEFAULT_FROM_EMAIL", "")),
    }

    return {
        "generated_at": now,
        "start_date": start_date,
        "end_date": end_date,
        "start_dt": start_dt,
        "end_dt": end_dt,
        "period_days": (end_date - start_date).days + 1,
        "call_scope": call_scope,
        "call_scope_label": call_scope_label,
        "users": {
            "total": user_model.objects.count(),
            "active": user_model.objects.filter(is_active=True).count(),
            "staff": user_model.objects.filter(is_staff=True).count(),
            "superusers": user_model.objects.filter(is_superuser=True).count(),
            "seen_24h": active_24h_users.count(),
            "seen_period": active_period_users.count(),
            "called_today": sum(1 for row in user_call_rows if row["calls_today"] > 0),
        },
        "calls": call_totals,
        "call_sources": call_sources,
        "hourly_rows": hourly_rows,
        "best_hour": best_hour,
        "bugs": bug_summary,
        "anomalies": anomaly_summary,
        "analysis": analysis_summary,
        "cga_dimensions": cga_dimensions,
        "formateurs_summary": formateurs_summary,
        "cga_outcomes": cga_outcomes,
        "user_call_rows": user_call_rows,
        "mail_status": mail_status,
        "daily_reference_date": end_date,
        "daily_reference_is_today": daily_reference_is_today,
        "daily_reference_label": daily_reference_label,
        "selected_class_code": selected_class.code if selected_class else selected_class_code,
    }


# ---------------------------------------------------------------------------
# Section 4 — Export Word
# ---------------------------------------------------------------------------


def export_application_report_word(report: dict) -> bytes:
    if not HAS_DOCX:
        raise ImportError("python-docx est requis pour l'export Word.")

    document = Document()
    _configure_report_document(document)

    header_table = document.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    logo_cell, text_cell = header_table.rows[0].cells
    _clear_cell_border(logo_cell)
    _clear_cell_border(text_cell)
    logo_path = _get_report_logo_path()
    if logo_path and logo_path.exists():
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(1.35))

    title = text_cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("RAPPORT JOURNALIER")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(76, 29, 149)

    subtitle = text_cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle.add_run(
        f"Periode: du {report['start_date'].strftime('%d/%m/%Y')} au {report['end_date'].strftime('%d/%m/%Y')}"
    )
    subtitle_run.font.size = Pt(11)

    meta = text_cell.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.add_run(
        f"Date de generation: {timezone.localtime(report['generated_at']).strftime('%d/%m/%Y a %H:%M')}"
    )

    best_hour_label = report["best_hour"]["label"] if report["best_hour"] else "Aucune donnee"
    best_hour_value = report["best_hour"]["completed"] if report["best_hour"] else 0

    document.add_paragraph("")
    _add_heading(document, "1. Resume")
    summary_rows = [
        ("Utilisateurs ayant appele", report["users"]["called_today"]),
        ("Appels termines", report["calls"]["completed"]),
        ("Heure la plus performante", f"{best_hour_label} ({best_hour_value} appels termines)"),
    ]
    if report.get("call_scope") == "cga":
        cga_dimensions = report["cga_dimensions"]
        summary_rows.extend(
            [
                ("Regimes", cga_dimensions["regimes_count"]),
                ("Centres", cga_dimensions["centres_count"]),
                ("CRI", cga_dimensions["cris_count"]),
                ("Villes", cga_dimensions["villes_count"]),
                ("Interesses", cga_dimensions["interesses"]),
                ("Pas interesses", cga_dimensions["pas_interesses"]),
                ("Indisponibles", cga_dimensions["indisponibles"]),
                ("Faux numeros", cga_dimensions["faux_numeros"]),
            ]
        )
    else:
        summary_rows.extend(
            [
                ("Classes analysees", report["analysis"]["classes_count"]),
                ("Prestations analysees", report["analysis"]["prestations_count"]),
                ("Prestataires analyses", report["analysis"]["prestataires_count"]),
                ("Beneficiaires analyses", report["analysis"]["beneficiaires_count"]),
            ]
        )
    _add_metric_table(
        document,
        summary_rows,
    )

    _add_heading(document, "2. Situation des appels")
    calls_table = document.add_table(rows=1, cols=5)
    calls_table.style = "Table Grid"
    calls_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Source", "Total", "Effectues", "Termines", "Audios"]
    for idx, label in enumerate(headers):
        _write_header_cell(calls_table.rows[0].cells[idx], label)
    for source in report["call_sources"]:
        row = calls_table.add_row().cells
        row[0].text = source["label"]
        row[1].text = str(source["total"])
        row[2].text = str(source["processed"])
        row[3].text = str(source["completed"])
        row[4].text = str(source["with_audio"])

    _add_heading(document, "3. Activite par utilisateur")
    users_table = document.add_table(rows=1, cols=6)
    users_table.style = "Table Grid"
    users_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(
        [
            "Utilisateur",
            "Appels effectues",
            "Appels termines",
            "Temps estime",
            "Premiere activite",
            "Derniere activite",
        ]
    ):
        _write_header_cell(users_table.rows[0].cells[idx], label)
    for item in report["user_call_rows"]:
        row = users_table.add_row().cells
        row[0].text = item["username"]
        row[1].text = str(item["calls_made"])
        row[2].text = str(item["completed_calls"])
        row[3].text = item["time_spent_label"]
        row[4].text = item["first_activity_label"]
        row[5].text = item["last_activity_label"]

    _add_heading(document, "4. Repartition horaire")
    hourly_table = document.add_table(rows=1, cols=3)
    hourly_table.style = "Table Grid"
    hourly_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(["Heure", "Total", "Termines"]):
        _write_header_cell(hourly_table.rows[0].cells[idx], label)
    top_hour_rows = sorted(
        report["hourly_rows"], key=lambda row: (-row["completed"], -row["total"], row["hour"])
    )[:8]
    for row_data in top_hour_rows:
        row = hourly_table.add_row().cells
        row[0].text = row_data["label"]
        row[1].text = str(row_data["total"])
        row[2].text = str(row_data["completed"])

    if report.get("call_scope") == "cga":
        cga_dimensions = report["cga_dimensions"]
        _add_heading(document, "5. Analyse CGA")
        _add_metric_table(
            document,
            [
                ("Regimes", cga_dimensions["regimes_count"]),
                ("Centres", cga_dimensions["centres_count"]),
                ("CRI", cga_dimensions["cris_count"]),
                ("Villes", cga_dimensions["villes_count"]),
                ("Interesses", cga_dimensions["interesses"]),
                ("Pas interesses", cga_dimensions["pas_interesses"]),
                ("Indisponibles", cga_dimensions["indisponibles"]),
                ("Faux numeros", cga_dimensions["faux_numeros"]),
            ],
        )
    else:
        _add_heading(document, "5. Analyse satisfaction – Apprenants")
        _add_metric_table(
            document,
            [
                ("Classes analysees", report["analysis"]["classes_count"]),
                ("Prestations analysees", report["analysis"]["prestations_count"]),
                ("Prestataires analyses", report["analysis"]["prestataires_count"]),
                ("Beneficiaires analyses", report["analysis"]["beneficiaires_count"]),
            ],
        )

    fs = report.get("formateurs_summary", {})
    if report.get("call_scope") != "cga" and fs:
        _add_heading(document, "6. Analyse satisfaction – Formateurs")
        _add_metric_table(
            document,
            [
                ("Appels formateurs termines (periode)", fs.get("total_termines", 0)),
                ("Avec scores Q1-Q3", fs.get("with_scores", 0)),
                ("Moy. Q1 – Prerequis apprenants", fs.get("avg_q1", 0)),
                ("Moy. Q2 – Interaction apprenants", fs.get("avg_q2", 0)),
                ("Moy. Q3 – Competences acquises", fs.get("avg_q3", 0)),
            ],
        )
        if fs.get("by_prestataire"):
            form_table = document.add_table(rows=1, cols=5)
            form_table.style = "Table Grid"
            form_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for idx, label in enumerate(["Prestataire", "Nb appels", "Q1", "Q2", "Q3"]):
                _write_header_cell(form_table.rows[0].cells[idx], label)
            for row_data in fs["by_prestataire"]:
                row = form_table.add_row().cells
                row[0].text = row_data["prestataire"]
                row[1].text = str(row_data["nb"])
                row[2].text = str(row_data["avg_q1"])
                row[3].text = str(row_data["avg_q2"])
                row[4].text = str(row_data["avg_q3"])

    _add_heading(document, "7. Incidents")
    _add_metric_table(
        document,
        [
            ("Bugs signales", report["bugs"]["alarms_total"]),
            ("Bugs resolus", report["bugs"]["alarms_resolved"]),
            ("Bugs ouverts", report["bugs"]["alarms_open"]),
            ("Erreurs logs", report["bugs"]["log_errors"]),
            ("Erreurs critiques", report["bugs"]["log_critical"]),
        ],
    )
    if report["bugs"]["recent_alarms"]:
        incidents_table = document.add_table(rows=1, cols=4)
        incidents_table.style = "Table Grid"
        incidents_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, label in enumerate(["Incident", "Module", "Date", "Etat"]):
            _write_header_cell(incidents_table.rows[0].cells[idx], label)
        for alarm in report["bugs"]["recent_alarms"]:
            row = incidents_table.add_row().cells
            row[0].text = alarm["title"]
            row[1].text = alarm["module"]
            row[2].text = alarm["created_at"]
            row[3].text = alarm["status"]

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


# ---------------------------------------------------------------------------
# Section 5 — Envoi mail + HTML
# ---------------------------------------------------------------------------


def send_report_by_email(report: dict) -> dict:
    recipients = get_report_email_recipients()
    if not recipients:
        return {"ok": False, "detail": "Aucun destinataire email configure"}
    if not getattr(settings, "EMAIL_HOST", ""):
        return {"ok": False, "detail": "EMAIL_HOST non configure"}
    if not getattr(settings, "DEFAULT_FROM_EMAIL", ""):
        return {"ok": False, "detail": "DEFAULT_FROM_EMAIL non configure"}
    connection, connection_error = _build_report_email_connection()
    if connection_error:
        return {"ok": False, "detail": connection_error}

    subject = (
        f"Rapport application NAUMUR CALL APP - {report['start_date']} au {report['end_date']}"
    )
    body = build_report_email_html(report)
    message = EmailMessage(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=recipients,
        connection=connection,
    )
    message.content_subtype = "html"
    logo_path = _get_report_logo_path()
    if logo_path and logo_path.exists():
        with logo_path.open("rb") as handle:
            logo_image = MIMEImage(handle.read())
            logo_image.add_header("Content-ID", "<logo_cga_naumur>")
            logo_image.add_header("Content-Disposition", "inline", filename=logo_path.name)
            message.attach(logo_image)
    if HAS_DOCX:
        try:
            word_payload = export_application_report_word(report)
            message.attach(
                f"rapport_application_{report['start_date']}_{report['end_date']}.docx",
                word_payload,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception:
            pass
    try:
        message.send(fail_silently=False)
    except Exception as exc:
        logger.exception("Echec d'envoi du rapport journalier par email.")
        return {"ok": False, "detail": f"Echec de l'envoi SMTP: {exc}"}
    return {"ok": True, "detail": f"Mail envoye a {len(recipients)} destinataire(s)"}


def _format_file_size(size: Any) -> str:
    try:
        value = float(size)
    except (TypeError, ValueError):
        return ""
    if value < 1024:
        return f"{int(value)} o"
    if value < 1024 * 1024:
        return f"{value / 1024:.1f} Ko"
    if value < 1024 * 1024 * 1024:
        return f"{value / (1024 * 1024):.1f} Mo"
    return f"{value / (1024 * 1024 * 1024):.1f} Go"


def _format_iso_datetime_label(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    try:
        parsed = datetime.fromisoformat(raw)
    except ValueError:
        return raw
    if timezone.is_aware(parsed):
        parsed = timezone.localtime(parsed)
    return parsed.strftime("%d/%m/%Y %H:%M")


def build_backup_status_summary(backup_job_id: str | None = None, *, backup_error: str = "") -> dict:
    from App_PADESCE.core import backup_manager

    error_text = str(backup_error or "").strip()
    if error_text:
        return {
            "status": "error",
            "label": "Erreur",
            "job_id": backup_job_id or "",
            "detail": error_text,
            "backup_file": "",
            "file_size_label": "",
            "finished_at_label": "",
            "retention_days": None,
            "purged_backups": [],
        }

    entry = None
    if backup_job_id:
        entry = backup_manager.get_job(backup_job_id) or backup_manager.get_history_entry(backup_job_id)
    if entry is None:
        entry = backup_manager.get_latest_history_entry()
    if entry is None:
        return {
            "status": "unknown",
            "label": "Indisponible",
            "job_id": backup_job_id or "",
            "detail": "Aucun historique de backup disponible.",
            "backup_file": "",
            "file_size_label": "",
            "finished_at_label": "",
            "retention_days": None,
            "purged_backups": [],
        }

    status = str(entry.get("status", "unknown") or "unknown")
    label_map = {
        "success": "Succes",
        "partial": "Partiel",
        "error": "Erreur",
        "running": "En cours",
        "pending": "En attente",
    }
    file_size_label = _format_file_size(entry.get("file_size"))
    finished_at_label = _format_iso_datetime_label(entry.get("finished_at") or entry.get("started_at"))
    backup_file = str(entry.get("backup_file", "") or "").strip()
    purged_backups = list(entry.get("purged_backups") or [])
    detail_parts = []
    if backup_file:
        detail_parts.append(f"Fichier: {backup_file}")
    if file_size_label:
        detail_parts.append(f"Taille: {file_size_label}")
    if finished_at_label:
        detail_parts.append(f"Heure: {finished_at_label}")
    if purged_backups:
        detail_parts.append(f"Purge: {len(purged_backups)} fichier(s)")
    if entry.get("error"):
        detail_parts.append(str(entry["error"]))
    if not detail_parts:
        detail_parts.append(str(entry.get("message", "") or "Statut de backup disponible."))

    return {
        "status": status,
        "label": label_map.get(status, status.capitalize() or "Inconnu"),
        "job_id": str(entry.get("id", "") or backup_job_id or ""),
        "detail": " | ".join(detail_parts),
        "backup_file": backup_file,
        "file_size_label": file_size_label,
        "finished_at_label": finished_at_label,
        "retention_days": entry.get("retention_days"),
        "purged_backups": purged_backups,
    }


def build_daily_digest_email_html(report: dict, backup_summary: dict) -> str:
    best_hour_label = report["best_hour"]["label"] if report.get("best_hour") else "Aucune donnee"
    best_hour_value = report["best_hour"]["completed"] if report.get("best_hour") else 0
    status_color = {
        "success": "#166534",
        "partial": "#92400e",
        "error": "#991b1b",
        "running": "#1d4ed8",
        "pending": "#6b7280",
        "unknown": "#6b7280",
    }.get(backup_summary.get("status"), "#6b7280")
    purge_line = ""
    if backup_summary.get("purged_backups"):
        purge_line = (
            f"<p style='margin:8px 0 0;color:#5b6472;'>Purge automatique: "
            f"{len(backup_summary['purged_backups'])} ancien(s) backup(s) supprime(s).</p>"
        )

    return f"""
<div style="margin:0;padding:24px;background:#f4f4f6;font-family:Calibri,Arial,sans-serif;color:#222222;">
  <div style="max-width:900px;margin:0 auto;background:#ffffff;border:1px solid #d9d9e3;">
    <div style="padding:24px 28px;border-bottom:3px solid #4c1d95;background:#ffffff;">
      <table role="presentation" style="width:100%;border-collapse:collapse;">
        <tr>
          <td style="width:120px;vertical-align:top;text-align:left;padding-right:16px;">
            <img src="cid:logo_cga_naumur" alt="Logo CGA Naumur" style="max-width:88px;height:auto;display:block;">
          </td>
          <td style="vertical-align:top;">
            <div style="font-size:12px;letter-spacing:.08em;text-transform:uppercase;color:#4c1d95;font-weight:700;">Digest quotidien</div>
            <h1 style="margin:8px 0 6px;font-size:24px;line-height:1.25;color:#1f1630;">Backup + rapports PADESCE / CGA</h1>
            <p style="margin:0;font-size:14px;color:#555555;">Periode du {report["start_date"].strftime("%d/%m/%Y")} au {report["end_date"].strftime("%d/%m/%Y")}</p>
            <p style="margin:4px 0 0;font-size:13px;color:#666666;">Genere le {timezone.localtime(report["generated_at"]).strftime("%d/%m/%Y a %H:%M")}</p>
          </td>
        </tr>
      </table>
    </div>

    <div style="padding:24px 28px;">
      <div style="border:1px solid #dfe1ea;padding:18px 20px;margin-bottom:18px;border-left:6px solid {status_color};">
        <div style="font-size:12px;text-transform:uppercase;letter-spacing:.08em;color:#6b7280;font-weight:700;">Statut backup</div>
        <div style="font-size:24px;font-weight:800;color:{status_color};margin:8px 0 6px;">{backup_summary["label"]}</div>
        <p style="margin:0;color:#4b5563;">{backup_summary["detail"]}</p>
        {purge_line}
      </div>

      <div style="display:flex;flex-wrap:wrap;gap:14px;margin-bottom:22px;">
        {_email_card("Appels termines", report["calls"]["completed"])}
        {_email_card("Utilisateurs actifs", report["users"]["called_today"])}
        {_email_card("Heure forte", f"{best_hour_label} ({best_hour_value})")}
      </div>

      <div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:12px;margin-bottom:20px;">
        {_email_stat("Classes analysees", report["analysis"]["classes_count"])}
        {_email_stat("Prestations analysees", report["analysis"]["prestations_count"])}
        {_email_stat("Prestataires analyses", report["analysis"]["prestataires_count"])}
        {_email_stat("Beneficiaires analyses", report["analysis"]["beneficiaires_count"])}
      </div>

      <div style="border:1px solid #ece5fb;padding:16px 18px;margin-bottom:18px;background:#faf7ff;">
        <div style="font-size:16px;font-weight:700;color:#2f2550;margin-bottom:8px;">Pieces jointes incluses</div>
        <p style="margin:0;color:#5b6472;">Le mail contient le rapport Word d'activite, le rapport Excel PADESCE et le rapport Excel CGA.</p>
      </div>

      <div style="border:1px solid #ece5fb;padding:16px 18px;background:#ffffff;">
        <div style="font-size:16px;font-weight:700;color:#2f2550;margin-bottom:8px;">Incidents</div>
        <p style="margin:0;color:#5b6472;">Bugs signales: {report["bugs"]["alarms_total"]} | Bugs ouverts: {report["bugs"]["alarms_open"]} | Erreurs logs: {report["bugs"]["log_errors"]}</p>
      </div>
    </div>
  </div>
</div>
"""


def send_daily_digest_email(
    start_date: date,
    end_date: date,
    *,
    backup_job_id: str | None = None,
    backup_error: str = "",
    recipients: str | Iterable[str] | None = None,
) -> dict:
    resolved_recipients = get_report_email_recipients(recipients)
    if not resolved_recipients:
        return {"ok": False, "detail": "Aucun destinataire email configure"}
    if not getattr(settings, "EMAIL_HOST", ""):
        return {"ok": False, "detail": "EMAIL_HOST non configure"}
    if not getattr(settings, "DEFAULT_FROM_EMAIL", ""):
        return {"ok": False, "detail": "DEFAULT_FROM_EMAIL non configure"}

    connection, connection_error = _build_report_email_connection()
    if connection_error:
        return {"ok": False, "detail": connection_error}

    report = build_application_report(start_date, end_date)
    backup_summary = build_backup_status_summary(backup_job_id, backup_error=backup_error)
    subject = f"Digest quotidien NAUMUR CALL APP - {start_date} au {end_date}"
    body = build_daily_digest_email_html(report, backup_summary)
    message = EmailMessage(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=resolved_recipients,
        connection=connection,
    )
    message.content_subtype = "html"

    logo_path = _get_report_logo_path()
    if logo_path and logo_path.exists():
        with logo_path.open("rb") as handle:
            logo_image = MIMEImage(handle.read())
            logo_image.add_header("Content-ID", "<logo_cga_naumur>")
            logo_image.add_header("Content-Disposition", "inline", filename=logo_path.name)
            message.attach(logo_image)

    attachments: list[str] = []
    if HAS_DOCX:
        try:
            filename = f"rapport_application_{start_date}_{end_date}.docx"
            message.attach(
                filename,
                export_application_report_word(report),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            attachments.append(filename)
        except Exception:
            logger.exception("Impossible de joindre le rapport Word du digest quotidien.")

    try:
        with TemporaryDirectory() as temp_dir:
            padesce_path = Path(temp_dir) / f"rapport_appels_padesce_{start_date}_{end_date}.xlsx"
            build_padesce_calls_report(padesce_path)
            if padesce_path.exists():
                filename = padesce_path.name
                message.attach(
                    filename,
                    padesce_path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
                attachments.append(filename)
    except Exception:
        logger.exception("Impossible de joindre le rapport PADESCE du digest quotidien.")

    try:
        cga_filename = get_cga_calls_report_filename()
        message.attach(
            cga_filename,
            build_cga_calls_report_workbook(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        attachments.append(cga_filename)
    except Exception:
        logger.exception("Impossible de joindre le rapport CGA du digest quotidien.")

    try:
        message.send(fail_silently=False)
    except Exception as exc:
        logger.exception("Echec d'envoi du digest quotidien par email.")
        return {
            "ok": False,
            "detail": f"Echec de l'envoi SMTP: {exc}",
            "recipients": resolved_recipients,
            "backup_status": backup_summary.get("status"),
            "attachments": attachments,
        }

    return {
        "ok": True,
        "detail": f"Digest envoye a {len(resolved_recipients)} destinataire(s)",
        "recipients": resolved_recipients,
        "backup_status": backup_summary.get("status"),
        "attachments": attachments,
    }


def build_report_text(report: dict) -> str:
    best_hour = report.get("best_hour")
    best_hour_text = (
        f"Heure la plus performante: {best_hour['label']} ({best_hour['completed']} appels termines)"
        if best_hour
        else "Heure la plus performante: aucune donnee"
    )
    lines = [
        "Rapport quotidien NAUMUR CALL APP",
        f"Periode: {report['start_date']} au {report['end_date']}",
        "",
        "Suivi journee",
        f"- Utilisateurs ayant appele: {report['users']['called_today']}",
        f"- Termines: {report['calls']['completed']}",
        best_hour_text,
        "",
    ]
    if report.get("call_scope") == "cga":
        cga_dimensions = report["cga_dimensions"]
        lines.extend(
            [
                "Analyse CGA",
                f"- Regimes: {cga_dimensions['regimes_count']}",
                f"- Centres: {cga_dimensions['centres_count']}",
                f"- CRI: {cga_dimensions['cris_count']}",
                f"- Villes: {cga_dimensions['villes_count']}",
                f"- Interesses: {cga_dimensions['interesses']}",
                f"- Pas interesses: {cga_dimensions['pas_interesses']}",
                f"- Indisponibles: {cga_dimensions['indisponibles']}",
                f"- Faux numeros: {cga_dimensions['faux_numeros']}",
                "",
            ]
        )
    else:
        lines.extend(
            [
                "Analyse satisfaction",
                f"- Classes analysees: {report['analysis']['classes_count']}",
                f"- Prestations analysees: {report['analysis']['prestations_count']}",
                f"- Prestataires analyses: {report['analysis']['prestataires_count']}",
                f"- Beneficiaires analyses: {report['analysis']['beneficiaires_count']}",
                "",
            ]
        )
    lines.extend(
        [
            "Incidents",
            f"- Bugs signales: {report['bugs']['alarms_total']}",
            f"- Bugs non resolus: {report['bugs']['alarms_open']}",
            f"- Erreurs logs: {report['bugs']['log_errors']}",
            f"- Critiques logs: {report['bugs']['log_critical']}",
        ]
    )
    if report["user_call_rows"]:
        lines.extend(["", "Temps estime par utilisateur"])
        for row in report["user_call_rows"]:
            lines.append(
                f"- {row['username']}: {row['calls_made']} appels, "
                f"{row['completed_calls']} termines, temps estime {row['time_spent_label']}"
            )
    return "\n".join(lines)


def build_report_email_html(report: dict) -> str:
    best_hour_label = report["best_hour"]["label"] if report["best_hour"] else "Aucune donnee"
    best_hour_value = report["best_hour"]["completed"] if report["best_hour"] else 0
    user_rows = "".join(
        [
            (
                "<tr>"
                f"<td style='padding:10px;border-bottom:1px solid #eee7fb;'>{row['username']}</td>"
                f"<td style='padding:10px;border-bottom:1px solid #eee7fb;text-align:center;'>{row['calls_made']}</td>"
                f"<td style='padding:10px;border-bottom:1px solid #eee7fb;text-align:center;'>{row['completed_calls']}</td>"
                f"<td style='padding:10px;border-bottom:1px solid #eee7fb;text-align:center;'>{row['time_spent_label']}</td>"
                "</tr>"
            )
            for row in report["user_call_rows"][:8]
        ]
    ) or (
        "<tr><td colspan='4' style='padding:12px;text-align:center;color:#6b7280;'>"
        "Aucune activite utilisateur sur la periode.</td></tr>"
    )

    if report.get("call_scope") == "cga":
        cga_dimensions = report["cga_dimensions"]
        analysis_html = f"""
      <div style="border:1px solid #dfe1ea;padding:16px 18px;margin-bottom:18px;">
        <div style="font-size:16px;font-weight:700;color:#4c1d95;margin-bottom:10px;">2. Analyse CGA</div>
        <table role="presentation" style="width:100%;border-collapse:collapse;">
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Regimes</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["regimes_count"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Centres</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["centres_count"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">CRI</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["cris_count"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Villes</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["villes_count"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Interesses</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["interesses"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Pas interesses</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["pas_interesses"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Indisponibles</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{cga_dimensions["indisponibles"]}</td></tr>
          <tr><td style="padding:8px 0;">Faux numeros</td><td style="padding:8px 0;text-align:right;font-weight:700;">{cga_dimensions["faux_numeros"]}</td></tr>
        </table>
      </div>
"""
    else:
        analysis_html = f"""
      <div style="border:1px solid #dfe1ea;padding:16px 18px;margin-bottom:18px;">
        <div style="font-size:16px;font-weight:700;color:#4c1d95;margin-bottom:10px;">2. Analyse satisfaction – Apprenants</div>
        <table role="presentation" style="width:100%;border-collapse:collapse;">
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Classes analysees</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["analysis"]["classes_count"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Prestations analysees</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["analysis"]["prestations_count"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Prestataires analyses</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["analysis"]["prestataires_count"]}</td></tr>
          <tr><td style="padding:8px 0;">Beneficiaires analyses</td><td style="padding:8px 0;text-align:right;font-weight:700;">{report["analysis"]["beneficiaires_count"]}</td></tr>
        </table>
      </div>

      <div style="border:1px solid #dfe1ea;padding:16px 18px;margin-bottom:18px;">
        <div style="font-size:16px;font-weight:700;color:#4c1d95;margin-bottom:10px;">3. Analyse satisfaction – Formateurs</div>
        <table role="presentation" style="width:100%;border-collapse:collapse;">
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Appels formateurs termines (periode)</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["formateurs_summary"]["total_termines"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Avec scores Q1-Q3</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["formateurs_summary"]["with_scores"]}</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Moy. Q1 – Prerequis apprenants</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["formateurs_summary"]["avg_q1"]}/5</td></tr>
          <tr><td style="padding:8px 0;border-bottom:1px solid #ececf2;">Moy. Q2 – Interaction apprenants</td><td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["formateurs_summary"]["avg_q2"]}/5</td></tr>
          <tr><td style="padding:8px 0;">Moy. Q3 – Competences acquises</td><td style="padding:8px 0;text-align:right;font-weight:700;">{report["formateurs_summary"]["avg_q3"]}/5</td></tr>
        </table>
      </div>
"""

    return f"""
<div style="margin:0;padding:24px;background:#f4f4f6;font-family:Calibri,Arial,sans-serif;color:#222222;">
  <div style="max-width:880px;margin:0 auto;background:#ffffff;border:1px solid #d9d9e3;">
    <div style="padding:24px 28px;border-bottom:3px solid #4c1d95;background:#ffffff;">
      <table role="presentation" style="width:100%;border-collapse:collapse;">
        <tr>
          <td style="width:120px;vertical-align:top;text-align:left;padding-right:16px;">
            <img src="cid:logo_cga_naumur" alt="Logo CGA Naumur" style="max-width:88px;height:auto;display:block;">
          </td>
          <td style="vertical-align:top;">
            <div style="font-size:12px;letter-spacing:.08em;text-transform:uppercase;color:#4c1d95;font-weight:700;">Rapport journalier</div>
            <h1 style="margin:8px 0 6px;font-size:24px;line-height:1.25;color:#1f1630;">Synthese d'activite</h1>
            <p style="margin:0;font-size:14px;color:#555555;">Periode du {report["start_date"].strftime("%d/%m/%Y")} au {report["end_date"].strftime("%d/%m/%Y")}</p>
            <p style="margin:4px 0 0;font-size:13px;color:#666666;">Genere le {timezone.localtime(report["generated_at"]).strftime("%d/%m/%Y a %H:%M")}</p>
          </td>
        </tr>
      </table>
    </div>

    <div style="padding:24px 28px;">
      <div style="border:1px solid #dfe1ea;padding:16px 18px;margin-bottom:18px;">
        <div style="font-size:16px;font-weight:700;color:#4c1d95;margin-bottom:10px;">1. Resume</div>
        <table role="presentation" style="width:100%;border-collapse:collapse;">
          <tr>
            <td style="padding:8px 0;border-bottom:1px solid #ececf2;">Utilisateurs ayant appele</td>
            <td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["users"]["called_today"]}</td>
          </tr>
          <tr>
            <td style="padding:8px 0;border-bottom:1px solid #ececf2;">Appels termines</td>
            <td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{report["calls"]["completed"]}</td>
          </tr>
          <tr>
            <td style="padding:8px 0;border-bottom:1px solid #ececf2;">Heure la plus performante</td>
            <td style="padding:8px 0;border-bottom:1px solid #ececf2;text-align:right;font-weight:700;">{best_hour_label}</td>
          </tr>
          <tr>
            <td style="padding:8px 0;">Appels termines sur cette tranche</td>
            <td style="padding:8px 0;text-align:right;font-weight:700;">{best_hour_value}</td>
          </tr>
        </table>
      </div>

{analysis_html}

      <div style="border:1px solid #dfe1ea;padding:16px 18px;margin-bottom:18px;">
        <div style="font-size:16px;font-weight:700;color:#4c1d95;margin-bottom:10px;">3. Activite par utilisateur</div>
        <table style="width:100%;border-collapse:collapse;border:1px solid #e3e4eb;">
          <thead>
            <tr style="background:#f5f5fa;color:#3b2a59;">
              <th style="padding:10px;text-align:left;">Utilisateur</th>
              <th style="padding:10px;text-align:center;">Appels effectues</th>
              <th style="padding:10px;text-align:center;">Appels termines</th>
              <th style="padding:10px;text-align:center;">Temps estime</th>
            </tr>
          </thead>
          <tbody>{user_rows}</tbody>
        </table>
      </div>

      <div style="border:1px solid #dfe1ea;padding:16px 18px;">
        <div style="font-size:16px;font-weight:700;color:#4c1d95;margin-bottom:10px;">4. Observation</div>
        <p style="margin:0;font-size:14px;line-height:1.7;color:#333333;">
          La tranche horaire la plus performante est <strong>{best_hour_label}</strong>, avec <strong>{best_hour_value} appels termines</strong>.
          Le document Word joint reprend le rapport sous une forme plus complete pour impression ou partage.
        </p>
      </div>
    </div>
  </div>
</div>
""".strip()


# ---------------------------------------------------------------------------
# Section 6 — Helpers de calcul
# ---------------------------------------------------------------------------


def _source_class_apprenant_counts(source_bundle: dict | None) -> dict[str, int]:
    counts: dict[str, int] = {}
    for record in (source_bundle or {}).get("records", {}).values():
        classe_code = str(record.get("classe_id") or "").strip()
        if not classe_code:
            continue
        counts[classe_code] = counts.get(classe_code, 0) + 1
    return counts


def _qualified_prestation_codes(source_bundle: dict | None) -> set[str]:
    if not source_bundle:
        return set()
    source_classes = list((source_bundle.get("classes") or {}).values())
    if not source_classes:
        return set()

    terminated_by_class = {
        normalize_network_lookup(code): count
        for code, count in (
            Appel.objects.filter(is_active=True, status__in=CALL_ANALYSIS_THRESHOLD_STATUSES)
            .exclude(classe_label="")
            .values("classe_label")
            .annotate(total=Count("id"))
            .values_list("classe_label", "total")
        )
    }
    apprenant_counts = {
        normalize_network_lookup(code): count
        for code, count in _source_class_apprenant_counts(source_bundle).items()
    }

    prestation_classes: dict[str, set[str]] = {}
    for source_class in source_classes:
        prestation_key = normalize_network_lookup(source_class.get("prestation_id", ""))
        classe_key = normalize_network_lookup(source_class.get("classe_id", ""))
        if not prestation_key or not classe_key:
            continue
        if int(apprenant_counts.get(classe_key) or 0) <= 0:
            continue
        prestation_classes.setdefault(prestation_key, set()).add(classe_key)

    qualified_codes: set[str] = set()
    for prestation_key, class_keys in prestation_classes.items():
        all_reached = True
        for class_key in class_keys:
            total_apprenants = int(apprenant_counts.get(class_key) or 0)
            if total_apprenants <= 0:
                all_reached = False
                break
            threshold_target = analysis_threshold_target(total_apprenants)
            total_termines = int(terminated_by_class.get(class_key) or 0)
            if total_termines < threshold_target:
                all_reached = False
                break
        if all_reached:
            qualified_codes.add(prestation_key)
    return qualified_codes


def _get_selected_class(selected_class_code: str | None) -> Classe | None:
    code = (selected_class_code or "").strip()
    if not code:
        return None
    return (
        Classe.objects.filter(actif=True)
        .select_related("prestation__prestataire", "formation")
        .filter(code__iexact=code)
        .first()
    )


def _extract_class_code(classe, classe_label: str | None) -> str:
    if classe and getattr(classe, "code", ""):
        return str(classe.code).strip()
    raw_label = str(classe_label or "").strip()
    if not raw_label:
        return ""
    match = re.match(r"^([A-Za-z0-9_]+)", raw_label)
    return (match.group(1) if match else raw_label.split()[0]).strip()


def _selected_class_code_matches(
    class_code: str, selected_class, selected_class_code: str | None
) -> bool:
    effective_code = (
        (selected_class.code if selected_class else (selected_class_code or "")).strip().lower()
    )
    if not effective_code:
        return True
    return str(class_code or "").strip().lower() == effective_code


def _build_not_formed_rows(
    selected_class=None,
    selected_class_code: str | None = None,
    source_records: dict[str, dict] | None = None,
) -> list[dict]:
    queryset = (
        Appel.objects.filter(is_active=True, flag_pas_forme=True)
        .select_related("classe")
        .order_by("classe_label", "nom")
    )
    rows: list[dict] = []
    for appel in queryset:
        class_code = _extract_class_code(appel.classe, appel.classe_label)
        if not _selected_class_code_matches(class_code, selected_class, selected_class_code):
            continue
        source_record = (source_records or {}).get(normalize_network_lookup(appel.code or "")) or {}
        rows.append(
            _build_anomaly_row(
                appel,
                source_record,
                class_code=class_code,
                anomaly_label="Pas forme",
                anomaly_detail="Appel marque comme apprenant non forme",
            )
        )
    return rows


def _build_false_name_rows(
    selected_class=None,
    selected_class_code: str | None = None,
    source_records: dict[str, dict] | None = None,
) -> list[dict]:
    queryset = (
        Appel.objects.filter(is_active=True, flag_faux_nom=True)
        .select_related("classe")
        .order_by("classe_label", "nom")
    )
    rows: list[dict] = []
    for appel in queryset:
        class_code = _extract_class_code(appel.classe, appel.classe_label)
        if not _selected_class_code_matches(class_code, selected_class, selected_class_code):
            continue
        source_record = (source_records or {}).get(normalize_network_lookup(appel.code or "")) or {}
        rows.append(
            _build_anomaly_row(
                appel,
                source_record,
                class_code=class_code,
                anomaly_label="Faux nom",
                anomaly_detail=appel.nom or "-",
            )
        )
    return rows


def _build_duplicate_phone_rows(
    selected_class=None,
    selected_class_code: str | None = None,
    source_records: dict[str, dict] | None = None,
) -> list[dict]:
    queryset = (
        Appel.objects.filter(is_active=True, flag_numero_double=True)
        .select_related("classe")
        .order_by("classe_label", "nom")
    )
    duplicate_rows: list[dict] = []
    for appel in queryset:
        normalized_phone = _normalize_phone_number(appel.telephone1)
        if not normalized_phone:
            continue
        class_code = _extract_class_code(appel.classe, appel.classe_label)
        if not _selected_class_code_matches(class_code, selected_class, selected_class_code):
            continue
        source_record = (source_records or {}).get(normalize_network_lookup(appel.code or "")) or {}
        duplicate_rows.append(
            _build_anomaly_row(
                appel,
                source_record,
                class_code=class_code,
                anomaly_label="Doublon",
                anomaly_detail=(appel.telephone1 or "").strip() or normalized_phone,
            )
        )
    return duplicate_rows


def _build_anomaly_row(
    appel,
    source_record: dict[str, str] | None,
    *,
    class_code: str,
    anomaly_label: str,
    anomaly_detail: str,
) -> dict:
    source_record = source_record or {}
    phone_display = (appel.telephone1 or "").strip() or "-"
    return {
        "call_id": appel.id,
        "code": appel.code,
        "apprenant_id": source_record.get("apprenant_id") or "",
        "individu_id": source_record.get("individu_id") or "",
        "beneficiaire_id": source_record.get("beneficiaire_id") or "",
        "name": appel.nom or "-",
        "source_name": source_record.get("nom_individu") or "",
        "phone_display": phone_display,
        "phone_normalized": _normalize_phone_number(appel.telephone1),
        "phone_source_1": source_record.get("telephone1") or "",
        "phone_source_2": source_record.get("telephone2") or "",
        "class_code": class_code,
        "class_label": _class_display_label(appel.classe, appel.classe_label),
        "prestation_id": source_record.get("prestation_id") or "",
        "prestataire": source_record.get("prestataire") or "",
        "beneficiaire": source_record.get("beneficiaire") or "",
        "formation": source_record.get("formation") or "",
        "cohorte": source_record.get("cohorte") or "",
        "statut_apprenant": source_record.get("statut_apprenant") or "",
        "fenetre": source_record.get("fenetre") or "",
        "lieu": source_record.get("lieu") or "",
        "ville": source_record.get("ville") or "",
        "arrondissement": source_record.get("arrondissement") or "",
        "departement": source_record.get("departement") or "",
        "region": source_record.get("region") or "",
        "statut_prestation": source_record.get("statut_prestation") or "",
        "anomaly_label": anomaly_label,
        "anomaly_detail": anomaly_detail,
    }


def _build_anomaly_summary(selected_class=None, selected_class_code: str | None = None) -> dict:
    classes_qs = Classe.objects.filter(actif=True)
    classes_not_trained_count = classes_qs.exclude(statut="termine").count()

    try:
        source_records = (build_padesce_source_index() or {}).get("records", {})
    except Exception:
        source_records = {}

    not_formed_rows_all = _build_not_formed_rows(source_records=source_records)
    false_name_rows_all = _build_false_name_rows(source_records=source_records)
    duplicate_rows_all = _build_duplicate_phone_rows(source_records=source_records)

    selected_not_formed = _build_not_formed_rows(
        selected_class,
        selected_class_code,
        source_records=source_records,
    )
    selected_false_name = _build_false_name_rows(
        selected_class,
        selected_class_code,
        source_records=source_records,
    )
    selected_duplicate = _build_duplicate_phone_rows(
        selected_class,
        selected_class_code,
        source_records=source_records,
    )

    anomaly_options: dict[str, dict] = {}
    for row in not_formed_rows_all + false_name_rows_all + duplicate_rows_all:
        code = row["class_code"]
        if not code:
            continue
        if code not in anomaly_options:
            anomaly_options[code] = {
                "code": code,
                "label": row["class_label"],
                "has_not_formed": False,
                "has_duplicate": False,
                "has_false_name": False,
            }
        if row["anomaly_label"] == "Pas forme":
            anomaly_options[code]["has_not_formed"] = True
        if row["anomaly_label"] == "Doublon":
            anomaly_options[code]["has_duplicate"] = True
        if row["anomaly_label"] == "Faux nom":
            anomaly_options[code]["has_false_name"] = True

    anomaly_rows = sorted(
        selected_duplicate + selected_false_name + selected_not_formed,
        key=lambda row: (
            {"Doublon": 0, "Faux nom": 1, "Pas forme": 2}.get(row["anomaly_label"], 9),
            row["class_label"],
            row["code"],
            row["call_id"],
        ),
    )

    effective_selected_code = (
        selected_class.code if selected_class else (selected_class_code or "")
    ).strip()
    selected_class_summary = {
        "code": effective_selected_code,
        "label": _class_display_label(selected_class, effective_selected_code),
        "not_formed_count": len(selected_not_formed),
        "duplicate_numbers_count": len(selected_duplicate),
        "false_name_count": len(selected_false_name),
        "total_anomalies_count": len(selected_not_formed)
        + len(selected_duplicate)
        + len(selected_false_name),
    }

    return {
        "classes_not_trained_count": classes_not_trained_count,
        "duplicate_numbers_count": len(duplicate_rows_all),
        "duplicate_rows": selected_duplicate,
        "not_formed_rows": selected_not_formed,
        "false_name_rows": selected_false_name,
        "anomaly_rows": anomaly_rows,
        "selected_class": selected_class_summary,
        "class_options": sorted(anomaly_options.values(), key=lambda item: item["code"]),
    }


def _class_display_label(classe, label: str | None) -> str:
    if classe:
        return f"{classe.code} - {classe.intitule_formation}"
    return str(label or "-").strip()


def _normalize_phone_number(val) -> str:
    if not val:
        return ""
    digits = "".join(filter(str.isdigit, str(val)))
    if len(digits) == 9 and digits.startswith(("6", "2")):
        return f"237{digits}"
    return digits


def _build_analysis_summary(padesce_qs) -> dict:
    try:
        from App_PADESCE.satisfaction_apprenants.views import _build_satisfaction_dashboard_data

        query = QueryDict("", mutable=True)
        query["source"] = "cutoff"
        payload = _build_satisfaction_dashboard_data(SimpleNamespace(GET=query))
        context = payload["context"]
        return {
            "classes_count": context["analyzed_classes_count"],
            "prestations_count": context["analyzed_prestations_count"],
            "prestataires_count": context["analyzed_prestataires_count"],
            "beneficiaires_count": context["analyzed_beneficiaires_count"],
        }
    except Exception:
        pass

    terminated_rows = list(
        padesce_qs.filter(status__in=CALL_COMPLETED_STATUSES).select_related("classe__prestation")
    )
    classes = sorted(
        {
            (row.classe_label or "").strip()
            for row in terminated_rows
            if (row.classe_label or "").strip()
        }
    )
    prestataires = sorted(
        {
            (row.prestataire or "").strip()
            for row in terminated_rows
            if (row.prestataire or "").strip()
        }
    )
    beneficiaires = sorted(
        {
            (row.beneficiaire or "").strip()
            for row in terminated_rows
            if (row.beneficiaire or "").strip()
        }
    )
    try:
        source_bundle = build_padesce_source_index()
    except Exception:
        source_bundle = None
    qualified_codes = _qualified_prestation_codes(source_bundle)
    prestations = sorted(
        {
            normalize_network_lookup(getattr(row.classe.prestation, "code", ""))
            for row in terminated_rows
            if (
                getattr(row, "classe", None)
                and getattr(row.classe, "prestation", None)
                and getattr(row.classe.prestation, "code", "")
                and normalize_network_lookup(row.classe.prestation.code) in qualified_codes
            )
        }
    )
    return {
        "classes_count": len(classes),
        "prestations_count": len(prestations),
        "prestataires_count": len(prestataires),
        "beneficiaires_count": len(beneficiaires),
    }


def _build_formateurs_satisfaction_summary(start_dt, end_dt) -> dict:
    """Résumé satisfaction formateurs (Q1-Q3) pour le rapport quotidien."""
    from django.db.models import Avg as DjAvg

    qs = AppelFormateur.objects.filter(
        is_active=True,
        status__in=CALL_COMPLETED_STATUSES,
        updated_at__gte=start_dt,
        updated_at__lte=end_dt,
    )
    total = qs.count()
    with_scores = qs.filter(
        q1_prerequis_apprenants__isnull=False,
        q2_interaction_apprenants__isnull=False,
        q3_competences_acquises__isnull=False,
    ).count()

    avgs = qs.aggregate(
        avg_q1=DjAvg("q1_prerequis_apprenants"),
        avg_q2=DjAvg("q2_interaction_apprenants"),
        avg_q3=DjAvg("q3_competences_acquises"),
    )

    # Par prestataire
    by_prestataire = list(
        qs.values("prestataire")
        .annotate(
            nb=Count("id"),
            avg_q1=DjAvg("q1_prerequis_apprenants"),
            avg_q2=DjAvg("q2_interaction_apprenants"),
            avg_q3=DjAvg("q3_competences_acquises"),
        )
        .order_by("prestataire")[:10]
    )

    return {
        "total_termines": total,
        "with_scores": with_scores,
        "avg_q1": round(avgs["avg_q1"] or 0, 2),
        "avg_q2": round(avgs["avg_q2"] or 0, 2),
        "avg_q3": round(avgs["avg_q3"] or 0, 2),
        "by_prestataire": [
            {
                "prestataire": row["prestataire"] or "—",
                "nb": row["nb"],
                "avg_q1": round(row["avg_q1"] or 0, 2),
                "avg_q2": round(row["avg_q2"] or 0, 2),
                "avg_q3": round(row["avg_q3"] or 0, 2),
            }
            for row in by_prestataire
        ],
    }


def _build_user_call_rows(*querysets, day_start=None, day_end=None) -> list[dict]:
    day_filter = Q()
    has_day_window = day_start is not None or day_end is not None
    if day_start is not None:
        day_filter &= Q(updated_at__gte=day_start)
    if day_end is not None:
        day_filter &= Q(updated_at__lte=day_end)

    merged: dict[int, dict] = {}
    for queryset in querysets:
        rows = (
            queryset.exclude(status="en_attente")
            .filter(locked_by__isnull=False)
            .values("locked_by_id", "locked_by__username")
            .annotate(
                calls_made=Count("id"),
                calls_today=Count("id", filter=day_filter) if has_day_window else Count("id"),
                completed_calls=Count("id", filter=Q(status__in=CALL_COMPLETED_STATUSES)),
                first_activity=Min("updated_at"),
                last_activity=Max("updated_at"),
            )
        )
        for row in rows:
            user_id = row["locked_by_id"]
            if user_id not in merged:
                merged[user_id] = {
                    "username": row["locked_by__username"] or f"user-{user_id}",
                    "calls_made": 0,
                    "calls_today": 0,
                    "completed_calls": 0,
                    "first_activity": row["first_activity"],
                    "last_activity": row["last_activity"],
                }
            merged[user_id]["calls_made"] += int(row["calls_made"] or 0)
            merged[user_id]["calls_today"] += int(row["calls_today"] or 0)
            merged[user_id]["completed_calls"] += int(row["completed_calls"] or 0)
            if row["first_activity"] and (
                not merged[user_id]["first_activity"]
                or row["first_activity"] < merged[user_id]["first_activity"]
            ):
                merged[user_id]["first_activity"] = row["first_activity"]
            if row["last_activity"] and (
                not merged[user_id]["last_activity"]
                or row["last_activity"] > merged[user_id]["last_activity"]
            ):
                merged[user_id]["last_activity"] = row["last_activity"]

    results = []
    for item in merged.values():
        first_activity = item["first_activity"]
        last_activity = item["last_activity"]
        duration = (
            (last_activity - first_activity) if (first_activity and last_activity) else timedelta(0)
        )
        results.append(
            {
                "username": item["username"],
                "calls_made": item["calls_made"],
                "calls_today": item["calls_today"],
                "completed_calls": item["completed_calls"],
                "time_spent_seconds": int(max(duration.total_seconds(), 0)),
                "time_spent_label": _format_duration(duration),
                "first_activity_label": (
                    timezone.localtime(first_activity).strftime("%H:%M") if first_activity else "-"
                ),
                "last_activity_label": (
                    timezone.localtime(last_activity).strftime("%H:%M") if last_activity else "-"
                ),
            }
        )
    return sorted(results, key=lambda row: (-row["calls_made"], row["username"].lower()))


def _format_duration(duration) -> str:
    total_seconds = int(max(duration.total_seconds(), 0))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    if hours:
        return f"{hours}h {minutes:02d}min"
    if minutes:
        return f"{minutes}min {seconds:02d}s"
    return f"{seconds}s"


def _email_card(label: str, value) -> str:
    return (
        "<div style='flex:1 1 180px;min-width:180px;padding:16px 18px;border-radius:18px;"
        "background:#ffffff;border:1px solid #ece5fb;box-shadow:0 10px 24px rgba(76,29,149,0.08);'>"
        f"<div style='font-size:13px;color:#7a67a5;margin-bottom:8px;'>{label}</div>"
        f"<div style='font-size:28px;font-weight:800;color:#4c1d95;line-height:1.2;'>{value}</div>"
        "</div>"
    )


def _email_stat(label: str, value) -> str:
    return (
        "<div style='padding:12px 14px;border-radius:14px;background:#ffffff;border:1px solid #ece5fb;'>"
        f"<div style='font-size:12px;color:#7a67a5;margin-bottom:5px;'>{label}</div>"
        f"<div style='font-size:20px;font-weight:800;color:#4c1d95;'>{value}</div>"
        "</div>"
    )


def _get_report_logo_path() -> Path | None:
    candidates = [
        Path(settings.BASE_DIR) / "LogoCGANaumur.png",
        Path(settings.BASE_DIR) / "media" / "LogoCGANaumur.png",
        Path(settings.BASE_DIR) / "media" / "img" / "LogoCGANaumur.png",
        Path(settings.BASE_DIR) / "static" / "LogoCGANaumur.png",
    ]
    for path in candidates:
        if path.exists():
            return path
    return None


def _clear_cell_border(cell) -> None:
    if not HAS_DOCX:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        tc_borders.append(element)


# ---------------------------------------------------------------------------
# Helpers internes — construction du rapport
# ---------------------------------------------------------------------------


def _build_call_source_summary(label: str, queryset) -> dict:
    total = queryset.count()
    completed = queryset.filter(status__in=CALL_COMPLETED_STATUSES).count()
    pending = queryset.filter(status="en_attente").count()
    in_progress = queryset.filter(status__in=CALL_TENTATIVE_STATUSES).count()
    callbacks = queryset.filter(status="a_rappeler").count()
    processed = total - pending
    with_audio = queryset.exclude(audio_file="").exclude(audio_file__isnull=True).count()
    return {
        "label": label,
        "total": total,
        "completed": completed,
        "pending": pending,
        "in_progress": in_progress,
        "callbacks": callbacks,
        "processed": processed,
        "with_audio": with_audio,
    }


def _count_distinct_non_empty(queryset, field_name: str) -> int:
    return queryset.exclude(**{field_name: ""}).values(field_name).distinct().count()


def _build_cga_dimension_summary(cga_qs) -> dict:
    return {
        "regimes_count": _count_distinct_non_empty(cga_qs, "regime"),
        "centres_count": _count_distinct_non_empty(cga_qs, "centre_de_rattachement"),
        "cris_count": _count_distinct_non_empty(cga_qs, "cri"),
        "villes_count": _count_distinct_non_empty(cga_qs, "ville"),
        "interesses": cga_qs.filter(interet="OUI").count(),
        "pas_interesses": cga_qs.filter(interet="NON").count(),
        "indisponibles": cga_qs.filter(indisponible="OUI").count(),
        "faux_numeros": cga_qs.filter(mauvais_numero="OUI").count(),
    }


def _build_cga_anomaly_summary(cga_qs) -> dict:
    rows = []
    for appel in (
        cga_qs.filter(mauvais_numero="OUI")
        .select_related("locked_by")
        .order_by("-updated_at", "raison_sociale")
    ):
        rows.append(
            {
                "anomaly_label": "Faux numero",
                "anomaly_detail": "Faux numero CGA",
                "niu": appel.niu,
                "raison_sociale": appel.raison_sociale,
                "telephone": appel.telephone,
                "regime": appel.regime,
                "centre": appel.centre_de_rattachement,
                "cri": appel.cri,
                "ville": appel.ville,
                "username": appel.locked_by.get_username() if appel.locked_by else "",
                "updated_at": appel.updated_at,
                "updated_at_label": (
                    timezone.localtime(appel.updated_at).strftime("%d/%m/%Y %H:%M")
                    if appel.updated_at
                    else "-"
                ),
            }
        )
    total = len(rows)
    return {
        "type": "cga",
        "false_number_count": total,
        "anomaly_rows": rows,
        "class_options": [],
        "selected_class": {
            "not_formed_count": 0,
            "duplicate_numbers_count": 0,
            "false_name_count": 0,
            "total_anomalies_count": total,
        },
    }


def _hourly_completed_counts(queryset) -> dict[int, int]:
    """Compte les appels termines par heure (compatible SQLite et PostgreSQL)."""
    counts: dict[int, int] = {}
    tz = timezone.get_current_timezone()
    for updated_at in queryset.filter(status__in=CALL_COMPLETED_STATUSES).values_list(
        "updated_at", flat=True
    ):
        if updated_at is None:
            continue
        local_dt = (
            timezone.localtime(updated_at, tz) if timezone.is_aware(updated_at) else updated_at
        )
        hour = local_dt.hour
        counts[hour] = counts.get(hour, 0) + 1
    return counts


def _build_hourly_rows(*hourly_counts_list) -> list[dict]:
    merged: dict[int, dict] = {}
    for hourly_counts in hourly_counts_list:
        for hour, count in hourly_counts.items():
            if hour not in merged:
                merged[hour] = {
                    "hour": hour,
                    "label": f"{hour:02d}h00 - {hour:02d}h59",
                    "total": 0,
                    "completed": 0,
                }
            merged[hour]["completed"] += count
            merged[hour]["total"] += count
    return sorted(merged.values(), key=lambda row: row["hour"])


def _build_bug_summary(start_dt, end_dt) -> dict:
    """Résumé des bugs/alarmes sur la période."""
    try:
        from App_PADESCE.core.models import Alarm

        alarms_qs = Alarm.objects.filter(created_at__gte=start_dt, created_at__lte=end_dt)
        alarms_total = alarms_qs.count()
        alarms_resolved = alarms_qs.filter(status="resolved").count()
        alarms_open = alarms_total - alarms_resolved
        recent_alarms = [
            {
                "title": str(alarm.title or ""),
                "module": str(getattr(alarm, "module", "") or ""),
                "created_at": timezone.localtime(alarm.created_at).strftime("%d/%m/%Y %H:%M"),
                "status": str(alarm.status or ""),
            }
            for alarm in alarms_qs.order_by("-created_at")[:10]
        ]
    except Exception:
        alarms_total = 0
        alarms_resolved = 0
        alarms_open = 0
        recent_alarms = []

    log_errors = _count_log_level("ERROR", start_dt, end_dt)
    log_critical = _count_log_level("CRITICAL", start_dt, end_dt)

    return {
        "alarms_total": alarms_total,
        "alarms_resolved": alarms_resolved,
        "alarms_open": alarms_open,
        "log_errors": log_errors,
        "log_critical": log_critical,
        "recent_alarms": recent_alarms,
    }


def _count_log_level(level: str, start_dt, end_dt) -> int:
    """Compte les lignes de log contenant le niveau donné sur la période."""
    try:
        log_path = Path(settings.BASE_DIR) / "logs" / "app.log"
        if not log_path.exists():
            return 0
        start_str = timezone.localtime(start_dt).strftime("%Y-%m-%d")
        end_str = timezone.localtime(end_dt).strftime("%Y-%m-%d")
        count = 0
        with log_path.open("r", encoding="utf-8", errors="replace") as f:
            for line in f:
                if level in line and (start_str <= line[:10] <= end_str):
                    count += 1
        return count
    except Exception:
        return 0


# ---------------------------------------------------------------------------
# Helpers Word — document
# ---------------------------------------------------------------------------


def _configure_report_document(document) -> None:
    if not HAS_DOCX:
        return
    from docx.shared import Pt as _Pt

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = _Pt(11)


def _add_heading(document, text: str) -> None:
    if not HAS_DOCX:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(76, 29, 149)
    paragraph.space_before = Pt(14)
    paragraph.space_after = Pt(6)


def _add_metric_table(document, rows: list[tuple]) -> None:
    if not HAS_DOCX:
        return
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = str(value)


def _write_header_cell(cell, text: str) -> None:
    if not HAS_DOCX:
        return
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(76, 29, 149)


# ---------------------------------------------------------------------------
# Exports CSV / Excel
# ---------------------------------------------------------------------------


ANOMALY_EXPORT_COLUMNS = [
    ("anomaly_label", "Anomalie"),
    ("anomaly_detail", "Detail anomalie"),
    ("apprenant_id", "ApprenantID"),
    ("individu_id", "IndividuID"),
    ("beneficiaire_id", "BeneficiaireID"),
    ("code", "Code appel"),
    ("name", "Nom sur appel"),
    ("source_name", "Nom apprenant source"),
    ("class_code", "Classe"),
    ("class_label", "Classe libelle"),
    ("prestation_id", "Prestation ID"),
    ("prestataire", "Prestataire"),
    ("beneficiaire", "Beneficiaire"),
    ("formation", "Formation"),
    ("cohorte", "Cohorte"),
    ("statut_apprenant", "Statut apprenant"),
    ("phone_display", "Telephone appel"),
    ("phone_source_1", "Telephone source 1"),
    ("phone_source_2", "Telephone source 2"),
    ("fenetre", "Fenetre"),
    ("lieu", "Lieu"),
    ("ville", "Ville"),
    ("arrondissement", "Arrondissement"),
    ("departement", "Departement"),
    ("region", "Region"),
    ("statut_prestation", "Statut prestation"),
]

CGA_ANOMALY_EXPORT_COLUMNS = [
    ("anomaly_label", "Anomalie"),
    ("niu", "NIU"),
    ("raison_sociale", "Raison sociale"),
    ("telephone", "Telephone"),
    ("regime", "Regime"),
    ("centre", "Centre"),
    ("cri", "CRI"),
    ("ville", "Ville"),
    ("username", "Utilisateur"),
    ("updated_at_label", "Date"),
]


def export_application_report_anomalies_excel(report: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Anomalies"

    columns = (
        CGA_ANOMALY_EXPORT_COLUMNS
        if report.get("call_scope") == "cga"
        else ANOMALY_EXPORT_COLUMNS
    )
    worksheet.append([label for _key, label in columns])
    for cell in worksheet[1]:
        cell.font = Font(bold=True, color="4C1D95")

    for row in report.get("anomalies", {}).get("anomaly_rows", []):
        worksheet.append([row.get(key, "") for key, _label in columns])

    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def export_application_report_csv(report: dict) -> bytes:
    import csv
    import io as _io

    output = _io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Section", "Indicateur", "Valeur"])
    writer.writerow(["Appels", "Termines", report["calls"]["completed"]])
    writer.writerow(["Appels", "Effectues", report["calls"]["processed"]])
    writer.writerow(["Appels", "En attente", report["calls"]["pending"]])
    writer.writerow(["Appels", "En cours", report["calls"]["in_progress"]])
    writer.writerow(["Appels", "A rappeler", report["calls"]["callbacks"]])
    writer.writerow(["Appels", "Avec audio", report["calls"]["with_audio"]])
    writer.writerow(["Utilisateurs", "Ayant appele", report["users"]["called_today"]])
    if report.get("call_scope") == "cga":
        cga_dimensions = report["cga_dimensions"]
        writer.writerow(["Analyse CGA", "Regimes", cga_dimensions["regimes_count"]])
        writer.writerow(["Analyse CGA", "Centres", cga_dimensions["centres_count"]])
        writer.writerow(["Analyse CGA", "CRI", cga_dimensions["cris_count"]])
        writer.writerow(["Analyse CGA", "Villes", cga_dimensions["villes_count"]])
        writer.writerow(["Analyse CGA", "Interesses", cga_dimensions["interesses"]])
        writer.writerow(["Analyse CGA", "Pas interesses", cga_dimensions["pas_interesses"]])
        writer.writerow(["Analyse CGA", "Indisponibles", cga_dimensions["indisponibles"]])
        writer.writerow(["Analyse CGA", "Faux numeros", cga_dimensions["faux_numeros"]])
    else:
        writer.writerow(["Analyse", "Classes", report["analysis"]["classes_count"]])
        writer.writerow(["Analyse", "Prestations", report["analysis"]["prestations_count"]])
        writer.writerow(["Analyse", "Prestataires", report["analysis"]["prestataires_count"]])
        writer.writerow(["Analyse", "Beneficiaires", report["analysis"]["beneficiaires_count"]])
    for row in report["user_call_rows"]:
        writer.writerow(
            [
                "Utilisateur",
                row["username"],
                f"{row['calls_made']} appels, {row['completed_calls']} termines, {row['time_spent_label']}",
            ]
        )
    return output.getvalue().encode("utf-8-sig")


def export_application_report_excel(report: dict) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        ws = wb.active
        ws.title = "Rapport"

        def add_section(title, rows_data):
            ws.append([title])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=13, color="4C1D95")
            for row in rows_data:
                ws.append(list(row))

        add_section(
            "Appels",
            [
                ("Termines", report["calls"]["completed"]),
                ("Effectues", report["calls"]["processed"]),
                ("En attente", report["calls"]["pending"]),
                ("En cours", report["calls"]["in_progress"]),
                ("A rappeler", report["calls"]["callbacks"]),
                ("Avec audio", report["calls"]["with_audio"]),
            ],
        )
        ws.append([])
        if report.get("call_scope") == "cga":
            cga_dimensions = report["cga_dimensions"]
            add_section(
                "Analyse CGA",
                [
                    ("Regimes", cga_dimensions["regimes_count"]),
                    ("Centres", cga_dimensions["centres_count"]),
                    ("CRI", cga_dimensions["cris_count"]),
                    ("Villes", cga_dimensions["villes_count"]),
                    ("Interesses", cga_dimensions["interesses"]),
                    ("Pas interesses", cga_dimensions["pas_interesses"]),
                    ("Indisponibles", cga_dimensions["indisponibles"]),
                    ("Faux numeros", cga_dimensions["faux_numeros"]),
                ],
            )
        else:
            add_section(
                "Analyse satisfaction",
                [
                    ("Classes", report["analysis"]["classes_count"]),
                    ("Prestations", report["analysis"]["prestations_count"]),
                    ("Prestataires", report["analysis"]["prestataires_count"]),
                    ("Beneficiaires", report["analysis"]["beneficiaires_count"]),
                ],
            )
        ws.append([])
        add_section(
            "Activite utilisateurs",
            [
                (
                    row["username"],
                    row["calls_made"],
                    row["completed_calls"],
                    row["time_spent_label"],
                )
                for row in report["user_call_rows"]
            ],
        )

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
    except ImportError:
        return export_application_report_csv(report)
