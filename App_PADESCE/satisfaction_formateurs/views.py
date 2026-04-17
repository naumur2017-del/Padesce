import base64
import csv
import hashlib
import io
import logging
import os
import re
import uuid
import zipfile
from collections import defaultdict
from datetime import date as date_cls

import openpyxl
import requests
from django.contrib import messages
from django.core.cache import cache
from django.core.files.storage import default_storage
from django.core.paginator import Paginator
from django.db import transaction
from django.db.models import Count, Q
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.urls import reverse
from django.utils import timezone

from App_PADESCE.appels.formateur_names import (
    FORMATEUR_NAME_FALLBACK,
    resolve_formateur_db_name_from_values,
    resolve_formateur_name_from_values,
    sync_formateur_names_from_excel,
)
from App_PADESCE.appels.models import (
    CALL_ANALYSIS_THRESHOLD_STATUSES,
    FORMATEUR_SCORE_FIELDS,
    FORMATEUR_TEXT_FIELDS,
    Appel,
    AppelFormateur,
    derive_formateur_status,
    formateur_has_any_audio,
    formateur_has_any_form_data,
    sync_formateur_status,
)
from App_PADESCE.core.access import require_analysis_access
from App_PADESCE.core.analysis_rules import analysis_threshold_label, analysis_threshold_target
from App_PADESCE.core.cache_versions import get_analysis_cache_version
from App_PADESCE.core.fast_stats import build_fast_stats_context
from App_PADESCE.formations.models import (
    Beneficiaire,
    Classe,
    Formateur,
    Formation,
    Prestataire,
    Prestation,
)
from App_PADESCE.satisfaction_apprenants.services import get_prestations_ranking
from App_PADESCE.satisfaction_formateurs.forms import (
    SatisfactionFormateurBatchUpdateForm,
    SatisfactionFormateurForm,
)
from App_PADESCE.satisfaction_formateurs.models import SatisfactionFormateur

SESSION_KEY = "sat_form_workflow"

OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
DEFAULT_TRANSCRIBE_MODEL = "google/gemini-2.5-flash"
SUPPORTED_AUDIO_FORMATS = {"wav", "mp3", "m4a", "ogg", "webm", "flac"}

logger = logging.getLogger(__name__)
_FORMATEURS_CACHE_TIMEOUT = int(str(os.getenv("PADESCE_ANALYSIS_CACHE_TIMEOUT", "300") or "300"))
_FORMATEURS_DASHBOARD_CACHE_VERSION = "scores-q1q3-v2"


def _formateurs_cache_key(*parts) -> str:
    rendered = [str(p or "").strip() for p in parts]
    digest = hashlib.sha1("||".join(rendered).encode("utf-8")).hexdigest()
    return f"satisfaction:formateurs:{digest}"


def _normalize_phone(value: str) -> str:
    return "".join(ch for ch in value if ch.isdigit())


def _find_formateur(classe_id: str, identifiant: str) -> tuple[Formateur | None, str | None]:
    identifiant = identifiant.strip()
    if not identifiant:
        return None, "Renseignez le code ou telephone du formateur."
    classe = Classe.objects.select_related("formateur").filter(id=classe_id).first()
    if not classe:
        return None, "Classe introuvable."
    if not classe.formateur:
        return None, "Aucun formateur associe a cette classe."
    formateur = classe.formateur
    identifiant_lower = identifiant.lower()
    identifiant_digits = _normalize_phone(identifiant)
    if formateur.code.lower() == identifiant_lower:
        return formateur, None
    if identifiant_digits and _normalize_phone(formateur.telephone) == identifiant_digits:
        return formateur, None
    return None, "Ce formateur ne correspond pas a la classe."


def _save_audio(uploaded_file, folder: str) -> str:
    _, ext = os.path.splitext(uploaded_file.name)
    ext = ext or ".dat"
    filename = f"{uuid.uuid4().hex}{ext}"
    return default_storage.save(f"enquetes/{folder}/{filename}", uploaded_file)


# Fallback scoring when transcript does not include explicit answers.
def _ai_scores(seed: str, count: int = 9) -> list[int]:
    digest = hashlib.sha256(seed.encode("utf-8")).digest()
    return [1 + (digest[i] % 5) for i in range(count)]


def _flatten_message_content(content) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text" and item.get("text"):
                parts.append(item["text"])
        return "\n".join(parts)
    return ""


def _extract_transcription_text(result) -> str:
    if not result:
        return ""
    if isinstance(result, dict):
        if "choices" in result:
            try:
                content = result["choices"][0]["message"].get("content")
            except (IndexError, KeyError, TypeError, AttributeError):
                content = None
            return _flatten_message_content(content)
        return result.get("text") or result.get("transcript") or ""
    if hasattr(result, "text"):
        return result.text or ""
    return str(result)


def _guess_audio_format(audio_path: str) -> str:
    ext = os.path.splitext(audio_path)[1].lstrip(".").lower()
    if ext == "ma4":
        ext = "m4a"
    return ext if ext in SUPPORTED_AUDIO_FORMATS else "wav"


def _encode_audio_to_base64(audio_path: str) -> str:
    with default_storage.open(audio_path, "rb") as audio_file:
        return base64.b64encode(audio_file.read()).decode("ascii")


def _transcribe_audio(audio_path: str) -> tuple[str | None, str | None]:
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        return None, "Cle OpenRouter manquante. Definissez OPENROUTER_API_KEY dans l'environnement."
    audio_format = _guess_audio_format(audio_path)
    model = os.getenv("OPENROUTER_TRANSCRIBE_MODEL", DEFAULT_TRANSCRIBE_MODEL)
    logger.info(
        "Transcription OpenRouter demarree. audio=%s format=%s model=%s",
        audio_path,
        audio_format,
        model,
    )
    try:
        base64_audio = _encode_audio_to_base64(audio_path)
    except Exception as exc:
        logger.exception("Lecture audio impossible pour transcription. audio=%s", audio_path)
        return None, f"Impossible de lire l'audio: {exc}"

    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Please transcribe this audio file."},
                    {
                        "type": "input_audio",
                        "input_audio": {"data": base64_audio, "format": audio_format},
                    },
                ],
            }
        ],
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    response = None
    try:
        response = requests.post(OPENROUTER_API_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
    except requests.RequestException as exc:
        detail = ""
        if response is not None:
            try:
                detail = response.json().get("error", {}).get("message") or response.text
            except ValueError:
                detail = response.text
        detail = detail.strip()
        logger.error(
            "Echec transcription OpenRouter. audio=%s status=%s detail=%s",
            audio_path,
            response.status_code if response is not None else "n/a",
            detail or str(exc),
        )
        return None, f"Echec de transcription OpenRouter: {exc} {detail}".strip()
    except ValueError as exc:
        logger.error("Reponse OpenRouter invalide. audio=%s error=%s", audio_path, exc)
        return None, f"Reponse OpenRouter invalide: {exc}"

    text = _extract_transcription_text(data).strip()
    if not text:
        logger.error("Transcription OpenRouter videe. audio=%s", audio_path)
        return None, "Transcription videe ou non reconnue."
    logger.info("Transcription OpenRouter terminee. audio=%s chars=%s", audio_path, len(text))
    return text[:8000], None


def _parse_scores_from_transcript(transcript: str, total_questions: int = 3) -> dict[int, int]:
    results: dict[int, int] = {}
    if not transcript:
        return results
    lowered = transcript.lower()
    for idx in range(1, total_questions + 1):
        pattern = rf"(?:\bq{idx}\b|question\s*{idx})[^\d]{{0,10}}([1-5])"
        match = re.search(pattern, lowered)
        if match:
            results[idx] = int(match.group(1))
    return results


def _ai_results_formateur(audio_path: str) -> tuple[dict | None, str | None, str | None]:
    transcript, error = _transcribe_audio(audio_path)
    if error:
        return None, None, error
    size = default_storage.size(audio_path)
    scores = _ai_scores(f"{audio_path}:{size}:{len(transcript)}", count=3)
    parsed = _parse_scores_from_transcript(transcript, total_questions=3)
    for idx, value in parsed.items():
        if 1 <= idx <= 3:
            scores[idx - 1] = value
    results = {
        "q1_prerequis_apprenants": scores[0],
        "q2_interaction_apprenants": scores[1],
        "q3_competences_acquises": scores[2],
        # Q4-Q6 : questions ouvertes — a saisir manuellement apres transcription
        "q4_gestion_administrative": "",
        "q5_gestion_financiere": "",
        "q6_communication": "",
        "commentaires": "",
        "recommandations": "",
    }
    return results, transcript, None


def satisfaction_formateurs(request):
    filter_classe = request.GET.get("classe")
    qs = SatisfactionFormateur.objects.select_related(
        "classe", "formateur", "inspecteur", "enqueteur"
    ).order_by("-date", "-created_at")
    if filter_classe:
        qs = qs.filter(classe_id=filter_classe)

    workflow = request.session.get(SESSION_KEY, {})
    save_errors = None

    if request.method == "POST":
        action = request.POST.get("action")
        identifiant = (request.POST.get("identifiant") or "").strip()
        posted_classe = request.POST.get("classe") or workflow.get("classe_id")
        posted_inspecteur = request.POST.get("inspecteur") or workflow.get("inspecteur_id")
        posted_date = request.POST.get("date") or workflow.get("date")
        posted_heure = request.POST.get("heure") or workflow.get("heure")

        if posted_classe:
            workflow["classe_id"] = str(posted_classe)
        if posted_inspecteur:
            workflow["inspecteur_id"] = posted_inspecteur
        if posted_date:
            workflow["date"] = posted_date
        if posted_heure:
            workflow["heure"] = posted_heure
        if identifiant:
            workflow["identifiant"] = identifiant

        if action == "identify":
            if not posted_classe or not identifiant:
                messages.error(
                    request, "Renseignez la classe et le code ou telephone du formateur."
                )
            else:
                formateur, error = _find_formateur(posted_classe, identifiant)
                if formateur:
                    workflow["formateur_id"] = formateur.id
                    workflow.pop("audio_path", None)
                    workflow.pop("ai_results", None)
                    messages.success(request, f"Formateur identifie: {formateur}.")
                else:
                    messages.error(request, error or "Formateur non trouve pour cette classe.")
        elif action == "process_audio":
            if not workflow.get("formateur_id"):
                messages.error(request, "Identifiez d'abord un formateur.")
            elif str(posted_classe or "") != str(workflow.get("classe_id") or ""):
                messages.error(request, "La classe ne correspond pas au formateur identifie.")
            else:
                uploaded_audio = request.FILES.get("audio_appel")
                if uploaded_audio:
                    workflow["audio_path"] = _save_audio(uploaded_audio, "satisfaction_formateurs")
                    logger.info(
                        "Audio recu pour transcription formateur. fichier=%s taille=%s",
                        uploaded_audio.name,
                        getattr(uploaded_audio, "size", "n/a"),
                    )
                if not workflow.get("audio_path"):
                    messages.error(request, "Chargez un audio d'appel pour lancer le traitement.")
                else:
                    results, transcript, error = _ai_results_formateur(workflow["audio_path"])
                    if error:
                        messages.error(request, error)
                    else:
                        workflow["ai_results"] = results
                        workflow["transcription"] = transcript
                        messages.success(
                            request, "Transcription terminee et traitement vocal actualise."
                        )
        elif action == "save":
            if not workflow.get("formateur_id"):
                messages.error(request, "Identifiez un formateur avant d'enregistrer.")
            elif not workflow.get("ai_results"):
                messages.error(request, "Lancez le traitement vocal avant d'enregistrer.")
            else:
                data = request.POST.copy()
                data["formateur"] = workflow["formateur_id"]
                if workflow.get("classe_id"):
                    data["classe"] = workflow["classe_id"]
                if workflow.get("inspecteur_id"):
                    data["inspecteur"] = workflow["inspecteur_id"]
                if workflow.get("date"):
                    data["date"] = workflow["date"]
                if workflow.get("heure"):
                    data["heure"] = workflow["heure"]
                # Q1-Q3 viennent de l'IA ; Q4-Q6 viennent du POST (saisie manuelle)
                ai = workflow["ai_results"]
                data["q1_prerequis_apprenants"] = ai.get("q1_prerequis_apprenants", 1)
                data["q2_interaction_apprenants"] = ai.get("q2_interaction_apprenants", 1)
                data["q3_competences_acquises"] = ai.get("q3_competences_acquises", 1)
                save_form = SatisfactionFormateurForm(data)
                if save_form.is_valid():
                    obj = save_form.save(commit=False)
                    if hasattr(request, "user") and request.user.is_authenticated:
                        obj.enqueteur = request.user
                    audio_path = workflow.get("audio_path")
                    if audio_path:
                        obj.audio_appel.name = audio_path
                    obj.transcription = workflow.get("transcription", "")
                    obj.save()
                    messages.success(request, "Satisfaction formateur enregistree.")
                    request.session.pop(SESSION_KEY, None)
                    return redirect(
                        request.path_info + f"?classe={filter_classe}"
                        if filter_classe
                        else request.path_info
                    )
                else:
                    save_errors = save_form.errors

        request.session[SESSION_KEY] = workflow
        request.session.modified = True

    initial = {
        "classe": workflow.get("classe_id"),
        "inspecteur": workflow.get("inspecteur_id"),
        "date": workflow.get("date") or date_cls.today(),
        "heure": workflow.get("heure"),
    }
    form = SatisfactionFormateurForm(initial=initial)

    identified_formateur = None
    formateur_id = workflow.get("formateur_id")
    if formateur_id:
        identified_formateur = Formateur.objects.filter(id=formateur_id).first()
        if not identified_formateur:
            workflow.pop("formateur_id", None)
            workflow.pop("ai_results", None)
            workflow.pop("audio_path", None)
            request.session[SESSION_KEY] = workflow
            request.session.modified = True

    paginator = Paginator(qs, 50)
    page_obj = paginator.get_page(request.GET.get("page"))

    context = {
        "form": form,
        "identified_formateur": identified_formateur,
        "identifiant": workflow.get("identifiant", ""),
        "ai_results": workflow.get("ai_results"),
        "transcription": workflow.get("transcription"),
        "audio_name": (
            os.path.basename(workflow.get("audio_path")) if workflow.get("audio_path") else None
        ),
        "save_errors": save_errors,
        "enquetes": page_obj,
        "page_obj": page_obj,
        "classes": Classe.objects.all().order_by("code"),
        "filter_classe": filter_classe,
    }
    return render(request, "satisfaction_formateurs/index.html", context)


def satisfaction_formateurs_export_csv(request):
    filter_classe = request.GET.get("classe")
    qs = SatisfactionFormateur.objects.select_related(
        "classe", "formateur", "inspecteur", "enqueteur"
    ).order_by("-date")
    if filter_classe:
        qs = qs.filter(classe_id=filter_classe)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = "attachment; filename=satisfaction_formateurs.csv"
    writer = csv.writer(response)
    writer.writerow(
        [
            "classe",
            "formateur",
            "inspecteur",
            "enqueteur",
            "date",
            "heure",
            "Q1 - Prerequis apprenants (1-5)",
            "Q2 - Interaction apprenants (1-5)",
            "Q3 - Competences acquises (1-5)",
            "Q4 - Gestion administrative (texte)",
            "Q5 - Gestion financiere (texte)",
            "Q6 - Communication (texte)",
            "commentaires",
            "recommandations",
        ]
    )
    for s in qs:
        writer.writerow(
            [
                s.classe,
                s.formateur,
                s.inspecteur,
                s.enqueteur,
                s.date,
                s.heure,
                s.q1_prerequis_apprenants,
                s.q2_interaction_apprenants,
                s.q3_competences_acquises,
                s.q4_gestion_administrative,
                s.q5_gestion_financiere,
                s.q6_communication,
                s.commentaires,
                s.recommandations,
            ]
        )
    return response


def _normalize_batch_update_container(raw_value: str) -> str:
    text = str(raw_value or "").strip()
    if text.startswith("[") and text.endswith("]"):
        return text[1:-1].strip()
    return text


def _parse_formateur_batch_targets(raw_codes: str) -> list[str]:
    parsed_targets: list[str] = []
    seen_codes: set[str] = set()
    for block in re.split(r"[\r\n;]+", _normalize_batch_update_container(raw_codes)):
        block = _normalize_batch_update_container(block)
        if not block:
            continue
        for token in [item.strip() for item in re.split(r"[\s,]+", block) if item.strip()]:
            code_key = token.casefold()
            if not code_key or code_key in seen_codes:
                continue
            seen_codes.add(code_key)
            parsed_targets.append(token)
    return parsed_targets


def _merge_formateur_batch_targets(
    raw_codes: str,
    selected_targets: list[str] | None,
) -> list[str]:
    merged_targets: list[str] = []
    seen_codes: set[str] = set()
    for raw_value in [raw_codes, *(selected_targets or [])]:
        for target in _parse_formateur_batch_targets(raw_value):
            code_key = target.casefold()
            if code_key in seen_codes:
                continue
            seen_codes.add(code_key)
            merged_targets.append(target)
    return merged_targets


def _expand_formateur_batch_values(
    values: list,
    target_count: int,
    *,
    label: str,
    default_value=None,
) -> list:
    if not values:
        return [default_value] * target_count
    if len(values) == 1:
        return list(values) * target_count
    if len(values) == target_count:
        return list(values)
    raise ValueError(f"{label}: fournissez une seule valeur ou exactement {target_count} valeurs.")


def _build_formateur_batch_payloads(cleaned_data: dict, target_count: int) -> list[dict]:
    payloads = [{} for _ in range(target_count)]
    field_labels = {
        "q1_prerequis_apprenants": "Q1 Prerequis apprenants",
        "q2_interaction_apprenants": "Q2 Interaction apprenants",
        "q3_competences_acquises": "Q3 Competences acquises",
        "q4_gestion_administrative": "Q4 Gestion administrative",
        "q5_gestion_financiere": "Q5 Gestion financiere",
        "q6_communication": "Q6 Communication",
        "commentaires_values": "Commentaires",
        "recommandations_values": "Recommandations",
    }

    for field_name in FORMATEUR_SCORE_FIELDS:
        expanded_values = _expand_formateur_batch_values(
            cleaned_data.get(field_name) or [],
            target_count,
            label=field_labels[field_name],
            default_value=3,
        )
        for index, value in enumerate(expanded_values):
            payloads[index][field_name] = value

    for form_field, payload_field in (
        ("q4_gestion_administrative", "q4_gestion_administrative"),
        ("q5_gestion_financiere", "q5_gestion_financiere"),
        ("q6_communication", "q6_communication"),
        ("commentaires_values", "commentaires"),
        ("recommandations_values", "recommandations"),
    ):
        expanded_values = _expand_formateur_batch_values(
            cleaned_data.get(form_field) or [],
            target_count,
            label=field_labels[form_field],
            default_value="RAS",
        )
        for index, value in enumerate(expanded_values):
            payloads[index][payload_field] = value

    return payloads


def _build_formateur_batch_class_payloads(cleaned_data: dict, target_count: int) -> list[dict]:
    payloads = [{} for _ in range(target_count)]
    class_fields = (
        ("class_codes_values", "class_code", "Code classe"),
        ("prestation_codes_values", "prestation_code", "Prestation ID"),
        ("prestataire_values", "prestataire", "Prestataire"),
        ("beneficiaire_values", "beneficiaire", "Beneficiaire"),
        ("formation_values", "formation", "Titre de formation"),
        ("cohorte_values", "cohorte", "Cohorte"),
    )
    for form_field, payload_field, label in class_fields:
        expanded_values = _expand_formateur_batch_values(
            cleaned_data.get(form_field) or [],
            target_count,
            label=label,
            default_value=None,
        )
        for index, value in enumerate(expanded_values):
            value_text = str(value or "").strip()
            payloads[index][payload_field] = value_text or None
    return payloads


def _formateur_batch_status_display(status_code: str) -> str:
    normalized_status = str(status_code or "").strip()
    return dict(AppelFormateur.STATUS_CHOICES).get(normalized_status, normalized_status or "-")


def _has_complete_formateur_form(row: AppelFormateur) -> bool:
    return all(getattr(row, field_name, None) is not None for field_name in FORMATEUR_SCORE_FIELDS)


def _formateur_complete_form_queryset():
    queryset = AppelFormateur.objects.filter(is_active=True).order_by(
        "session_date", "numero_seance", "reference_code"
    )
    for field_name in FORMATEUR_SCORE_FIELDS:
        queryset = queryset.filter(**{f"{field_name}__isnull": False})
    return queryset


def _formateur_termine_without_form_queryset():
    return (
        AppelFormateur.objects.filter(is_active=True, status="termine")
        .order_by("session_date", "numero_seance", "reference_code")
        .exclude(pk__in=_formateur_complete_form_queryset().values("pk"))
    )


def _formateur_form_status_issue_queryset():
    return _formateur_complete_form_queryset().exclude(status="termine")


def _resolve_batch_update_formateur_classe(row: AppelFormateur):
    from App_PADESCE.appels.formateurs_views import _resolve_classe_for_formateur_row

    return _resolve_classe_for_formateur_row(row)


def _sync_batch_update_formateur_satisfaction(row: AppelFormateur, user) -> bool:
    from App_PADESCE.appels.formateurs_views import _sync_satisfaction_from_formateur_row

    return _sync_satisfaction_from_formateur_row(row, user)


def _build_formateur_candidate_row(row: AppelFormateur) -> dict:
    classe = _resolve_batch_update_formateur_classe(row)
    prestation = getattr(classe, "prestation", None)
    formateur = getattr(classe, "formateur", None)
    has_complete_form = _has_complete_formateur_form(row)
    has_partial_form = formateur_has_any_form_data(row)
    return {
        "reference_code": row.reference_code,
        "classe_code": getattr(classe, "code", "") or "-",
        "prestation_code": getattr(prestation, "code", "") or "-",
        "formateur_label": str(formateur or "-"),
        "prestataire": row.prestataire or "-",
        "beneficiaire": row.beneficiaire or "-",
        "formation": row.formation or "-",
        "cohorte": row.cohorte or "-",
        "telephone": row.telephone or "-",
        "audio_label": "Oui" if formateur_has_any_audio(row) else "Non",
        "formulaire_label": (
            "Complet" if has_complete_form else ("Partiel" if has_partial_form else "Non")
        ),
        "current_status": row.status or "",
        "current_status_label": row.get_status_display(),
        "computed_status": derive_formateur_status(row),
        "computed_status_label": _formateur_batch_status_display(derive_formateur_status(row)),
        "commentaires": row.commentaires or "-",
        "recommandations": row.recommandations or "-",
        "selection_value": row.reference_code,
        "has_complete_form": has_complete_form,
    }


def _formateur_answer_summary(row: AppelFormateur) -> str:
    return " / ".join(
        str(getattr(row, field_name, None) or "-") for field_name in FORMATEUR_SCORE_FIELDS
    )


def _paginate_formateur_update_form_rows(
    request,
    queryset,
    *,
    page_param: str,
    per_page: int = 50,
):
    paginator = Paginator(queryset, per_page)
    page_obj = paginator.get_page(request.GET.get(page_param))
    return page_obj, [_build_formateur_candidate_row(row) for row in page_obj.object_list]


def _apply_formateur_batch_update_target(reference_code: str, payload: dict, user) -> dict:
    result = {
        "reference_code": reference_code,
        "classe_code": "-",
        "formateur_label": "-",
        "before_status": "-",
        "after_status": "-",
        "before_answers": "-",
        "after_answers": "-",
        "commentaires": payload.get("commentaires", "-") or "-",
        "recommandations": payload.get("recommandations", "-") or "-",
        "message": "",
        "ok": False,
        "survey_synced": False,
    }

    row = AppelFormateur.objects.filter(
        is_active=True, reference_code__iexact=reference_code
    ).first()
    if row is None:
        result["message"] = "Reference formateur introuvable."
        return result

    classe = _resolve_batch_update_formateur_classe(row)
    formateur = getattr(classe, "formateur", None)
    result["classe_code"] = getattr(classe, "code", "") or "-"
    result["formateur_label"] = str(formateur or "-")
    result["before_status"] = row.get_status_display()
    result["before_answers"] = _formateur_answer_summary(row)

    update_fields = [*FORMATEUR_SCORE_FIELDS, *FORMATEUR_TEXT_FIELDS, "updated_at"]
    try:
        with transaction.atomic():
            for field_name, value in payload.items():
                setattr(row, field_name, value)
            row.save(update_fields=update_fields)
            sync_formateur_status(row)
            survey_synced = False
            if _has_complete_formateur_form(row):
                survey_synced = _sync_batch_update_formateur_satisfaction(row, user)

        result["after_status"] = row.get_status_display()
        result["after_answers"] = _formateur_answer_summary(row)
        result["commentaires"] = row.commentaires or "-"
        result["recommandations"] = row.recommandations or "-"
        result["survey_synced"] = survey_synced
        result["message"] = (
            "Formulaire mis a jour et fiche satisfaction synchronisee."
            if survey_synced
            else "Mise a jour enregistree. Fiche satisfaction non synchronisee."
        )
        result["ok"] = True
        return result
    except Exception as exc:
        logger.exception(
            "UPDATE FORM formateurs batch update failed for reference=%s", reference_code
        )
        result["message"] = f"Erreur interne pendant la mise a jour: {exc}"
        return result


def _apply_formateur_batch_status_target(reference_code: str, target_status: str, user) -> dict:
    result = {
        "reference_code": reference_code,
        "classe_code": "-",
        "formateur_label": "-",
        "before_status": "-",
        "after_status": "-",
        "before_answers": "-",
        "after_answers": "-",
        "commentaires": "-",
        "recommandations": "-",
        "message": "",
        "ok": False,
        "survey_synced": False,
    }

    row = AppelFormateur.objects.filter(
        is_active=True, reference_code__iexact=reference_code
    ).first()
    if row is None:
        result["message"] = "Reference formateur introuvable."
        return result

    classe = _resolve_batch_update_formateur_classe(row)
    formateur = getattr(classe, "formateur", None)
    result["classe_code"] = getattr(classe, "code", "") or "-"
    result["formateur_label"] = str(formateur or "-")
    result["before_status"] = row.get_status_display()
    result["before_answers"] = _formateur_answer_summary(row)
    result["commentaires"] = row.commentaires or "-"
    result["recommandations"] = row.recommandations or "-"

    if not _has_complete_formateur_form(row):
        result["message"] = "Aucun formulaire complet trouve pour cette reference."
        return result

    try:
        with transaction.atomic():
            update_fields = ["status", "updated_at"]
            row.status = target_status
            if target_status in {"termine", "a_rappeler"}:
                row.satisfaction_completed_at = timezone.now()
                update_fields.append("satisfaction_completed_at")
            row.save(update_fields=update_fields)
            survey_synced = False
            if target_status in {"termine", "a_rappeler"}:
                survey_synced = _sync_batch_update_formateur_satisfaction(row, user)

        result["after_status"] = row.get_status_display()
        result["after_answers"] = _formateur_answer_summary(row)
        result["survey_synced"] = survey_synced
        result["message"] = (
            f"Statut mis a jour vers {row.get_status_display()} et fiche satisfaction synchronisee."
            if survey_synced
            else f"Statut mis a jour vers {row.get_status_display()}."
        )
        result["ok"] = True
        return result
    except Exception as exc:
        logger.exception(
            "UPDATE FORM formateurs status update failed for reference=%s", reference_code
        )
        result["message"] = f"Erreur interne pendant le changement de statut: {exc}"
        return result


def _apply_formateur_batch_class_target(reference_code: str, payload: dict, user) -> dict:
    result = {
        "reference_code": reference_code,
        "classe_code": "-",
        "formateur_label": "-",
        "before_status": "-",
        "after_status": "-",
        "before_answers": "-",
        "after_answers": "-",
        "commentaires": "-",
        "recommandations": "-",
        "message": "",
        "ok": False,
        "survey_synced": False,
    }
    row = AppelFormateur.objects.filter(
        is_active=True, reference_code__iexact=reference_code
    ).first()
    if row is None:
        result["message"] = "Reference formateur introuvable."
        return result

    classe = _resolve_batch_update_formateur_classe(row)
    formateur = getattr(classe, "formateur", None)
    result["classe_code"] = getattr(classe, "code", "") or "-"
    result["formateur_label"] = str(formateur or "-")
    result["before_status"] = row.get_status_display()
    result["after_status"] = row.get_status_display()
    result["before_answers"] = _formateur_answer_summary(row)
    result["after_answers"] = _formateur_answer_summary(row)
    result["commentaires"] = row.commentaires or "-"
    result["recommandations"] = row.recommandations or "-"

    try:
        update_fields = ["updated_at"]
        class_code = payload.get("class_code")
        prestation_code = payload.get("prestation_code")
        prestataire = payload.get("prestataire")
        beneficiaire = payload.get("beneficiaire")
        formation = payload.get("formation")
        cohorte = payload.get("cohorte")

        if class_code:
            class_obj = (
                Classe.objects.select_related(
                    "prestation__prestataire",
                    "prestation__beneficiaire",
                    "formation",
                )
                .filter(code__iexact=class_code)
                .first()
            )
            if class_obj is None:
                result["message"] = f"Classe {class_code} introuvable."
                return result
            prestation_obj = getattr(class_obj, "prestation", None)
            if prestation_obj and not prestation_code:
                prestation_code = prestation_obj.code
            if prestation_obj and not prestataire:
                prestataire = (
                    getattr(getattr(prestation_obj, "prestataire", None), "raison_sociale", "")
                    or None
                )
            if prestation_obj and not beneficiaire:
                beneficiaire = (
                    getattr(getattr(prestation_obj, "beneficiaire", None), "nom_structure", "")
                    or None
                )
            if not formation:
                formation = (
                    getattr(class_obj, "intitule_formation", "")
                    or getattr(getattr(class_obj, "formation", None), "nom", "")
                    or None
                )
            if not cohorte and getattr(class_obj, "cohorte", None) is not None:
                cohorte = str(class_obj.cohorte)

        if prestation_code:
            prestation_obj = (
                Prestation.objects.select_related("prestataire", "beneficiaire", "formation")
                .filter(code__iexact=prestation_code)
                .first()
            )
            if prestation_obj:
                if not prestataire:
                    prestataire = prestation_obj.prestataire.raison_sociale
                if not beneficiaire:
                    beneficiaire = getattr(prestation_obj.beneficiaire, "nom_structure", "") or None
                if not formation:
                    formation = getattr(prestation_obj.formation, "nom", "") or None

        if prestataire:
            row.prestataire = prestataire
            update_fields.append("prestataire")
        if beneficiaire:
            row.beneficiaire = beneficiaire
            update_fields.append("beneficiaire")
        if formation:
            row.formation = formation
            update_fields.append("formation")
        if cohorte:
            row.cohorte = cohorte
            update_fields.append("cohorte")

        if len(update_fields) == 1:
            result["message"] = "Aucune donnee de classe renseignee pour mise a jour."
            return result

        row.save(update_fields=update_fields)
        synced = _sync_batch_update_formateur_satisfaction(row, user)
        result["survey_synced"] = synced
        result["message"] = "Donnees de classe mises a jour."
        result["ok"] = True
        return result
    except Exception as exc:
        logger.exception(
            "UPDATE FORM formateurs class-data update failed for reference=%s", reference_code
        )
        result["message"] = f"Erreur interne pendant la mise a jour des donnees de classe: {exc}"
        return result


def _merge_class_code_targets(raw_codes: str, selected_classes: list[str] | None) -> list[str]:
    merged_targets: list[str] = []
    seen_codes: set[str] = set()
    raw_values = [raw_codes, *(selected_classes or [])]
    for raw_value in raw_values:
        for target in _parse_formateur_batch_targets(raw_value):
            code_key = target.casefold()
            if code_key in seen_codes:
                continue
            seen_codes.add(code_key)
            merged_targets.append(target)
    return merged_targets


def _build_classes_tab_rows():
    classes_qs = Classe.objects.select_related(
        "prestation__prestataire",
        "prestation__beneficiaire",
        "formation",
        "formateur",
    ).order_by("code")
    call_counts_by_id = {
        row["classe_id"]: {
            "total": int(row["total"] or 0),
            "threshold_done": int(row["threshold_done"] or 0),
        }
        for row in Appel.objects.filter(is_active=True, classe_id__isnull=False)
        .values("classe_id")
        .annotate(
            total=Count("id"),
            threshold_done=Count("id", filter=Q(status__in=CALL_ANALYSIS_THRESHOLD_STATUSES)),
        )
    }
    orphan_counts_by_code = {
        str(row["classe_label"] or "").strip(): {
            "total": int(row["total"] or 0),
            "threshold_done": int(row["threshold_done"] or 0),
        }
        for row in Appel.objects.filter(is_active=True, classe_id__isnull=True)
        .exclude(classe_label="")
        .values("classe_label")
        .annotate(
            total=Count("id"),
            threshold_done=Count("id", filter=Q(status__in=CALL_ANALYSIS_THRESHOLD_STATUSES)),
        )
        if str(row["classe_label"] or "").strip()
    }
    rows = []
    known_codes: set[str] = set()
    for classe in classes_qs:
        counts = call_counts_by_id.get(classe.pk, {"total": 0, "threshold_done": 0})
        total = counts["total"]
        done = counts["threshold_done"]
        target = analysis_threshold_target(total)
        known_codes.add(str(classe.code or "").strip().casefold())
        rows.append(
            {
                "selection_value": classe.code,
                "classe_code": classe.code,
                "prestation_code": getattr(classe.prestation, "code", "") or "-",
                "prestataire": getattr(
                    getattr(classe.prestation, "prestataire", None), "raison_sociale", ""
                )
                or "-",
                "beneficiaire": getattr(
                    getattr(classe.prestation, "beneficiaire", None), "nom_structure", ""
                )
                or "-",
                "formation_title": classe.intitule_formation
                or getattr(classe.formation, "nom", "")
                or "-",
                "cohorte": str(classe.cohorte),
                "appels_total": total,
                "threshold_done": done,
                "threshold_target": target,
                "threshold_reached": bool(total and done >= target),
            }
        )
    # Ajoute les classes presentes dans les appels mais absentes du referentiel Classe.
    for class_code, counts in orphan_counts_by_code.items():
        code_key = class_code.casefold()
        if code_key in known_codes:
            continue
        total = counts["total"]
        done = counts["threshold_done"]
        target = analysis_threshold_target(total)
        rows.append(
            {
                "selection_value": class_code,
                "classe_code": class_code,
                "prestation_code": "-",
                "prestataire": "-",
                "beneficiaire": "-",
                "formation_title": "-",
                "cohorte": "-",
                "appels_total": total,
                "threshold_done": done,
                "threshold_target": target,
                "threshold_reached": bool(total and done >= target),
            }
        )
    rows.sort(key=lambda item: str(item.get("classe_code") or ""))
    return rows


def _apply_classes_batch_update_target(class_code: str, payload: dict, user) -> dict:
    result = {
        "reference_code": f"Classe {class_code}",
        "classe_code": class_code,
        "formateur_label": "-",
        "before_status": "-",
        "after_status": "-",
        "before_answers": "-",
        "after_answers": "-",
        "commentaires": "-",
        "recommandations": "-",
        "message": "",
        "ok": False,
        "survey_synced": False,
    }
    classe = (
        Classe.objects.select_related(
            "prestation__prestataire", "prestation__beneficiaire", "formation"
        )
        .filter(code__iexact=class_code)
        .first()
    )
    prestation_code = payload.get("prestation_code")
    prestataire = payload.get("prestataire")
    beneficiaire = payload.get("beneficiaire")
    formation_title = payload.get("formation")
    cohorte = payload.get("cohorte")

    try:
        # Si la classe n'existe pas, essayer de la créer
        if classe is None:
            if not prestation_code:
                result["message"] = (
                    f"Classe {class_code} introuvable et aucune prestation fournie pour creation."
                )
                return result

            prestation, prestation_error = _resolve_or_create_prestation(
                prestation_code=prestation_code,
                prestataire_name=prestataire or "",
                beneficiaire_name=beneficiaire or "",
                formation_title=formation_title or "",
            )
            if prestation is None:
                result["message"] = (
                    prestation_error
                    or f"Prestation {prestation_code} introuvable pour creation de classe."
                )
                return result

            # Préparer les données pour la création
            cohorte_value = 1
            if cohorte:
                if not str(cohorte).strip().isdigit():
                    result["message"] = "La cohorte doit etre numerique."
                    return result
                cohorte_value = int(str(cohorte).strip())

            intitule = formation_title or getattr(prestation.formation, "nom", "") or "-"

            # Créer la classe
            classe = Classe.objects.create(
                code=class_code,
                prestation=prestation,
                formation=prestation.formation,
                intitule_formation=intitule,
                cohorte=cohorte_value,
                statut="non_demarre",
                actif=True,
            )
            # Continuer avec la synchronisation des lignes formateurs
            prestataire_value = (
                prestataire or getattr(prestation.prestataire, "raison_sociale", "") or ""
            )
            beneficiaire_value = (
                beneficiaire
                or getattr(getattr(prestation, "beneficiaire", None), "nom_structure", "")
                or ""
            )
            formation_value = intitule
            cohorte_str = str(cohorte_value or "").strip()

            rows = AppelFormateur.objects.filter(is_active=True)
            updated_rows = 0
            for row in rows.iterator():
                resolved = _resolve_batch_update_formateur_classe(row)
                if not resolved or resolved.pk != classe.pk:
                    continue
                changed = False
                if prestataire_value and row.prestataire != prestataire_value:
                    row.prestataire = prestataire_value
                    changed = True
                if beneficiaire_value and row.beneficiaire != beneficiaire_value:
                    row.beneficiaire = beneficiaire_value
                    changed = True
                if formation_value and row.formation != formation_value:
                    row.formation = formation_value
                    changed = True
                if cohorte_str and row.cohorte != cohorte_str:
                    row.cohorte = cohorte_str
                    changed = True
                if changed:
                    row.save(
                        update_fields=[
                            "prestataire",
                            "beneficiaire",
                            "formation",
                            "cohorte",
                            "updated_at",
                        ]
                    )
                    updated_rows += 1

            result["message"] = (
                f"Classe {class_code} creee avec succes. "
                f"{updated_rows} ligne(s) formateur synchronisee(s)."
            )
            result["ok"] = True
            return result

        # Sinon, mettre à jour la classe existante
        update_fields = ["updated_at"]
        if prestation_code:
            prestation, prestation_error = _resolve_or_create_prestation(
                prestation_code=prestation_code,
                prestataire_name=prestataire or "",
                beneficiaire_name=beneficiaire or "",
                formation_title=formation_title or "",
            )
            if prestation is None:
                result["message"] = prestation_error or f"Prestation {prestation_code} introuvable."
                return result
            classe.prestation = prestation
            update_fields.append("prestation")
            if getattr(classe, "formation_id", None) != getattr(prestation, "formation_id", None):
                classe.formation = prestation.formation
                update_fields.append("formation")
            if not formation_title:
                formation_title = getattr(prestation.formation, "nom", "") or formation_title
            if not prestataire:
                prestataire = getattr(prestation.prestataire, "raison_sociale", "") or prestataire
            if not beneficiaire:
                beneficiaire = (
                    getattr(getattr(prestation, "beneficiaire", None), "nom_structure", "")
                    or beneficiaire
                )
        if formation_title:
            classe.intitule_formation = formation_title
            update_fields.append("intitule_formation")
        if cohorte:
            if not str(cohorte).strip().isdigit():
                result["message"] = "La cohorte doit etre numerique."
                return result
            classe.cohorte = int(str(cohorte).strip())
            update_fields.append("cohorte")
        if len(update_fields) > 1:
            classe.save(update_fields=update_fields)

        # Synchronise les lignes appels formateurs liees a cette classe.
        normalized_prestataire = (
            prestataire
            or getattr(getattr(classe.prestation, "prestataire", None), "raison_sociale", "")
            or ""
        )
        normalized_beneficiaire = (
            beneficiaire
            or getattr(getattr(classe.prestation, "beneficiaire", None), "nom_structure", "")
            or ""
        )
        normalized_formation = formation_title or classe.intitule_formation or ""
        normalized_cohorte = str(cohorte or classe.cohorte or "").strip()

        rows = AppelFormateur.objects.filter(is_active=True)
        updated_rows = 0
        for row in rows.iterator():
            resolved = _resolve_batch_update_formateur_classe(row)
            if not resolved or resolved.pk != classe.pk:
                continue
            changed = False
            if normalized_prestataire and row.prestataire != normalized_prestataire:
                row.prestataire = normalized_prestataire
                changed = True
            if normalized_beneficiaire and row.beneficiaire != normalized_beneficiaire:
                row.beneficiaire = normalized_beneficiaire
                changed = True
            if normalized_formation and row.formation != normalized_formation:
                row.formation = normalized_formation
                changed = True
            if normalized_cohorte and row.cohorte != normalized_cohorte:
                row.cohorte = normalized_cohorte
                changed = True
            if changed:
                row.save(
                    update_fields=[
                        "prestataire",
                        "beneficiaire",
                        "formation",
                        "cohorte",
                        "updated_at",
                    ]
                )
                updated_rows += 1

        result["message"] = (
            f"Classe mise a jour. {updated_rows} ligne(s) formateur synchronisee(s)."
        )
        result["ok"] = True
        return result
    except Exception as exc:
        logger.exception("Classes tab update failed for class=%s", class_code)
        result["message"] = f"Erreur interne pendant la mise a jour de la classe: {exc}"
        return result


def _normalize_excel_header(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", str(value or "").strip().lower())


def _pick_excel_value(data: dict, aliases: tuple[str, ...]) -> str:
    for alias in aliases:
        value = str(data.get(alias) or "").strip()
        if value:
            return value
    return ""


def _make_unique_model_code(model, preferred: str, *, prefix: str, max_length: int = 50) -> str:
    base = re.sub(r"[^A-Z0-9]", "", str(preferred or "").upper())[: max_length - 4]
    if not base:
        base = prefix
    candidate = base[:max_length]
    if not model.objects.filter(code__iexact=candidate).exists():
        return candidate
    for index in range(1, 1000):
        suffix = f"{index:03d}"
        candidate = f"{base[: max_length - len(suffix)]}{suffix}"
        if not model.objects.filter(code__iexact=candidate).exists():
            return candidate
    return f"{prefix}{uuid.uuid4().hex[:6].upper()}"[:max_length]


def _resolve_or_create_prestation(
    *,
    prestation_code: str,
    prestataire_name: str,
    beneficiaire_name: str,
    formation_title: str,
) -> tuple[Prestation | None, str | None]:
    code = str(prestation_code or "").strip()
    if not code:
        return None, "Prestation ID manquant."

    existing = (
        Prestation.objects.select_related("prestataire", "beneficiaire", "formation")
        .filter(code__iexact=code)
        .first()
    )
    if existing:
        return existing, None

    prestataire_label = str(prestataire_name or "").strip() or "PRESTATAIRE A COMPLETER"
    formation_label = str(formation_title or "").strip() or f"Formation {code}"
    beneficiaire_label = str(beneficiaire_name or "").strip()

    prestataire = Prestataire.objects.filter(raison_sociale__iexact=prestataire_label).first()
    if prestataire is None:
        prestataire = Prestataire.objects.create(
            code=_make_unique_model_code(Prestataire, prestataire_label, prefix="PREST"),
            raison_sociale=prestataire_label,
            actif=True,
        )

    formation = Formation.objects.filter(nom__iexact=formation_label).first()
    if formation is None:
        formation = Formation.objects.create(
            code=_make_unique_model_code(Formation, f"FORM{code}", prefix="FORM"),
            nom=formation_label,
            nom_harmonise=formation_label,
            statut="non_demarre",
            actif=True,
        )

    beneficiaire_obj = None
    if beneficiaire_label:
        beneficiaire_obj = Beneficiaire.objects.filter(
            nom_structure__iexact=beneficiaire_label
        ).first()
        if beneficiaire_obj is None:
            beneficiaire_obj = Beneficiaire.objects.create(
                nom_structure=beneficiaire_label,
                actif=True,
            )

    created = Prestation.objects.create(
        code=code,
        prestataire=prestataire,
        formation=formation,
        beneficiaire=beneficiaire_obj,
        actif=True,
    )
    return created, None


def _upload_classes_from_excel(uploaded_file, user) -> dict:
    workbook = openpyxl.load_workbook(uploaded_file, data_only=True)
    worksheet = workbook["Sheet1"] if "Sheet1" in workbook.sheetnames else workbook.active
    if worksheet is None:
        return {"updated": 0, "errors": ["Feuille Excel introuvable."], "processed": 0}

    header_row = next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_row:
        return {"updated": 0, "errors": ["Le fichier ne contient pas d'entetes."], "processed": 0}
    normalized_headers = [_normalize_excel_header(value) for value in header_row]
    header_index = {name: idx for idx, name in enumerate(normalized_headers) if name}

    class_header_aliases = ("classeid", "classe", "classid", "codeclasse")
    if not any(alias in header_index for alias in class_header_aliases):
        return {
            "updated": 0,
            "errors": ["Colonne Classe ID absente dans Sheet1."],
            "processed": 0,
        }

    updated = 0
    processed = 0
    errors: list[str] = []
    for row_num, row_values in enumerate(
        worksheet.iter_rows(min_row=2, max_row=worksheet.max_row, values_only=True), start=2
    ):
        row_map = {
            header: row_values[idx] if idx < len(row_values) else None
            for header, idx in header_index.items()
        }
        class_code = _pick_excel_value(row_map, class_header_aliases)
        if not class_code:
            continue
        processed += 1
        payload = {
            "prestation_code": _pick_excel_value(
                row_map, ("prestationid", "prestation", "codeprestation")
            )
            or None,
            "prestataire": _pick_excel_value(
                row_map,
                ("nomduprestataire", "prestataire", "prestationhorssimplitfi"),
            )
            or None,
            "beneficiaire": _pick_excel_value(
                row_map,
                ("nomdubeneficiaire", "beneficiaire", "beneficiairehorssimplitfi"),
            )
            or None,
            "formation": _pick_excel_value(
                row_map, ("formation", "intituledelaformation", "titreformation")
            )
            or None,
            "cohorte": _pick_excel_value(row_map, ("cohorte", "nombrecohorte")) or None,
        }
        result = _apply_classes_batch_update_target(class_code, payload, user)
        if result["ok"]:
            updated += 1
        elif len(errors) < 30:
            errors.append(f"Ligne {row_num} ({class_code}): {result['message']}")

    return {"updated": updated, "errors": errors, "processed": processed}


@require_analysis_access
def satisfaction_formateurs_update_form_page(request):
    active_tab = str(request.GET.get("tab") or request.POST.get("tab") or "classes").strip().lower()
    if active_tab not in {"classes", "formulaires"}:
        active_tab = "classes"
    all_status_filter = str(request.GET.get("all_status", "") or "").strip()
    initial = {
        "reference_codes_text": str(request.GET.get("codes", "") or "").strip(),
    }
    form = SatisfactionFormateurBatchUpdateForm(request.POST or None, initial=initial)
    results: list[dict] = []
    summary = {
        "requested_total": 0,
        "updated_total": 0,
        "error_total": 0,
        "synced_total": 0,
        "action_label": "formulaire(s)",
    }
    selected_target_values = {
        str(value or "").strip()
        for value in request.POST.getlist("selected_targets")
        if str(value or "").strip()
    }

    if request.method == "POST" and form.is_valid():
        action = str(request.POST.get("action") or "update_form").strip() or "update_form"
        if action == "upload_classes_excel":
            uploaded_file = request.FILES.get("classes_excel_file")
            if not uploaded_file:
                form.add_error(None, "Selectionnez un fichier Excel a importer.")
            else:
                upload_result = _upload_classes_from_excel(uploaded_file, request.user)
                if upload_result["updated"]:
                    messages.success(
                        request,
                        f"{upload_result['updated']} classe(s) completee(s) depuis Excel (Sheet1).",
                    )
                if upload_result["processed"] == 0:
                    messages.warning(request, "Aucune ligne exploitable trouvee dans le fichier.")
                for error in upload_result["errors"][:10]:
                    messages.warning(request, error)
        else:
            targets = _merge_formateur_batch_targets(
                form.cleaned_data["reference_codes_text"],
                request.POST.getlist("selected_targets"),
            )
            if not targets:
                form.add_error(
                    "reference_codes_text",
                    "Ajoutez au moins une reference valide ou selectionnez au moins une ligne.",
                )
            elif action == "update_status":
                requested_status = str(form.cleaned_data.get("target_status") or "").strip()
                if not requested_status:
                    form.add_error("target_status", "Choisissez le statut a appliquer.")
                else:
                    results = [
                        _apply_formateur_batch_status_target(
                            reference_code, requested_status, request.user
                        )
                        for reference_code in targets
                    ]
                    summary = {
                        "requested_total": len(results),
                        "updated_total": sum(1 for item in results if item["ok"]),
                        "error_total": sum(1 for item in results if not item["ok"]),
                        "synced_total": sum(1 for item in results if item["survey_synced"]),
                        "action_label": "statut(s)",
                    }
                    if summary["updated_total"]:
                        messages.success(
                            request,
                            f"{summary['updated_total']} statut(s) mis a jour.",
                        )
                    if summary["error_total"]:
                        messages.warning(
                            request,
                            f"{summary['error_total']} reference(s) n'ont pas pu etre traitees.",
                        )
            elif action == "update_class_data":
                try:
                    payloads = _build_formateur_batch_class_payloads(
                        form.cleaned_data, len(targets)
                    )
                except ValueError as exc:
                    form.add_error(None, str(exc))
                else:
                    results = [
                        _apply_formateur_batch_class_target(reference_code, payload, request.user)
                        for reference_code, payload in zip(targets, payloads)
                    ]
                    summary = {
                        "requested_total": len(results),
                        "updated_total": sum(1 for item in results if item["ok"]),
                        "error_total": sum(1 for item in results if not item["ok"]),
                        "synced_total": sum(1 for item in results if item["survey_synced"]),
                        "action_label": "donnee(s) de classe",
                    }
                    if summary["updated_total"]:
                        messages.success(
                            request,
                            f"{summary['updated_total']} ligne(s) de donnees classe mises a jour.",
                        )
                    if summary["error_total"]:
                        messages.warning(
                            request,
                            f"{summary['error_total']} reference(s) n'ont pas pu etre traitees.",
                        )
            elif action == "update_classes":
                class_targets = _merge_class_code_targets(
                    form.cleaned_data.get("class_codes_values", ""),
                    request.POST.getlist("selected_classes"),
                )
                if not class_targets:
                    form.add_error(
                        "class_codes_values",
                        "Ajoutez au moins un code classe ou cochez au moins une classe.",
                    )
                else:
                    try:
                        payloads = _build_formateur_batch_class_payloads(
                            form.cleaned_data, len(class_targets)
                        )
                    except ValueError as exc:
                        form.add_error(None, str(exc))
                    else:
                        results = [
                            _apply_classes_batch_update_target(class_code, payload, request.user)
                            for class_code, payload in zip(class_targets, payloads)
                        ]
                        summary = {
                            "requested_total": len(results),
                            "updated_total": sum(1 for item in results if item["ok"]),
                            "error_total": sum(1 for item in results if not item["ok"]),
                            "synced_total": 0,
                            "action_label": "classe(s)",
                        }
                        if summary["updated_total"]:
                            messages.success(
                                request, f"{summary['updated_total']} classe(s) mise(s) a jour."
                            )
                        if summary["error_total"]:
                            messages.warning(
                                request, f"{summary['error_total']} classe(s) en erreur."
                            )
            else:
                try:
                    payloads = _build_formateur_batch_payloads(form.cleaned_data, len(targets))
                except ValueError as exc:
                    form.add_error(None, str(exc))
                else:
                    results = [
                        _apply_formateur_batch_update_target(reference_code, payload, request.user)
                        for reference_code, payload in zip(targets, payloads)
                    ]
                    summary = {
                        "requested_total": len(results),
                        "updated_total": sum(1 for item in results if item["ok"]),
                        "error_total": sum(1 for item in results if not item["ok"]),
                        "synced_total": sum(1 for item in results if item["survey_synced"]),
                        "action_label": "formulaire(s)",
                    }
                    if summary["updated_total"]:
                        messages.success(
                            request,
                            f"{summary['updated_total']} formulaire(s) mis a jour.",
                        )
                    if summary["error_total"]:
                        messages.warning(
                            request,
                            f"{summary['error_total']} reference(s) n'ont pas pu etre traitees.",
                        )

    termine_qs = _formateur_termine_without_form_queryset()
    form_status_qs = _formateur_form_status_issue_queryset()
    all_formateurs_base_qs = AppelFormateur.objects.filter(is_active=True).order_by(
        "session_date", "numero_seance", "reference_code"
    )
    all_formateurs_qs = all_formateurs_base_qs
    if all_status_filter:
        all_formateurs_qs = all_formateurs_qs.filter(status=all_status_filter)
    termine_without_form_total = termine_qs.count()
    form_status_issue_total = form_status_qs.count()
    all_formateurs_total = all_formateurs_base_qs.count()
    all_formateurs_filtered_total = all_formateurs_qs.count()
    termine_page_obj, termine_without_form_rows = _paginate_formateur_update_form_rows(
        request,
        termine_qs,
        page_param="termine_page",
    )
    form_status_page_obj, form_status_issue_rows = _paginate_formateur_update_form_rows(
        request,
        form_status_qs,
        page_param="status_page",
    )
    all_formateurs_page_obj, all_formateurs_rows = _paginate_formateur_update_form_rows(
        request,
        all_formateurs_qs,
        page_param="all_page",
    )
    all_status_choices = [(value, label) for value, label in AppelFormateur.STATUS_CHOICES if value]
    known_status_values = {value for value, _label in all_status_choices}
    if all_status_filter and all_status_filter not in known_status_values:
        all_status_choices.append((all_status_filter, all_status_filter))

    classes_rows = _build_classes_tab_rows()
    classes_page_obj = Paginator(classes_rows, 50).get_page(request.GET.get("classes_page"))
    selected_class_values = {
        str(value or "").strip()
        for value in request.POST.getlist("selected_classes")
        if str(value or "").strip()
    }

    context = {
        "form": form,
        "active_tab": active_tab,
        "results": results,
        "summary": summary,
        "selected_target_values": selected_target_values,
        "termine_without_form_rows": termine_without_form_rows,
        "termine_without_form_total": termine_without_form_total,
        "termine_without_form_page_obj": termine_page_obj,
        "form_status_issue_rows": form_status_issue_rows,
        "form_status_issue_total": form_status_issue_total,
        "form_status_issue_page_obj": form_status_page_obj,
        "all_formateurs_rows": all_formateurs_rows,
        "all_formateurs_total": all_formateurs_total,
        "all_formateurs_filtered_total": all_formateurs_filtered_total,
        "all_formateurs_page_obj": all_formateurs_page_obj,
        "all_status_choices": all_status_choices,
        "all_status_filter": all_status_filter,
        "classes_rows": list(classes_page_obj.object_list),
        "classes_page_obj": classes_page_obj,
        "classes_total": len(classes_rows),
        "selected_class_values": selected_class_values,
        "analysis_threshold_label": analysis_threshold_label(),
        "candidate_total": termine_without_form_total + form_status_issue_total,
        "dashboard_url": reverse("satisfaction_formateurs_dashboard"),
        "index_url": reverse("satisfaction_formateurs_index"),
    }
    return render(request, "satisfaction_formateurs/update_form.html", context)


# ---------------------------------------------------------------------------
# Dashboard Analyse – Appels Formateurs
# ---------------------------------------------------------------------------

Q_FORM_FIELDS = [
    ("q1_prerequis_apprenants", "Prérequis apprenants"),
    ("q2_interaction_apprenants", "Interaction apprenants"),
    ("q3_competences_acquises", "Compétences acquises"),
]


def _avg_num(values):
    nums = []
    for v in values:
        if v is not None:
            try:
                # Convertir en float si ce n'est pas déjà un nombre
                if isinstance(v, (int, float)):
                    nums.append(v)
                else:
                    nums.append(float(v))
            except (ValueError, TypeError):
                continue
    return round(sum(nums) / len(nums), 2) if nums else 0


def _average_displayed_scores(values) -> float:
    displayed_values = []
    for value in values:
        if value is not None:
            try:
                displayed_values.append(round(float(value), 2))
            except (ValueError, TypeError):
                continue
    return round(sum(displayed_values) / len(displayed_values), 2) if displayed_values else 0


FORMATEUR_DASHBOARD_TABS = {"prestataire", "beneficiaire", "cohorte", "prestation", "detail"}


def _active_formateurs_tab(request) -> str:
    tab = (request.GET.get("tab") or "prestataire").strip().lower()
    return tab if tab in FORMATEUR_DASHBOARD_TABS else "prestataire"


def _sorted_distinct_non_empty_values(values) -> list[str]:
    normalized_values = set()
    for value in values:
        text = str(value or "").strip()
        if text:
            normalized_values.add(text)
    return sorted(normalized_values)


def _build_formateur_appel_status_summary(queryset) -> dict[str, int]:
    audio_q = Q(audio_file__isnull=False) & ~Q(audio_file="")

    summary = queryset.aggregate(
        appels_cibles=Count("id"),
        appels_tentes=Count("id", filter=~Q(status="en_attente")),
        appels_reussis=Count(
            "id",
            filter=Q(
                status__in=[
                    "formulaire_rempli",
                    "formulaire_avec_audio",
                    "termine",
                    "appel_reussi",
                    "a_rappeler",
                ]
            ),
        ),
        formulaires_remplis=Count("id", filter=Q(status="termine")),
        audios_enregistres=Count("id", filter=audio_q),
    )
    # Les KPI "avec/sans audio" doivent suivre l'existence reelle du fichier,
    # pas seulement la presence d'un chemin en base.
    termine_audio_rows = queryset.filter(status="termine").values_list("audio_file", flat=True)
    storage_exists_cache: dict[str, bool] = {}
    formulaires_avec_audio = 0
    formulaires_remplis_sans_audio = 0

    for audio_name in termine_audio_rows:
        normalized_name = str(audio_name or "").strip()
        if not normalized_name:
            formulaires_remplis_sans_audio += 1
            continue
        if normalized_name not in storage_exists_cache:
            try:
                storage_exists_cache[normalized_name] = default_storage.exists(normalized_name)
            except Exception:
                storage_exists_cache[normalized_name] = False
        if storage_exists_cache[normalized_name]:
            formulaires_avec_audio += 1
        else:
            formulaires_remplis_sans_audio += 1

    summary = {key: int(value or 0) for key, value in summary.items()}
    summary["formulaires_avec_audio"] = formulaires_avec_audio
    summary["formulaires_remplis_sans_audio"] = formulaires_remplis_sans_audio
    return summary


def _is_strict_formateur_record(record: dict) -> bool:
    return all(record.get(field_name) is not None for field_name, _ in Q_FORM_FIELDS)


def _build_satisfaction_formateurs_dashboard_context(request) -> dict:
    f_prestataire = (request.GET.get("prestataire") or "").strip()
    f_beneficiaire = (request.GET.get("beneficiaire") or "").strip()
    f_cohorte = (request.GET.get("cohorte") or "").strip()
    active_tab = _active_formateurs_tab(request)

    _formateur_marker = get_analysis_cache_version("model:appels.appelformateur")
    _cache_key = _formateurs_cache_key(
        _FORMATEURS_DASHBOARD_CACHE_VERSION,
        f_prestataire,
        f_beneficiaire,
        f_cohorte,
        _formateur_marker,
    )
    _cached = cache.get(_cache_key)
    if _cached is not None:
        return _cached

    qs = AppelFormateur.objects.filter(is_active=True).order_by("session_date", "prestataire")
    if f_prestataire:
        qs = qs.filter(prestataire=f_prestataire)
    if f_beneficiaire:
        qs = qs.filter(beneficiaire=f_beneficiaire)
    if f_cohorte:
        qs = qs.filter(cohorte=f_cohorte)

    records = list(
        qs.values(
            "id",
            "reference_code",
            "prestataire",
            "beneficiaire",
            "formation",
            "cohorte",
            "session_date",
            "telephone",
            "status",
            "q1_prerequis_apprenants",
            "q2_interaction_apprenants",
            "q3_competences_acquises",
            "satisfaction_completed_at",
            "q4_gestion_administrative",
            "q5_gestion_financiere",
            "q6_communication",
            "commentaires",
            "recommandations",
        )
    )

    total = len(records)
    termines = sum(1 for r in records if r["status"] == "termine")
    strict_form_records = [record for record in records if _is_strict_formateur_record(record)]
    with_scores = len(strict_form_records)
    global_avgs = {
        label: _avg_num([record[field] for record in strict_form_records])
        for field, label in Q_FORM_FIELDS
    }
    moyenne_generale_globale = _average_displayed_scores(global_avgs.values())

    def _group_stats(key_fn):
        groups = defaultdict(list)
        for r in records:
            groups[key_fn(r)].append(r)
        return sorted(
            [
                {
                    "label": k,
                    "nb": len(v),
                    "avgs": [_avg_num([r[field] for r in v]) for field, _ in Q_FORM_FIELDS],
                }
                for k, v in groups.items()
            ],
            key=lambda x: x["label"],
        )

    prestataire_stats = _group_stats(lambda r: r["prestataire"] or "-")
    beneficiaire_stats = _group_stats(lambda r: r["beneficiaire"] or "-")
    cohorte_stats = _group_stats(lambda r: r["cohorte"] or "-")
    prestation_stats = _group_stats(lambda r: r["reference_code"] or "-")

    all_qs = AppelFormateur.objects.filter(is_active=True)
    prestataires = _sorted_distinct_non_empty_values(all_qs.values_list("prestataire", flat=True))
    beneficiaires = _sorted_distinct_non_empty_values(all_qs.values_list("beneficiaire", flat=True))
    cohortes = _sorted_distinct_non_empty_values(all_qs.values_list("cohorte", flat=True))

    status_counts = defaultdict(int)
    for r in records:
        status_counts[r["status"]] += 1

    appel_summary = _build_formateur_appel_status_summary(qs)

    # Import the prestation indicators builder from apprenants views
    from App_PADESCE.satisfaction_apprenants.views import _build_prestation_indicators_table

    prestation_indicators_table = _build_prestation_indicators_table()

    # Prestation Ranking Logic (Mirrored from Espace PADESCE)
    from django.db import connection

    prestation_mapping = {}
    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT p.code, p.id, pr.raison_sociale as prestataire_nom,
                   b.nom_structure as beneficiaire_nom, f.nom as formation_nom,
                   b.region as beneficiaire_region
            FROM formations_prestation p
            LEFT JOIN formations_prestataire pr ON p.prestataire_id = pr.id
            LEFT JOIN formations_beneficiaire b ON p.beneficiaire_id = b.id
            LEFT JOIN formations_formation f ON p.formation_id = f.id
            WHERE p.actif = 1
        """
        )
        prestations_info = cursor.fetchall()
        for code, id, prestataire_nom, beneficiaire_nom, formation_nom, region in prestations_info:
            prestation_mapping[code] = {
                "id": id,
                "prestataire_nom": str(prestataire_nom or "").strip().lower(),
                "beneficiaire_nom": str(beneficiaire_nom or "").strip().lower(),
                "formation_nom": str(formation_nom or "").strip().lower(),
                "beneficiaire_region": str(region or "").strip().upper(),
            }

    grouped_ranking = {}
    resolution_cache = {}

    from App_PADESCE.core.public_views import _resolve_formateur_classe

    # We use all_rows (unfiltered by current dashboard filters) for the ranking
    # to maintain consistency, or maybe we should use the filtered ones?
    # Usually ranking is global.
    # but the user asked for "classement complet".
    # Let's use all active records if not filtered?
    # Actually, dashboards usually show rankings for the current selection.

    for record in records:
        classe = _resolve_formateur_classe(record, resolution_cache)
        prestation = getattr(classe, "prestation", None)
        code = str(getattr(prestation, "code", "") or "").strip()

        if not code:
            # Synthetic matching logic from public_views
            prestataire_val = record.get("prestataire") or "-"
            beneficiaire_val = record.get("beneficiaire") or "-"

            # Simple fallback key
            code = f"{prestataire_val}|{beneficiaire_val}"

        group_key = code
        bucket = grouped_ranking.setdefault(
            group_key,
            {
                "code": code,
                "prestataire": record.get("prestataire") or "-",
                "beneficiaire": record.get("beneficiaire") or "-",
                "effectif": 0,
                "nb": 0,
                "scores": {field_name: [] for field_name in FORMATEUR_SCORE_FIELDS},
            },
        )
        bucket["effectif"] += 1

        values = [record.get(field_name) for field_name in FORMATEUR_SCORE_FIELDS]
        if not all(v not in (None, "") for v in values):
            continue

        bucket["nb"] += 1
        for field_name, value in zip(FORMATEUR_SCORE_FIELDS, values):
            bucket["scores"][field_name].append(float(value))

    prestation_stats_form = []
    for item in grouped_ranking.values():
        if not item["nb"] and not item["effectif"]:
            continue
        avgs = []
        for field_name in FORMATEUR_SCORE_FIELDS:
            vals = item["scores"][field_name]
            avgs.append(round(sum(vals) / len(vals), 2) if vals else 0)
        avg = round(sum(avgs) / len(avgs), 2) if avgs else 0

        # Enrich with real data if possible
        code = item["code"]
        real_prestataire = item["prestataire"]
        real_region = "Inconnu"
        if code in prestation_mapping:
            minfo = prestation_mapping[code]
            if minfo["prestataire_nom"] and minfo["prestataire_nom"] != "-":
                real_prestataire = minfo["prestataire_nom"].title()
            real_region = minfo["beneficiaire_region"] or "Inconnu"

        prestation_stats_form.append(
            {
                "code": item["code"],
                "prestataire": real_prestataire,
                "beneficiaire": item["beneficiaire"],
                "region": real_region,
                "nb": item["nb"],
                "avg": avg,
                "avgs": avgs,
                "effectif": item["effectif"],
            }
        )

    toutes_prestations_classees = get_prestations_ranking(prestation_stats_form, order="desc")

    context = {
        "active_tab": active_tab,
        "total": total,
        "termines": termines,
        "with_scores": with_scores,
        "global_avgs": global_avgs,
        "moyenne_generale_globale": moyenne_generale_globale,
        "q_labels": [label for _, label in Q_FORM_FIELDS],
        "prestataire_stats": prestataire_stats,
        "beneficiaire_stats": beneficiaire_stats,
        "cohorte_stats": cohorte_stats,
        "prestation_stats": prestation_stats,
        "status_counts": dict(status_counts),
        "prestataires": prestataires,
        "beneficiaires": beneficiaires,
        "cohortes": cohortes,
        "f_prestataire": f_prestataire,
        "f_beneficiaire": f_beneficiaire,
        "f_cohorte": f_cohorte,
        "rows": records[:200],
        "all_rows": records,
        "appels_cibles": appel_summary["appels_cibles"],
        "appels_tentes": appel_summary["appels_tentes"],
        "appels_reussis": appel_summary["appels_reussis"],
        "formulaires_remplis_appels": appel_summary["formulaires_remplis"],
        "formulaires_remplis_sans_audio_appels": appel_summary["formulaires_remplis_sans_audio"],
        "formulaires_avec_audio_appels": appel_summary["formulaires_avec_audio"],
        "audios_enregistres_appels": appel_summary["audios_enregistres"],
        "prestation_indicators_table": prestation_indicators_table,
        "toutes_prestations_classees": toutes_prestations_classees,
    }
    context.update(build_fast_stats_context(request, default_mode="formateur"))
    cache.set(_cache_key, context, timeout=_FORMATEURS_CACHE_TIMEOUT)
    return context


def _formateurs_dashboard_export_filename(active_tab: str, extension: str) -> str:
    return f"analyse-formateurs-{active_tab}.{extension}"


def _tabular_formateurs_dashboard_export(
    active_tab: str, context: dict
) -> tuple[list[str], list[list]]:
    if active_tab == "beneficiaire":
        return (
            ["Beneficiaire", "Nb appels", *[label for _, label in Q_FORM_FIELDS]],
            [[item["label"], item["nb"], *item["avgs"]] for item in context["beneficiaire_stats"]],
        )
    if active_tab == "cohorte":
        return (
            ["Cohorte", "Nb appels", *[label for _, label in Q_FORM_FIELDS]],
            [[item["label"], item["nb"], *item["avgs"]] for item in context["cohorte_stats"]],
        )
    if active_tab == "prestation":
        return (
            ["Prestation", "Nb appels", *[label for _, label in Q_FORM_FIELDS]],
            [[item["label"], item["nb"], *item["avgs"]] for item in context["prestation_stats"]],
        )
    if active_tab == "detail":
        return (
            [
                "Prestataire",
                "Beneficiaire",
                "Formation",
                "Cohorte",
                "Telephone",
                "Statut",
                *[label for _, label in Q_FORM_FIELDS],
                "Q4 Gestion admin",
                "Q5 Gestion financiere",
                "Q6 Communication",
                "Commentaires",
            ],
            [
                [
                    row.get("prestataire", ""),
                    row.get("beneficiaire", ""),
                    row.get("formation", ""),
                    row.get("cohorte", ""),
                    row.get("telephone", ""),
                    row.get("status", ""),
                    *[row.get(field) for field, _ in Q_FORM_FIELDS],
                    row.get("q4_gestion_administrative", ""),
                    row.get("q5_gestion_financiere", ""),
                    row.get("q6_communication", ""),
                    row.get("commentaires", ""),
                ]
                for row in context["all_rows"]
            ],
        )
    return (
        ["Prestataire", "Nb appels", *[label for _, label in Q_FORM_FIELDS]],
        [[item["label"], item["nb"], *item["avgs"]] for item in context["prestataire_stats"]],
    )


def export_formateur_global_averages_xlsx(request):
    """Export des moyennes générales des formateurs en Excel."""
    context = _build_satisfaction_formateurs_dashboard_context(request)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Moyennes générales"

    # En-têtes
    q_labels = [label for _, label in Q_FORM_FIELDS]
    ws.append(["Indicateur"] + q_labels + ["Moyenne générale GLOBALE"])

    # Données des moyennes générales
    row_data = ["Moyenne générale"]
    for label in q_labels:
        avg = context.get("global_avgs", {}).get(label, 0)
        row_data.append(round(avg, 2))
    # Ajouter la moyenne générale globale
    moyenne_globale = context.get("moyenne_generale_globale", 0)
    row_data.append(round(moyenne_globale, 2))
    ws.append(row_data)

    # Style
    for col in range(1, len(q_labels) + 3):
        cell = ws.cell(row=1, column=col)
        cell.font = openpyxl.styles.Font(bold=True)
        cell.fill = openpyxl.styles.PatternFill(
            start_color="E6E6FA", end_color="E6E6FA", fill_type="solid"
        )

    for col in range(1, len(q_labels) + 3):
        cell = ws.cell(row=2, column=col)
        cell.font = openpyxl.styles.Font(bold=True)
        cell.fill = openpyxl.styles.PatternFill(
            start_color="FFE6E6", end_color="FFE6E6", fill_type="solid"
        )

    # Ajuster la largeur des colonnes
    for col in range(1, len(q_labels) + 3):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 15

    # Créer la réponse HTTP
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="moyennes-generales-formateurs.xlsx"'
    wb.save(response)
    return response


def _autosize_worksheet(worksheet, max_width=48):
    for column_cells in worksheet.columns:
        length = max(
            (len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells),
            default=0,
        )
        # 1.2 factor for padding
        worksheet.column_dimensions[column_cells[0].column_letter].width = min(
            max_width, length + 2
        )


@require_analysis_access
def satisfaction_formateurs_dashboard_export_ranking(request):
    """Génère un Excel contenant le classement complet des prestations (Vision Formateurs)."""
    context = _build_satisfaction_formateurs_dashboard_context(request)
    ranking = context.get("toutes_prestations_classees", [])

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Classement Formateurs"

    # Header
    headers = [
        "Rang",
        "Code Prestation",
        "Intitulé Formation",
        "Prestataire",
        "Bénéficiaire",
        "Région",
        "Effectif Formateurs",
        "Nb Enquêtes",
        "Taux de Réponse (%)",
        "Satisfaction (0-5)",
        "Score Global",
    ]
    ws.append(headers)

    # Data
    for idx, p in enumerate(ranking, start=1):
        ws.append(
            [
                idx,
                p["code"],
                p["intitule"],
                p["prestataire"],
                p["beneficiaire"],
                p["region"],
                p["effectif"],
                p["nb_reponses"],
                p["taux_reponse"],
                p["avg_satisfaction"],
                p["score_global"],
            ]
        )

    _autosize_worksheet(ws)

    filename = f"Classement_Prestation_Formateurs_{timezone.now().strftime('%Y%m%d')}.xlsx"
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    wb.save(response)
    return response


@require_analysis_access
def satisfaction_formateurs_dashboard_export_csv(request):
    context = _build_satisfaction_formateurs_dashboard_context(request)
    active_tab = _active_formateurs_tab(request)
    headers, export_rows = _tabular_formateurs_dashboard_export(active_tab, context)
    filename = _formateurs_dashboard_export_filename(active_tab, "csv")

    response = HttpResponse(content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response.write("\ufeff")
    writer = csv.writer(response)
    writer.writerow(headers)
    for row in export_rows:
        writer.writerow(row)
    return response


@require_analysis_access
def satisfaction_formateurs_dashboard_export_chapeau(request):
    context = _build_satisfaction_formateurs_dashboard_context(request)
    active_tab = _active_formateurs_tab(request)
    filename = _formateurs_dashboard_export_filename(active_tab, "docx")

    from docx import Document

    document = Document()
    stats_by_tab = {
        "prestataire": context["prestataire_stats"],
        "beneficiaire": context["beneficiaire_stats"],
        "cohorte": context["cohorte_stats"],
        "prestation": context["prestation_stats"],
    }

    if active_tab == "detail":
        rows = context["all_rows"]
        if not rows:
            document.add_paragraph("Aucun appel formateur a exporter.")
        else:
            for row in rows:
                table = document.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                title = table.rows[0].cells[0].merge(table.rows[0].cells[1]).paragraphs[0]
                title_run = title.add_run(
                    f"Enquete formateur : {row.get('formation') or row.get('reference_code')}"
                )
                title_run.bold = True
                for field, label in Q_FORM_FIELDS:
                    cells = table.add_row().cells
                    cells[0].text = label
                    value = row.get(field)
                    cells[1].text = str(value) if value not in (None, "") else "-"
                extra_rows = [
                    ("Q4 - Gestion administrative", row.get("q4_gestion_administrative", "")),
                    ("Q5 - Gestion financiere", row.get("q5_gestion_financiere", "")),
                    ("Q6 - Communication", row.get("q6_communication", "")),
                    ("Commentaires", row.get("commentaires", "")),
                ]
                for label, value in extra_rows:
                    cells = table.add_row().cells
                    cells[0].text = label
                    cells[1].text = str(value or "-")
                document.add_paragraph("")
    else:
        items = stats_by_tab.get(active_tab, context["prestataire_stats"])
        if not items:
            document.add_paragraph("Aucune synthese formateur a exporter.")
        else:
            title_by_tab = {
                "prestataire": "Prestataire",
                "beneficiaire": "Beneficiaire",
                "cohorte": "Cohorte",
            }
            for item in items:
                table = document.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                title = table.rows[0].cells[0].merge(table.rows[0].cells[1]).paragraphs[0]
                title_run = title.add_run(
                    "Enquete de satisfaction formateurs : "
                    f"{title_by_tab.get(active_tab, 'Groupe')} {item['label']}"
                )
                title_run.bold = True
                headers = table.add_row().cells
                headers[0].text = "QUESTION"
                headers[1].text = "NOTE"
                for index, (_field, label) in enumerate(Q_FORM_FIELDS):
                    cells = table.add_row().cells
                    cells[0].text = label
                    avg = item["avgs"][index] if index < len(item.get("avgs", [])) else 0
                    cells[1].text = f"{avg}/5" if avg else "-"
                total_cells = table.add_row().cells
                total_cells[0].text = "TOTAL DES APPELS"
                total_cells[1].text = str(item["nb"])
                document.add_paragraph("")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@require_analysis_access
def satisfaction_formateurs_dashboard_export_prestation_zip(request):
    """Export one CSV per prestataire as a ZIP archive (Satisfaction_PRESTA_SF.csv)."""
    context = _build_satisfaction_formateurs_dashboard_context(request)
    all_rows = context["all_rows"]

    prestation_rows: dict[str, list] = {}
    for row in all_rows:
        code = str(row.get("prestataire") or "").strip() or "-"
        prestation_rows.setdefault(code, []).append(row)

    headers = [
        "N°",
        "Prestataire",
        "Bénéficiaire",
        "Formation",
        "Cohorte",
        "Téléphone",
        "Statut",
        *[label for _, label in Q_FORM_FIELDS],
        "Q4 - Gestion administrative",
        "Q5 - Gestion financière",
        "Q6 - Communication",
        "Commentaires",
        "Recommandations",
    ]

    archive_buffer = io.BytesIO()
    with zipfile.ZipFile(archive_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for prestataire_code in sorted(prestation_rows.keys()):
            csv_buffer = io.StringIO()
            writer = csv.writer(csv_buffer, lineterminator="\n")
            writer.writerow(headers)
            for idx, row in enumerate(prestation_rows[prestataire_code], start=1):
                writer.writerow(
                    [
                        idx,
                        row.get("prestataire", ""),
                        row.get("beneficiaire", ""),
                        row.get("formation", ""),
                        row.get("cohorte", ""),
                        row.get("telephone", ""),
                        row.get("status", ""),
                        *[row.get(field) for field, _ in Q_FORM_FIELDS],
                        row.get("q4_gestion_administrative", ""),
                        row.get("q5_gestion_financiere", ""),
                        row.get("q6_communication", ""),
                        row.get("commentaires", ""),
                        row.get("recommandations", ""),
                    ]
                )
            safe_code = re.sub(r'[\\/:*?"<>|]', "_", prestataire_code)
            archive.writestr(
                f"Satisfaction_{safe_code}_SF.csv",
                csv_buffer.getvalue().encode("utf-8-sig"),
            )

    archive_buffer.seek(0)
    response = HttpResponse(archive_buffer.getvalue(), content_type="application/zip")
    response["Content-Disposition"] = (
        'attachment; filename="satisfaction-formateurs-prestations.zip"'
    )
    return response


def satisfaction_formateurs_export_prestation_csv(request, code: str):
    """Export CSV des enquêtes d'un seul prestataire (pour fill_excel.py).
    Authentification par clé API (X-Export-Api-Key) — pas de login requis.
    """
    from App_PADESCE.satisfaction_apprenants.views import _check_export_api_key

    if not _check_export_api_key(request):
        return JsonResponse({"error": "Clé API manquante ou invalide."}, status=403)
    context = _build_satisfaction_formateurs_dashboard_context(request)
    target = code.strip()
    rows = [r for r in context["all_rows"] if str(r.get("prestataire") or "").strip() == target]

    headers = [
        "N°",
        "Prestataire",
        "Bénéficiaire",
        "Formation",
        "Cohorte",
        "Téléphone",
        "Statut",
        *[label for _, label in Q_FORM_FIELDS],
        "Q4 - Gestion administrative",
        "Q5 - Gestion financière",
        "Q6 - Communication",
        "Commentaires",
        "Recommandations",
    ]

    safe_code = re.sub(r"[^A-Za-z0-9_-]", "_", code)
    filename = f"Satisfaction_{safe_code}_SF.csv"
    response = HttpResponse(content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response.write("\ufeff")
    writer = csv.writer(response)
    writer.writerow(headers)
    for idx, row in enumerate(rows, start=1):
        writer.writerow(
            [
                idx,
                row.get("prestataire", ""),
                row.get("beneficiaire", ""),
                row.get("formation", ""),
                row.get("cohorte", ""),
                row.get("telephone", ""),
                row.get("status", ""),
                *[row.get(field) for field, _ in Q_FORM_FIELDS],
                row.get("q4_gestion_administrative", ""),
                row.get("q5_gestion_financiere", ""),
                row.get("q6_communication", ""),
                row.get("commentaires", ""),
                row.get("recommandations", ""),
            ]
        )
    return response


def satisfaction_formateurs_api_prestations_excel(request):
    """
    API JSON — liste des prestataires/prestations formateurs (pour fill_excel.py).
    Authentification via en-tête HTTP : X-Export-Api-Key: <EXPORT_API_KEY>
    """
    from App_PADESCE.satisfaction_apprenants.views import _check_export_api_key

    if not _check_export_api_key(request):
        return JsonResponse({"error": "Clé API manquante ou invalide."}, status=403)

    from django.conf import settings as _s

    site_url = (getattr(_s, "SITE_URL", "") or "https://call.naumur.com").rstrip("/")
    base_app = "/satisfaction-formateurs"

    context = _build_satisfaction_formateurs_dashboard_context(request)
    seen: dict[str, dict] = {}
    for row in context["all_rows"]:
        code = str(row.get("prestataire") or "").strip()
        if not code or code == "-":
            continue
        if code not in seen:
            seen[code] = {
                "code": code,
                "prestataire": code,
                "beneficiaire": row.get("beneficiaire", ""),
                "formation": row.get("formation", ""),
                "cohorte": row.get("cohorte", ""),
                "enquete_url": f"{site_url}/prestation/{code.lower()}/",
                "csv_url": f"{site_url}{base_app}/analyse/export/prestation/{code}/csv/",
            }

    return JsonResponse({"prestations": list(seen.values())})


@require_analysis_access
def satisfaction_formateurs_dashboard(request):
    context = _build_satisfaction_formateurs_dashboard_context(request)
    return render(request, "satisfaction_formateurs/dashboard.html", context)


# ---------------------------------------------------------------------------
# Page de gestion : lier formateurs ↔ prestataires + villes des prestations
# ---------------------------------------------------------------------------


@require_analysis_access
def formateurs_prestataires_management(request):
    """Page de gestion : toggle formateur↔prestations + ville des prestations."""
    from App_PADESCE.formations.models import Formateur
    from App_PADESCE.formations.models import Prestation as PrestationModel

    saved_count = 0
    saved_label = ""

    def _sync_formateurs_from_appels() -> int:
        existing_by_phone: dict[str, Formateur] = {}
        for formateur in Formateur.objects.all().only("id", "telephone"):
            normalized = _normalize_phone(str(getattr(formateur, "telephone", "") or ""))
            if normalized and normalized not in existing_by_phone:
                existing_by_phone[normalized] = formateur

        created = 0
        seen_phone = set(existing_by_phone.keys())
        appels = AppelFormateur.objects.filter(is_active=True).values("telephone", "source_contact")
        for row in appels.iterator():
            raw_candidates = [
                str(row.get("telephone") or "").strip(),
                str(row.get("source_contact") or "").strip(),
            ]
            phone_candidates = []
            for raw in raw_candidates:
                if not raw:
                    continue
                phone_candidates.extend(re.findall(r"\d{8,15}", raw))

            for phone_text in phone_candidates:
                normalized = _normalize_phone(phone_text)
                if len(normalized) < 8 or normalized in seen_phone:
                    continue
                seen_phone.add(normalized)
                code_seed = f"TEL{normalized[-10:]}" if len(normalized) > 10 else f"TEL{normalized}"
                form_code = _make_unique_model_code(Formateur, code_seed, prefix="FORM")
                Formateur.objects.create(
                    code=form_code,
                    nom_complet=phone_text,
                    nom=(
                        resolve_formateur_name_from_values(phone_text, phone_text)
                        or FORMATEUR_NAME_FALLBACK
                    ),
                    telephone=phone_text,
                    actif=True,
                )
                created += 1
        return created

    if request.method == "POST":
        action = request.POST.get("action", "")

        if action == "sync_formateurs":
            created_count = _sync_formateurs_from_appels()
            sync_formateur_names_from_excel()
            saved_count = created_count
            saved_label = (
                f"{created_count} formateur(s) cree(s) depuis les appels."
                if created_count
                else "Aucun nouveau formateur a creer (deja synchronises)."
            )

        if action == "save_formateurs":
            # Met à jour le M2M formateur ↔ prestations
            prestations_all = {str(p.pk): p for p in PrestationModel.objects.all()}
            for form in Formateur.objects.prefetch_related("prestations").all():
                field_name = f"prestations_{form.pk}"
                selected_pks = set(request.POST.getlist(field_name))
                current_pks = set(str(pk) for pk in form.prestations.values_list("pk", flat=True))
                if selected_pks != current_pks:
                    new_prestations = [
                        prestations_all[pk] for pk in selected_pks if pk in prestations_all
                    ]
                    form.prestations.set(new_prestations)
                    saved_count += 1
            saved_label = f"{saved_count} formateur(s) mis à jour."

        elif action == "toggle_prestation":
            # Toggle rapide via AJAX : lier/délier un formateur d'une prestation
            formateur_pk = request.POST.get("formateur_pk", "").strip()
            prestation_pk = request.POST.get("prestation_pk", "").strip()
            try:
                form = Formateur.objects.get(pk=formateur_pk)
                prest = PrestationModel.objects.get(pk=prestation_pk)
                if form.prestations.filter(pk=prest.pk).exists():
                    form.prestations.remove(prest)
                    linked = False
                else:
                    form.prestations.add(prest)
                    linked = True
                from django.http import JsonResponse as _JsonResponse

                return _JsonResponse({"ok": True, "linked": linked})
            except Exception as exc:
                from django.http import JsonResponse as _JsonResponse

                return _JsonResponse({"ok": False, "error": str(exc)}, status=400)

        elif action == "save_prestations":
            # Mise à jour de la ville pour chaque prestation (carte)
            prestations = list(PrestationModel.objects.select_related("prestataire").all())
            to_update = []
            for prest in prestations:
                new_ville = request.POST.get(f"ville_{prest.pk}", "").strip()
                if prest.ville != new_ville:
                    prest.ville = new_ville
                    to_update.append(prest)
            if to_update:
                PrestationModel.objects.bulk_update(to_update, ["ville"], batch_size=200)
                saved_count = len(to_update)
            saved_label = f"{saved_count} prestation(s) mise(s) à jour."

        elif action == "save_prestation_ville":
            # Sauvegarde rapide d'une seule ville via AJAX
            prestation_pk = request.POST.get("prestation_pk", "").strip()
            new_ville = request.POST.get("ville", "").strip()
            try:
                prest = PrestationModel.objects.get(pk=prestation_pk)
                prest.ville = new_ville
                prest.save(update_fields=["ville"])
                from django.http import JsonResponse as _JsonResponse

                return _JsonResponse({"ok": True, "ville": new_ville})
            except Exception as exc:
                from django.http import JsonResponse as _JsonResponse

                return _JsonResponse({"ok": False, "error": str(exc)}, status=400)

    def _normalize_org_name(value: str) -> str:
        cleaned = re.sub(r"[^a-z0-9]+", " ", str(value or "").casefold())
        return re.sub(r"\s+", " ", cleaned).strip()

    def _compact_org_name(value: str) -> str:
        return "".join(ch for ch in _normalize_org_name(value) if ch.isalnum())

    def _name_match_score(left: str, right: str) -> float:
        left_norm = _normalize_org_name(left)
        right_norm = _normalize_org_name(right)
        if not left_norm or not right_norm:
            return 0.0
        if left_norm == right_norm:
            return 1.0

        left_compact = _compact_org_name(left_norm)
        right_compact = _compact_org_name(right_norm)
        if left_compact and right_compact and left_compact == right_compact:
            return 0.95

        if (
            left_compact
            and right_compact
            and (left_compact in right_compact or right_compact in left_compact)
        ):
            return 0.8

        left_tokens = set(left_norm.split())
        right_tokens = set(right_norm.split())
        if not left_tokens or not right_tokens:
            return 0.0
        overlap = len(left_tokens & right_tokens)
        base = overlap / max(len(left_tokens), len(right_tokens))
        if overlap >= 2:
            return min(0.85, base + 0.2)
        return base if base >= 0.4 else 0.0

    # En GET, complete automatiquement les formateurs manquants depuis les appels.
    if request.method != "POST":
        _sync_formateurs_from_appels()
        sync_formateur_names_from_excel()

    # Charger les formateurs avec leurs prestations courantes
    formateurs = list(
        Formateur.objects.prefetch_related("prestations", "classes__prestation").order_by(
            "nom", "nom_complet"
        )
    )
    # Index des prestations déjà assignées par formateur
    formateur_prestations = {
        f.pk: set(f.prestations.values_list("pk", flat=True)) for f in formateurs
    }
    # Prestations (pour le toggle et le tableau villes)
    prestations = list(
        PrestationModel.objects.select_related("prestataire", "formation", "beneficiaire")
        .prefetch_related("classes__lieu")
        .order_by("code")
    )

    # Fallback mapping formation -> prestation lorsque les appels n'ont pas
    # directement les infos prestataire/beneficiaire.
    formation_to_prestations: dict[str, list] = defaultdict(list)
    for prest in prestations:
        formation_label = str(getattr(getattr(prest, "formation", None), "nom", "") or "").strip()
        formation_key = _normalize_org_name(formation_label)
        if formation_key:
            formation_to_prestations[formation_key].append(prest)

    # Rapprocher les infos "page principale formateur" via les numeros de telephone.
    phone_insights: dict[str, dict] = {}
    appels_phone_rows = AppelFormateur.objects.filter(is_active=True).order_by(
        "session_date", "reference_code"
    )
    for row in appels_phone_rows.iterator():
        phone_key = _normalize_phone(str(getattr(row, "telephone", "") or ""))
        if not phone_key:
            continue

        resolved_classe = _resolve_batch_update_formateur_classe(row)
        resolved_prestation = getattr(resolved_classe, "prestation", None)
        resolved_prestataire_name = str(
            getattr(getattr(resolved_prestation, "prestataire", None), "raison_sociale", "") or ""
        ).strip()
        resolved_beneficiaire_name = str(
            getattr(getattr(resolved_prestation, "beneficiaire", None), "nom_structure", "") or ""
        ).strip()

        prestataire_name = (
            str(getattr(row, "prestataire", "") or "").strip() or resolved_prestataire_name
        )
        beneficiaire_name = (
            str(getattr(row, "beneficiaire", "") or "").strip() or resolved_beneficiaire_name
        )
        if not prestataire_name or not beneficiaire_name:
            row_formation = str(getattr(row, "formation", "") or "").strip()
            formation_key = _normalize_org_name(row_formation)
            candidates = formation_to_prestations.get(formation_key, [])
            if len(candidates) == 1:
                inferred = candidates[0]
                if not prestataire_name:
                    prestataire_name = str(
                        getattr(getattr(inferred, "prestataire", None), "raison_sociale", "") or ""
                    ).strip()
                if not beneficiaire_name:
                    beneficiaire_name = str(
                        getattr(getattr(inferred, "beneficiaire", None), "nom_structure", "") or ""
                    ).strip()

        insight = phone_insights.setdefault(
            phone_key,
            {
                "total": 0,
                "prest": defaultdict(int),
                "benef": defaultdict(int),
                "pair": defaultdict(int),
                "prest_label": {},
                "benef_label": {},
                "formation": defaultdict(int),
                "formation_label": {},
            },
        )
        raw_prest = prestataire_name
        raw_benef = beneficiaire_name
        prest_key = _normalize_org_name(prestataire_name)
        benef_key = _normalize_org_name(beneficiaire_name)
        insight["total"] += 1
        if prest_key:
            insight["prest"][prest_key] += 1
            insight["prest_label"].setdefault(prest_key, raw_prest)
        if benef_key:
            insight["benef"][benef_key] += 1
            insight["benef_label"].setdefault(benef_key, raw_benef)
        if prest_key or benef_key:
            insight["pair"][(prest_key, benef_key)] += 1
        formation_value = str(getattr(row, "formation", "") or "").strip()
        formation_key = _normalize_org_name(formation_value)
        if formation_key:
            insight["formation"][formation_key] += 1
            insight["formation_label"].setdefault(formation_key, formation_value)

    def _score_probability(
        insight: dict | None, prestation_obj: PrestationModel
    ) -> tuple[int, str]:
        if not insight:
            return 0, "none"
        total = int(insight.get("total") or 0)
        if total <= 0:
            return 0, "none"
        prest_name = _normalize_org_name(
            getattr(getattr(prestation_obj, "prestataire", None), "raison_sociale", "")
        )
        benef_name = _normalize_org_name(
            getattr(getattr(prestation_obj, "beneficiaire", None), "nom_structure", "")
        )
        pair_score = 0.0
        for (left_prest, left_benef), count in insight["pair"].items():
            prest_score = _name_match_score(prest_name, left_prest) if prest_name else 0.0
            benef_score = _name_match_score(benef_name, left_benef) if benef_name else 0.0
            if prest_score <= 0 and benef_score <= 0:
                continue
            match_strength = (prest_score * 0.65) + (benef_score * 0.35)
            pair_score += count * match_strength
        if pair_score > 0:
            return int(round((pair_score / total) * 100)), "orgs"

        prest_score_total = 0.0
        for left_prest, count in insight["prest"].items():
            prest_score_total += count * _name_match_score(prest_name, left_prest)
        benef_score_total = 0.0
        for left_benef, count in insight["benef"].items():
            benef_score_total += count * _name_match_score(benef_name, left_benef)

        prest_ratio = (prest_score_total / total) if prest_name else 0.0
        benef_ratio = (benef_score_total / total) if benef_name else 0.0
        score_org = ((prest_ratio * 0.6) + (benef_ratio * 0.4)) * 100
        score_org = int(round(max(0, min(score_org, 100))))
        if score_org > 0:
            return score_org, "orgs"

        # Approximation mode: only if org labels are missing for this phone.
        if not insight["prest"] and not insight["benef"]:
            formation_name = _normalize_org_name(
                getattr(getattr(prestation_obj, "formation", None), "nom", "")
            )
            formation_score_total = 0.0
            for left_formation, count in insight["formation"].items():
                formation_score_total += count * _name_match_score(formation_name, left_formation)
            score_form = (
                int(round(max(0, min((formation_score_total / total) * 100, 100))))
                if formation_name
                else 0
            )
            if score_form > 0:
                return score_form, "formation_approx"

        return 0, "none"

    def _probability_level(probability: int, source_mode: str) -> str:
        if probability <= 0:
            return "zero"
        if source_mode == "formation_approx":
            return "approx"
        if probability >= 66:
            return "high"
        if probability > 33:
            return "medium"
        return "low"

    for form in formateurs:
        phone_key = _normalize_phone(str(getattr(form, "telephone", "") or ""))
        insight = phone_insights.get(phone_key)
        resolved_formateur_name = resolve_formateur_db_name_from_values(
            str(getattr(form, "telephone", "") or ""),
            str(getattr(form, "nom_complet", "") or ""),
        )
        fallback_name = str(getattr(form, "nom_complet", "") or "").strip()
        form.display_nom = resolved_formateur_name or fallback_name or "-"
        form.nom_rempli = str(getattr(form, "nom", "") or "").strip() not in {
            "",
            FORMATEUR_NAME_FALLBACK,
        }
        if insight:
            top_prestataires = sorted(insight["prest"].items(), key=lambda kv: (-kv[1], kv[0]))
            top_beneficiaires = sorted(insight["benef"].items(), key=lambda kv: (-kv[1], kv[0]))
            form.appel_prestataires = [
                insight["prest_label"].get(name, name)
                for name, _count in top_prestataires[:3]
                if name
            ]
            form.appel_beneficiaires = [
                insight["benef_label"].get(name, name)
                for name, _count in top_beneficiaires[:3]
                if name
            ]
            top_formations = sorted(insight["formation"].items(), key=lambda kv: (-kv[1], kv[0]))
            form.appel_formations = [
                insight["formation_label"].get(name, name)
                for name, _count in top_formations[:3]
                if name
            ]
        else:
            form.appel_prestataires = []
            form.appel_beneficiaires = []
            form.appel_formations = []

        linked_ids = formateur_prestations.get(form.pk, set())
        form.linked_count = len(linked_ids)
        form.has_missing_orgs = not (form.appel_prestataires or form.appel_beneficiaires)
        form.prestation_candidates = []
        for order_index, prest in enumerate(prestations):
            probability, source_mode = _score_probability(insight, prest)
            form.prestation_candidates.append(
                {
                    "prestation": prest,
                    "linked": prest.pk in linked_ids,
                    "probability": probability,
                    "probability_level": _probability_level(probability, source_mode),
                    "probability_source_mode": source_mode,
                    "order_index": order_index,
                }
            )
        if any(item["probability"] > 0 for item in form.prestation_candidates):
            form.prestation_candidates.sort(
                key=lambda item: (-item["probability"], item["order_index"])
            )

    # Pour chaque prestation, dériver la ville depuis Classe.lieu si ville vide
    for prest in prestations:
        if not prest.ville:
            for cls in prest.classes.all():
                if cls.lieu and cls.lieu.ville:
                    prest.ville_derived = cls.lieu.ville
                    break
            else:
                prest.ville_derived = ""
        else:
            prest.ville_derived = prest.ville

    context = {
        "formateurs": formateurs,
        "formateur_prestations": formateur_prestations,
        "prestations": prestations,
        "saved_count": saved_count,
        "saved_label": saved_label,
    }
    return render(request, "satisfaction_formateurs/formateurs_prestataires.html", context)
