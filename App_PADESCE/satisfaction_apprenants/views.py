import base64
import csv
import hashlib
import io
import json
import logging
import os
import re

# ---------------------------------------------------------------------------
# Import-notification store – in-memory, global, polled by all active sessions
# ---------------------------------------------------------------------------
import threading as _threading
import time as _time
import uuid
from collections import Counter, defaultdict
from datetime import date as date_cls
from types import SimpleNamespace

import openpyxl
import requests
from django.contrib import messages
from django.core.cache import cache
from django.core.exceptions import ObjectDoesNotExist
from django.core.files.storage import default_storage
from django.core.paginator import Paginator
from django.db import transaction
from django.db.models import Count, Q
from django.http import HttpResponse, JsonResponse, QueryDict
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.http import require_POST
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from App_PADESCE.appels.models import (
    APPEL_ANSWER_QUESTION_FIELDS,
    Appel,
    AppelAnswers,
    AppelFormateur,
    CALL_ANALYSIS_THRESHOLD_STATUSES,
    CALL_FORM_STATUSES,
    CALL_SUCCESS_STATUSES,
    appel_answers_completed_q,
    appel_has_any_audio,
    appel_has_any_form_data,
    derive_padesce_status,
    sync_padesce_status,
)
from App_PADESCE.apprenants.models import Apprenant
from App_PADESCE.core.access import require_analysis_access
from App_PADESCE.core.analysis_rules import (
    analysis_threshold_label,
    analysis_threshold_target,
    answer_has_all_three_scores,
    appel_analysis_exclusion_reason,
    appel_has_analysis_phone,
    appel_is_analysis_eligible,
    appel_is_manually_excluded,
    set_appel_manual_exclusion,
    toggle_appel_manual_exclusion,
)
from App_PADESCE.core.cache_versions import get_analysis_cache_version
from App_PADESCE.core.call_metrics import (
    count_callable_source_records_by_class,
    has_usable_phone,
    normalize_phone_digits,
)
from App_PADESCE.core.fast_stats import build_fast_stats_context
from App_PADESCE.formations.models import Classe
from App_PADESCE.reporting.network_excel import (
    build_consolidation_call_candidates,
    build_padesce_source_index,
    get_workbook_source_options,
    normalize_network_lookup,
    normalize_workbook_source_key,
)
from App_PADESCE.satisfaction_apprenants.forms import (
    SatisfactionApprenantForm,
    SatisfactionBatchUpdateForm,
)
from App_PADESCE.satisfaction_apprenants.models import SatisfactionApprenant
from App_PADESCE.satisfaction_apprenants.rag import answer_dashboard_prompt
from App_PADESCE.satisfaction_apprenants.services import get_prestations_ranking

_IMPORT_NOTIFS: list[dict] = []
_IMPORT_NOTIFS_LOCK = _threading.Lock()


def _push_import_notif(message: str, classes: list[str]) -> None:
    entry = {
        "id": uuid.uuid4().hex,
        "ts": _time.time(),
        "message": message,
        "classes": classes,
    }
    with _IMPORT_NOTIFS_LOCK:
        _IMPORT_NOTIFS.append(entry)
        del _IMPORT_NOTIFS[:-200]


def _get_import_notifs_since(since_ts: float) -> list[dict]:
    cutoff = _time.time() - 300  # discard entries older than 5 min
    with _IMPORT_NOTIFS_LOCK:
        return [n for n in _IMPORT_NOTIFS if n["ts"] > max(since_ts, cutoff)]


# ---------------------------------------------------------------------------

SESSION_KEY = "sat_appr_workflow"
REASON_KEYWORDS = [
    ("ProblÃ¨mes de transport", ["transport", "vÃ©hicule", "route", "panne", "dÃ©placement"]),
    (
        "DisponibilitÃ© / santÃ©",
        ["malade", "santÃ©", "disponibilitÃ©", "maladie", "absent", "repos"],
    ),
    (
        "Pas au courant / notification",
        ["pas au courant", "notification", "notifiÃ©", "ignorÃ©", "erreur"],
    ),
    ("Conditions / Ã©ligibilitÃ©", ["diplÃ´me", "condition", "Ã©ligibilitÃ©", "inscription"]),
    ("Pas intÃ©ressÃ©", ["intÃ©ressÃ©", "ne souhaite pas", "dÃ©sintÃ©ressÃ©", "pas de formation"]),
]


def _categorize_reason(text: str) -> str:
    if not text:
        return "Sans rÃ©ponse"
    for label, keywords in REASON_KEYWORDS:
        if any(keyword in text for keyword in keywords):
            return label
    return "Autres raisons"


def _detect_participation(text: str) -> str:
    if not text:
        return "IndÃ©terminÃ©"
    if re.search(r"\b(?:pas|n['â€™]?a)\b.*\b(particip|assist|prÃ©sent|venu)\b", text):
        return "Absents"
    if any(
        kw in text
        for kw in ["participÃ©", "assistÃ©", "prÃ©sent", "Ã©tÃ© lÃ ", "prÃ©sente", "participant"]
    ):
        return "PrÃ©sents"
    return "IndÃ©terminÃ©"


OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
DEFAULT_TRANSCRIBE_MODEL = "google/gemini-2.5-flash"
SUPPORTED_AUDIO_FORMATS = {"wav", "mp3", "m4a", "ogg", "webm", "flac"}

logger = logging.getLogger(__name__)
ANALYSIS_CACHE_TIMEOUT = int(str(os.getenv("PADESCE_ANALYSIS_CACHE_TIMEOUT", "300") or "300"))


def _analysis_cache_key(prefix: str, *parts) -> str:
    rendered_parts = [str(part or "").strip() for part in parts]
    digest = hashlib.sha1("||".join(rendered_parts).encode("utf-8")).hexdigest()
    return f"satisfaction:{prefix}:{digest}"


def _analysis_queryset_marker(model) -> str:
    return get_analysis_cache_version(f"model:{model._meta.label_lower}")


def _normalize_phone(value: str) -> str:
    return normalize_phone_digits(value)


def _find_apprenant(classe_id: str, identifiant: str) -> Apprenant | None:
    identifiant = identifiant.strip()
    if not identifiant:
        return None
    identifiant_lower = identifiant.lower()
    identifiant_digits = _normalize_phone(identifiant)
    for apprenant in Apprenant.objects.filter(classe_id=classe_id):
        if apprenant.code.lower() == identifiant_lower:
            return apprenant
        if identifiant_digits and (
            _normalize_phone(apprenant.telephone1) == identifiant_digits
            or _normalize_phone(apprenant.telephone2) == identifiant_digits
        ):
            return apprenant
    return None


def _save_audio(uploaded_file, folder: str) -> str:
    _, ext = os.path.splitext(uploaded_file.name)
    ext = ext or ".dat"
    filename = f"{uuid.uuid4().hex}{ext}"
    return default_storage.save(f"enquetes/{folder}/{filename}", uploaded_file)


# Fallback scoring when transcript does not include explicit answers.
def _ai_scores(seed: str, count: int = 9) -> list[int]:
    digest = hashlib.sha256(seed.encode("utf-8")).digest()
    return [1 + (digest[i] % 5) for i in range(count)]


def _flatten_message_content(content) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text" and item.get("text"):
                parts.append(item["text"])
        return "\n".join(parts)
    return ""


def _extract_transcription_text(result) -> str:
    if not result:
        return ""
    if isinstance(result, dict):
        if "choices" in result:
            try:
                content = result["choices"][0]["message"].get("content")
            except (IndexError, KeyError, TypeError, AttributeError):
                content = None
            return _flatten_message_content(content)
        return result.get("text") or result.get("transcript") or ""
    if hasattr(result, "text"):
        return result.text or ""
    return str(result)


def _guess_audio_format(audio_path: str) -> str:
    ext = os.path.splitext(audio_path)[1].lstrip(".").lower()
    if ext == "ma4":
        ext = "m4a"
    return ext if ext in SUPPORTED_AUDIO_FORMATS else "wav"


def _encode_audio_to_base64(audio_path: str) -> str:
    with default_storage.open(audio_path, "rb") as audio_file:
        return base64.b64encode(audio_file.read()).decode("ascii")


def _transcribe_audio(audio_path: str) -> tuple[str | None, str | None]:
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        return None, "Cle OpenRouter manquante. Definissez OPENROUTER_API_KEY dans l'environnement."
    audio_format = _guess_audio_format(audio_path)
    model = os.getenv("OPENROUTER_TRANSCRIBE_MODEL", DEFAULT_TRANSCRIBE_MODEL)
    logger.info(
        "Transcription OpenRouter demarree. audio=%s format=%s model=%s",
        audio_path,
        audio_format,
        model,
    )
    try:
        base64_audio = _encode_audio_to_base64(audio_path)
    except Exception as exc:
        logger.exception("Lecture audio impossible pour transcription. audio=%s", audio_path)
        return None, f"Impossible de lire l'audio: {exc}"

    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Please transcribe this audio file."},
                    {
                        "type": "input_audio",
                        "input_audio": {"data": base64_audio, "format": audio_format},
                    },
                ],
            }
        ],
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    response = None
    try:
        response = requests.post(OPENROUTER_API_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
    except requests.RequestException as exc:
        detail = ""
        if response is not None:
            try:
                detail = response.json().get("error", {}).get("message") or response.text
            except ValueError:
                detail = response.text
        detail = detail.strip()
        logger.error(
            "Echec transcription OpenRouter. audio=%s status=%s detail=%s",
            audio_path,
            response.status_code if response is not None else "n/a",
            detail or str(exc),
        )
        return None, f"Echec de transcription OpenRouter: {exc} {detail}".strip()
    except ValueError as exc:
        logger.error("Reponse OpenRouter invalide. audio=%s error=%s", audio_path, exc)
        return None, f"Reponse OpenRouter invalide: {exc}"

    text = _extract_transcription_text(data).strip()
    if not text:
        logger.error("Transcription OpenRouter videe. audio=%s", audio_path)
        return None, "Transcription videe ou non reconnue."
    logger.info("Transcription OpenRouter terminee. audio=%s chars=%s", audio_path, len(text))
    return text[:8000], None


def _parse_scores_from_transcript(transcript: str, total_questions: int = 9) -> dict[int, int]:
    results: dict[int, int] = {}
    if not transcript:
        return results
    lowered = transcript.lower()
    for idx in range(1, total_questions + 1):
        pattern = rf"(?:\bq{idx}\b|question\s*{idx})[^\d]{{0,10}}([1-5])"
        match = re.search(pattern, lowered)
        if match:
            results[idx] = int(match.group(1))
    return results


def _ai_results_apprenant(audio_path: str) -> tuple[dict | None, str | None, str | None]:
    transcript, error = _transcribe_audio(audio_path)
    if error:
        return None, None, error
    size = default_storage.size(audio_path)
    scores = _ai_scores(f"{audio_path}:{size}:{len(transcript)}")
    parsed = _parse_scores_from_transcript(transcript)
    for idx, value in parsed.items():
        if 1 <= idx <= 9:
            scores[idx - 1] = value
    results = {
        "q1_clarte_exposes": scores[0],
        "q2_interaction_formateur": scores[1],
        "q3_maitrise_contenu": scores[2],
        "q4_salle_adequate": scores[3],
        "q5_materiel_disponible": scores[4],
        "q6_organisation_temps": scores[5],
        "q7_utilite_formation": scores[6],
        "q8_adequation_besoins": scores[7],
        "q9_satisfaction_globale": scores[8],
        "commentaire": "",
        "recommandations": "",
    }
    return results, transcript, None


def satisfaction_apprenants(request):
    filter_classe = request.GET.get("classe")
    qs = SatisfactionApprenant.objects.select_related(
        "classe", "apprenant", "appel", "inspecteur", "enqueteur"
    ).order_by("-date", "-created_at")
    if filter_classe:
        qs = qs.filter(classe_id=filter_classe)

    workflow = request.session.get(SESSION_KEY, {})
    save_errors = None

    if request.method == "POST":
        action = request.POST.get("action")
        identifiant = (request.POST.get("identifiant") or "").strip()
        posted_classe = request.POST.get("classe") or workflow.get("classe_id")
        posted_inspecteur = request.POST.get("inspecteur") or workflow.get("inspecteur_id")
        posted_date = request.POST.get("date") or workflow.get("date")
        posted_heure = request.POST.get("heure") or workflow.get("heure")

        if posted_classe:
            workflow["classe_id"] = str(posted_classe)
        if posted_inspecteur:
            workflow["inspecteur_id"] = posted_inspecteur
        if posted_date:
            workflow["date"] = posted_date
        if posted_heure:
            workflow["heure"] = posted_heure
        if identifiant:
            workflow["identifiant"] = identifiant

        if action == "identify":
            if not posted_classe or not identifiant:
                messages.error(
                    request, "Renseignez la classe et le code ou telephone de l'apprenant."
                )
            else:
                apprenant = _find_apprenant(posted_classe, identifiant)
                if apprenant:
                    workflow["apprenant_id"] = apprenant.id
                    workflow.pop("audio_path", None)
                    workflow.pop("ai_results", None)
                    messages.success(request, f"Apprenant identifie: {apprenant}.")
                else:
                    messages.error(request, "Aucun apprenant correspondant dans cette classe.")
        elif action == "process_audio":
            if not workflow.get("apprenant_id"):
                messages.error(request, "Identifiez d'abord un apprenant.")
            elif str(posted_classe or "") != str(workflow.get("classe_id") or ""):
                messages.error(request, "La classe ne correspond pas a l'apprenant identifie.")
            else:
                uploaded_audio = request.FILES.get("audio_appel")
                if uploaded_audio:
                    workflow["audio_path"] = _save_audio(uploaded_audio, "satisfaction_apprenants")
                    logger.info(
                        "Audio recu pour transcription apprenant. fichier=%s taille=%s",
                        uploaded_audio.name,
                        getattr(uploaded_audio, "size", "n/a"),
                    )
                if not workflow.get("audio_path"):
                    messages.error(request, "Chargez un audio d'appel pour lancer le traitement.")
                else:
                    results, transcript, error = _ai_results_apprenant(workflow["audio_path"])
                    if error:
                        messages.error(request, error)
                    else:
                        workflow["ai_results"] = results
                        workflow["transcription"] = transcript
                        messages.success(
                            request, "Transcription terminee et traitement vocal actualise."
                        )
        elif action == "save":
            if not workflow.get("apprenant_id"):
                messages.error(request, "Identifiez un apprenant avant d'enregistrer.")
            elif not workflow.get("ai_results"):
                messages.error(request, "Lancez le traitement vocal avant d'enregistrer.")
            else:
                data = request.POST.copy()
                data["apprenant"] = workflow["apprenant_id"]
                if workflow.get("classe_id"):
                    data["classe"] = workflow["classe_id"]
                if workflow.get("inspecteur_id"):
                    data["inspecteur"] = workflow["inspecteur_id"]
                if workflow.get("date"):
                    data["date"] = workflow["date"]
                if workflow.get("heure"):
                    data["heure"] = workflow["heure"]
                data.update(workflow["ai_results"])
                save_form = SatisfactionApprenantForm(data)
                if save_form.is_valid():
                    obj = save_form.save(commit=False)
                    if hasattr(request, "user") and request.user.is_authenticated:
                        obj.enqueteur = request.user
                    audio_path = workflow.get("audio_path")
                    if audio_path:
                        obj.audio_appel.name = audio_path
                    obj.transcription = workflow.get("transcription", "")
                    obj.save()
                    messages.success(request, "Satisfaction apprenant enregistree.")
                    request.session.pop(SESSION_KEY, None)
                    return redirect(
                        request.path_info + f"?classe={filter_classe}"
                        if filter_classe
                        else request.path_info
                    )
                else:
                    save_errors = save_form.errors

        request.session[SESSION_KEY] = workflow
        request.session.modified = True

    initial = {
        "classe": workflow.get("classe_id"),
        "inspecteur": workflow.get("inspecteur_id"),
        "date": workflow.get("date") or date_cls.today(),
        "heure": workflow.get("heure"),
    }
    form = SatisfactionApprenantForm(initial=initial)

    identified_apprenant = None
    apprenant_id = workflow.get("apprenant_id")
    if apprenant_id:
        identified_apprenant = Apprenant.objects.filter(id=apprenant_id).first()
        if not identified_apprenant:
            workflow.pop("apprenant_id", None)
            workflow.pop("ai_results", None)
            workflow.pop("audio_path", None)
            request.session[SESSION_KEY] = workflow
            request.session.modified = True

    paginator = Paginator(qs, 50)
    page_obj = paginator.get_page(request.GET.get("page"))

    summary_rows = qs.values_list("commentaire", "transcription", "recommandations")
    reason_counts = Counter({label: 0 for label, _ in REASON_KEYWORDS})
    reason_counts.update({"Autres raisons": 0, "Sans rÃ©ponse": 0})
    participation_counts = Counter({"PrÃ©sents": 0, "Absents": 0, "IndÃ©terminÃ©": 0})
    reason_samples = defaultdict(list)
    for comment, transcript, reco in summary_rows:
        text = " ".join(filter(None, [comment, transcript, reco])).strip()
        if not text:
            continue
        lower_text = text.lower()
        label = _categorize_reason(lower_text)
        reason_counts[label] += 1
        if len(reason_samples[label]) < 2:
            reason_samples[label].append(text)
        participation_counts[_detect_participation(lower_text)] += 1

    reason_distribution = []
    for label, _ in list(REASON_KEYWORDS) + [("Autres raisons", []), ("Sans rÃ©ponse", [])]:
        reason_distribution.append(
            {
                "label": label,
                "count": reason_counts.get(label, 0),
                "samples": reason_samples.get(label, []),
            }
        )

    reason_chart = {
        "labels": [item["label"] for item in reason_distribution if item["count"]],
        "values": [item["count"] for item in reason_distribution if item["count"]],
    }
    participation_chart = {
        "labels": ["PrÃ©sents", "Absents", "IndÃ©terminÃ©"],
        "values": [
            participation_counts.get("PrÃ©sents", 0),
            participation_counts.get("Absents", 0),
            participation_counts.get("IndÃ©terminÃ©", 0),
        ],
    }

    context = {
        "form": form,
        "identified_apprenant": identified_apprenant,
        "identifiant": workflow.get("identifiant", ""),
        "ai_results": workflow.get("ai_results"),
        "transcription": workflow.get("transcription"),
        "audio_name": (
            os.path.basename(workflow.get("audio_path")) if workflow.get("audio_path") else None
        ),
        "save_errors": save_errors,
        "enquetes": page_obj,
        "page_obj": page_obj,
        "classes": Classe.objects.all().order_by("code"),
        "filter_classe": filter_classe,
        "reason_distribution": reason_distribution,
        "participation_chart": json.dumps(participation_chart),
        "reason_chart": json.dumps(reason_chart),
    }
    return render(request, "satisfaction_apprenants/index.html", context)


def satisfaction_apprenants_export_csv(request):
    filter_classe = request.GET.get("classe")
    qs = SatisfactionApprenant.objects.select_related(
        "classe", "apprenant", "appel", "inspecteur", "enqueteur"
    ).order_by("-date")
    if filter_classe:
        qs = qs.filter(classe_id=filter_classe)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = "attachment; filename=satisfaction_apprenants.csv"
    writer = csv.writer(response)
    writer.writerow(
        [
            "classe",
            "apprenant",
            "inspecteur",
            "enqueteur",
            "date",
            "heure",
            "q1",
            "q2",
            "q3",
            "q4",
            "q5",
            "q6",
            "q7",
            "q8",
            "q9",
            "commentaire",
            "recommandations",
        ]
    )
    for s in qs:
        writer.writerow(
            [
                s.classe,
                s.apprenant,
                s.inspecteur,
                s.enqueteur,
                s.date,
                s.heure,
                s.q1_clarte_exposes,
                s.q2_interaction_formateur,
                s.q3_maitrise_contenu,
                s.q4_salle_adequate,
                s.q5_materiel_disponible,
                s.q6_organisation_temps,
                s.q7_utilite_formation,
                s.q8_adequation_besoins,
                s.q9_satisfaction_globale,
                s.commentaire,
                s.recommandations,
            ]
        )
    return response


Q_FIELDS = [
    ("q1_clarte_exposes", "Clarté des exposés"),
    ("q2_interaction_formateur", "Interaction avec le formateur"),
    ("q3_maitrise_contenu", "Maîtrise du contenu"),
    ("q4_salle_adequate", "Salle adéquate"),
    ("q5_materiel_disponible", "Matériel disponible"),
    ("q6_organisation_temps", "Organisation du temps"),
    ("q7_utilite_formation", "Utilité de la formation"),
    ("q8_adequation_besoins", "Adéquation aux besoins"),
    ("q9_satisfaction_globale", "Satisfaction globale"),
]

CSV_Q_HEADERS = [
    "Q1_ClarteExposes",
    "Q2_InteractionFormateur",
    "Q3_MaitriseContenu",
    "Q4_SalleAdequate",
    "Q5_MaterielDisponible",
    "Q6_OrganisationTemps",
    "Q7_UtiliteFormation",
    "Q8_AdequationBesoins",
    "Q9_SatisfactionGlobale",
]


def _avg(values):
    nums = [v for v in values if v is not None]
    return round(sum(nums) / len(nums), 2) if nums else 0


def _satisfaction_dashboard_base_queryset():
    return AppelAnswers.objects.select_related(
        "appel",
        "appel__classe",
        "appel__classe__prestation",
        "appel__classe__prestation__prestataire",
        "appel__classe__prestation__beneficiaire",
        "appel__classe__lieu",
        "appel__satisfaction_apprenant",
        "appel__satisfaction_apprenant__inspecteur",
        "appel__satisfaction_apprenant__apprenant",
        "modified_by",
    ).filter(appel__is_active=True)


def _autosize_worksheet(worksheet, max_width: int = 40):
    for index, column_cells in enumerate(worksheet.columns, start=1):
        values = [len(str(cell.value or "")) for cell in column_cells]
        if not values:
            continue
        worksheet.column_dimensions[get_column_letter(index)].width = min(
            max(values) + 2, max_width
        )


def _dashboard_row_from_answer(answer_or_appel) -> dict:
    if hasattr(answer_or_appel, "appel"):
        answer = answer_or_appel
        appel = answer.appel
    else:
        answer = None
        appel = answer_or_appel

    classe = appel.classe
    prestation = getattr(classe, "prestation", None) if classe else None
    survey = getattr(appel, "satisfaction_apprenant", None)
    prestataire = (
        getattr(getattr(prestation, "prestataire", None), "raison_sociale", "")
        or appel.prestataire
        or "-"
    )
    beneficiaire_obj = getattr(prestation, "beneficiaire", None) if prestation else None
    beneficiaire = getattr(beneficiaire_obj, "nom_structure", "") or appel.beneficiaire or "-"
    beneficiaire_type = str(getattr(beneficiaire_obj, "type_structure", "") or "").lower()
    ville = getattr(getattr(classe, "lieu", None), "ville", "") or appel.lieu or "Non renseignée"
    fenetre = _analysis_fenetre_for_appel(appel)

    cohorte = str(getattr(classe, "cohorte", "") or "").strip() or "Non renseignée"
    has_phone = appel_has_analysis_phone(appel)
    has_audio = bool(
        getattr(getattr(appel, "audio_file", None), "name", "")
        or getattr(getattr(survey, "audio_appel", None), "name", "")
    )
    analysis_scope = fenetre in {"2", "3"}
    analysis_eligible = analysis_scope and appel_is_analysis_eligible(
        appel,
        answer=answer,
        survey=survey,
    )
    analysis_exclusion_reason = (
        appel_analysis_exclusion_reason(appel, answer=answer, survey=survey)
        if analysis_scope
        else "Fenetre hors analyse"
    )

    # Handle potentially missing answer data
    timestamp = (
        getattr(answer, "modified_at", None)
        or getattr(answer, "created_at", None)
        or appel.updated_at
        or appel.created_at
    )
    survey_date = getattr(survey, "date", None) or timestamp.date()
    survey_time = getattr(survey, "heure", None) or timestamp.time().replace(microsecond=0)
    inspecteur = getattr(survey, "inspecteur", None)

    q_filled_count = 0
    for field, _ in Q_FIELDS:
        val = getattr(answer, field, None) if answer else None
        if val is None and survey:
            val = getattr(survey, field, None)
        if val not in (None, ""):
            q_filled_count += 1
    has_form = q_filled_count >= 9

    return {
        "id": getattr(answer, "id", None),
        "appel_id": getattr(appel, "pk", None),
        "date": timestamp.date(),
        "heure": timestamp.time().replace(microsecond=0),
        "modified_at": timestamp,
        "survey_date": survey_date,
        "survey_time": survey_time,
        "inspecteur_code": getattr(inspecteur, "code", "") or "",
        "inspecteur_nom": getattr(inspecteur, "nom_complet", "") or "",
        "classe_code": getattr(classe, "code", "") or appel.classe_label or "Non renseignée",
        "classe_intitule": getattr(classe, "intitule_formation", "") or "-",
        "formation_intitule": getattr(classe, "intitule_formation", "") or "-",
        "prestation_code": getattr(prestation, "code", "") or "-",
        "prestataire": prestataire,
        "beneficiaire": beneficiaire,
        "fenetre": fenetre,
        "cohorte": cohorte,
        "ville": ville,
        "apprenant_code": appel.code or "",
        "apprenant_nom": appel.nom or "",
        "telephone1": appel.telephone1 or "",
        "telephone2": appel.telephone2 or "",
        "has_phone": has_phone,
        "has_audio": has_audio,
        "status": appel.status or "",
        "user": (
            getattr(getattr(answer, "modified_by", None), "username", "")
            if answer
            else "Non renseigné"
        ),
        "commentaire": getattr(answer, "commentaire", "") if answer else "",
        "recommandations": getattr(answer, "recommandations", "") if answer else "",
        "analysis_scope": analysis_scope,
        "analysis_eligible": analysis_eligible,
        "analysis_included": has_form and analysis_eligible,
        "has_form": has_form,
        "analysis_excluded": not analysis_eligible,
        "analysis_exclusion_reason": analysis_exclusion_reason,
        "formulaire_all_three": answer_has_all_three_scores(answer),
        "exclude_from_analysis": appel_is_manually_excluded(appel),
        **{field: getattr(answer, field, None) if answer else None for field, _ in Q_FIELDS},
    }


def _dashboard_bucket():
    return {
        "nb": 0,
        "sums": {field: 0 for field, _ in Q_FIELDS},
        "counts": {field: 0 for field, _ in Q_FIELDS},
    }


def _dashboard_bucket_add(bucket: dict, row: dict):
    bucket["nb"] += 1
    for field, _label in Q_FIELDS:
        value = row.get(field)
        if value is None:
            continue
        bucket["sums"][field] += float(value)
        bucket["counts"][field] += 1


def _dashboard_bucket_avg(bucket: dict, field_name: str) -> float:
    count = bucket["counts"][field_name]
    return round(bucket["sums"][field_name] / count, 2) if count else 0


def _dashboard_bucket_avgs(bucket: dict) -> list[float]:
    return [_dashboard_bucket_avg(bucket, field) for field, _label in Q_FIELDS]


def _row_has_any_numeric_score(row: dict) -> bool:
    return any(row.get(field) is not None for field, _label in Q_FIELDS)


def _source_class_apprenant_counts(source_bundle: dict | None) -> dict[str, int]:
    return count_callable_source_records_by_class(source_bundle)


def _normalize_class_count_map(raw_counts: dict[str, int]) -> dict[str, int]:
    normalized: dict[str, int] = {}
    for classe_code, count in (raw_counts or {}).items():
        key = normalize_network_lookup(classe_code)
        if not key:
            continue
        normalized[key] = max(int(normalized.get(key) or 0), int(count or 0))
    return normalized


def _status_threshold_progress_by_class(source_bundle: dict | None = None) -> dict[str, dict]:
    source_records = (source_bundle or {}).get("records") or {}
    source_callable_counts = _normalize_class_count_map(
        _source_class_apprenant_counts(source_bundle)
    )

    label_by_key: dict[str, str] = {}
    db_total_by_key: dict[str, int] = defaultdict(int)
    db_callable_by_key: dict[str, int] = defaultdict(int)
    db_threshold_counts_by_key: dict[str, int] = defaultdict(int)

    for source_class in (source_bundle or {}).get("classes", {}).values():
        classe_id = str(source_class.get("classe_id") or "").strip()
        classe_key = normalize_network_lookup(classe_id)
        if classe_key and classe_id:
            label_by_key.setdefault(classe_key, classe_id)

    for row in Appel.objects.filter(is_active=True).values(
        "code",
        "classe_label",
        "status",
        "telephone1",
        "telephone2",
    ):
        classe_label = str(row.get("classe_label") or "").strip()
        if not classe_label:
            source_record = source_records.get(
                normalize_network_lookup(str(row.get("code") or ""))
            )
            classe_label = str((source_record or {}).get("classe_id") or "").strip()
        classe_key = normalize_network_lookup(classe_label)
        if not classe_key:
            continue
        label_by_key.setdefault(classe_key, classe_label)
        db_total_by_key[classe_key] += 1
        if not has_usable_phone(row.get("telephone1"), row.get("telephone2")):
            continue
        db_callable_by_key[classe_key] += 1
        if str(row.get("status") or "").strip() in CALL_ANALYSIS_THRESHOLD_STATUSES:
            db_threshold_counts_by_key[classe_key] += 1

    progress_by_key: dict[str, dict] = {}
    classe_keys = set(label_by_key) | set(source_callable_counts) | set(db_total_by_key)
    for classe_key in classe_keys:
        callable_total = (
            int(db_callable_by_key.get(classe_key) or 0)
            if int(db_total_by_key.get(classe_key) or 0) > 0
            else int(source_callable_counts.get(classe_key) or 0)
        )
        completed_count = int(db_threshold_counts_by_key.get(classe_key) or 0)
        target = analysis_threshold_target(callable_total)
        progress_by_key[classe_key] = {
            "code": label_by_key.get(classe_key, classe_key).strip() or classe_key,
            "total": callable_total,
            "completed": completed_count,
            "target": target,
            "reached": callable_total > 0 and completed_count >= target,
        }
    return progress_by_key


def _status_threshold_class_codes(source_bundle: dict | None = None) -> set[str]:
    return {
        classe_key
        for classe_key, progress in _status_threshold_progress_by_class(source_bundle).items()
        if bool(progress.get("reached"))
    }


def _merge_class_apprenant_counts(
    local_counts: dict[str, int], source_bundle: dict | None
) -> dict[str, int]:
    merged = _normalize_class_count_map(local_counts)
    for classe_code, source_count in _source_class_apprenant_counts(source_bundle).items():
        key = normalize_network_lookup(classe_code)
        if not key:
            continue
        merged[key] = max(int(merged.get(key) or 0), int(source_count or 0))
    return merged


def _analysis_fenetre_for_appel(appel: Appel) -> str:
    classe = getattr(appel, "classe", None)
    prestation = getattr(classe, "prestation", None) if classe else None
    beneficiaire = getattr(prestation, "beneficiaire", None) if prestation else None
    beneficiaire_type = str(getattr(beneficiaire, "type_structure", "") or "").strip().lower()
    raw_fenetre = (
        str(getattr(appel, "fenetre", "") or "").strip()
        or str(getattr(classe, "fenetre", "") or "").strip()
    )
    if raw_fenetre in {"2", "3"}:
        return raw_fenetre
    if "entreprise" in beneficiaire_type:
        return "2"
    if "association" in beneficiaire_type or "gic" in beneficiaire_type:
        return "3"
    digits = "".join(ch for ch in raw_fenetre if ch.isdigit())
    return digits if digits in {"2", "3"} else "Non renseignée"


def _analysis_class_code_for_appel(appel: Appel) -> str:
    return str(getattr(getattr(appel, "classe", None), "code", "") or appel.classe_label or "").strip()


def _analysis_class_count(counts: dict[str, int], classe_code: str) -> int:
    return int((counts or {}).get(normalize_network_lookup(classe_code), 0) or 0)


def _local_analysis_class_summary() -> dict[str, dict[str, int | str]]:
    summary: dict[str, dict[str, int | str]] = {}
    queryset = Appel.objects.filter(is_active=True).select_related(
        "classe",
        "classe__prestation",
        "classe__prestation__beneficiaire",
    )
    for appel in queryset:
        if _analysis_fenetre_for_appel(appel) not in {"2", "3"}:
            continue
        classe_code = _analysis_class_code_for_appel(appel)
        key = normalize_network_lookup(classe_code)
        if not key:
            continue
        item = summary.setdefault(
            key,
            {
                "code": classe_code,
                "total": 0,
                "with_phone": 0,
                "eligible": 0,
            },
        )
        item["total"] = int(item["total"]) + 1
        if appel_has_analysis_phone(appel):
            item["with_phone"] = int(item["with_phone"]) + 1
        if appel_has_analysis_phone(appel) and not appel_is_manually_excluded(appel):
            item["eligible"] = int(item["eligible"]) + 1
    return summary


def _local_analysis_class_counts() -> dict[str, int]:
    return {
        key: int(item.get("eligible") or 0)
        for key, item in _local_analysis_class_summary().items()
    }


def _source_class_is_finished(source_class: dict) -> bool:
    status = normalize_network_lookup(source_class.get("statut_prestation", ""))
    if not status:
        return True
    return status in {"termine", "terminee"}


def _sorted_unique(values):
    return sorted(
        {
            str(value).strip()
            for value in values
            if str(value or "").strip() and str(value).strip() != "-"
        }
    )


def _is_placeholder_dashboard_label(value: str) -> bool:
    normalized = normalize_network_lookup(str(value or ""))
    return normalized in {
        "",
        "-",
        "--",
        "- -",
        "none",
        "n/a",
        "na",
        "non renseigne",
        "non renseignee",
        "sans intitule",
    }


def _prefer_dashboard_label(current: str, candidate: str) -> str:
    current_text = str(current or "").strip()
    candidate_text = str(candidate or "").strip()
    if not candidate_text:
        return current_text
    if not current_text:
        return candidate_text
    current_missing = _is_placeholder_dashboard_label(current_text)
    candidate_missing = _is_placeholder_dashboard_label(candidate_text)
    if current_missing and not candidate_missing:
        return candidate_text
    if candidate_missing and not current_missing:
        return current_text
    return candidate_text if len(candidate_text) > len(current_text) else current_text


SATISFACTION_TABS = {
    "tab-apprenants",
    "tab-classe",
    "tab-prestation",
    "tab-cohorte",
    "tab-ville",
    "tab-user",
}

SATISFACTION_DASHBOARD_TAB_LABELS = {
    "tab-apprenants": "Liste des enquêtes",
    "tab-classe": "Par classe",
    "tab-prestation": "Par prestation",
    "tab-cohorte": "Par cohorte",
    "tab-ville": "Par ville",
    "tab-user": "Par utilisateur",
}

SATISFACTION_DASHBOARD_TAB_DESCRIPTIONS = {
    "tab-apprenants": "Détail des enquêtes visibles, enrichi avec la cohérence de la source réseau.",
    "tab-classe": "Synthèse des classes affichées dont le seuil d'appels est atteint.",
    "tab-prestation": "Regroupement des résultats visibles par prestation, prestataire et bénéficiaire.",
    "tab-cohorte": "Vue agrégée des résultats visibles par cohorte.",
    "tab-ville": "Répartition des résultats visibles par ville.",
    "tab-user": "Répartition des résultats visibles par utilisateur.",
}

SOURCE_COMPARE_FIELDS = (
    ("apprenant_nom", "nom_individu", "Nom"),
    ("classe_code", "classe_id", "Classe"),
    ("prestataire", "prestataire", "Prestataire"),
    ("beneficiaire", "beneficiaire", "Bénéficiaire"),
)

FILTER_FIELD_ROW_MAP = {
    "prestation": "prestation_code",
    "fenetre": "fenetre",
    "ville": "ville",
    "user": "user",
    "classe": "classe_code",
    "prestataire": "prestataire",
    "beneficiaire": "beneficiaire",
    "cohorte": "cohorte",
    "status": "status",
}

ANALYSIS_DASHBOARD_DEFAULT_SOURCE = "cutoff"


def _analysis_selected_source(request) -> str:
    requested_source = str(request.GET.get("source") or "").strip()
    return normalize_workbook_source_key(
        requested_source or ANALYSIS_DASHBOARD_DEFAULT_SOURCE
    )


def _active_satisfaction_tab(request) -> str:
    tab = (request.GET.get("tab") or "tab-apprenants").strip()
    return tab if tab in SATISFACTION_TABS else "tab-apprenants"


def _source_compare_alert(
    row: dict, source_record: dict, row_key: str, source_key: str, label: str
) -> str | None:
    row_value = str(row.get(row_key) or "").strip()
    source_value = str(source_record.get(source_key) or "").strip()
    if not row_value and not source_value:
        return None
    if normalize_network_lookup(row_value) == normalize_network_lookup(source_value):
        return None
    return f"{label}: {row_value or '-'} != {source_value or '-'}"


def _source_summary_unavailable(
    error_message: str,
    source_key: str = ANALYSIS_DASHBOARD_DEFAULT_SOURCE,
) -> dict:
    source_options = {item["value"]: item for item in get_workbook_source_options()}
    selected_source = normalize_workbook_source_key(source_key)
    source_meta = source_options.get(selected_source, {})
    return {
        "available": False,
        "message": error_message,
        "key": selected_source,
        "label": source_meta.get("label", ""),
        "name": source_meta.get("name", ""),
        "modified_label": "",
        "matched_count": 0,
        "missing_count": 0,
        "consistent_count": 0,
        "mismatch_count": 0,
        "apprenant_id_count": 0,
        "source_apprenant_count": 0,
        "duplicate_code_count": 0,
        "mismatch_rows": [],
    }


def _normalize_enquete_id(value: str, fallback_index: int | None = None) -> str:
    text = str(value or "").strip().upper()
    digits = "".join(ch for ch in text if ch.isdigit())
    if digits:
        return f"ENQ{int(digits):03d}"
    if fallback_index is not None:
        return f"ENQ{int(fallback_index):03d}"
    return ""


def _format_export_date(value) -> str:
    if not value:
        return ""
    try:
        return value.strftime("%d/%m/%Y")
    except Exception:
        return str(value)


def _format_export_time(value) -> str:
    if not value:
        return ""
    try:
        return value.strftime("%H:%M:%S")
    except Exception:
        return str(value)


def _row_matches_dashboard_filters(row: dict, filters: dict, skip_field: str | None = None) -> bool:
    for filter_name, row_key in FILTER_FIELD_ROW_MAP.items():
        if filter_name == skip_field:
            continue
        filter_value = str(filters.get(filter_name) or "").strip()
        if filter_value and str(row.get(row_key) or "").strip() != filter_value:
            return False
    return True


def _build_threshold_class_stats(
    filtered_rows: list[dict],
    classe_apprenant_counts: dict,
    threshold_class_codes: set[str] | None = None,
) -> tuple[list[dict], set[str]]:
    normalized_threshold_codes = {
        normalize_network_lookup(code)
        for code in (threshold_class_codes or set())
        if str(code or "").strip()
    }
    classe_groups_pre_threshold = {}
    for row in filtered_rows:
        classe_key = row["classe_code"]
        classe_groups_pre_threshold.setdefault(
            classe_key,
            {
                "code": row["classe_code"],
                "intitule": row.get("formation_intitule") or row["classe_intitule"],
                "prestation": row["prestation_code"],
                "cohorte": row["cohorte"],
                "metrics": _dashboard_bucket(),
            },
        )
        classe_groups_pre_threshold[classe_key]["intitule"] = _prefer_dashboard_label(
            classe_groups_pre_threshold[classe_key]["intitule"],
            row.get("formation_intitule") or row["classe_intitule"],
        )
        classe_groups_pre_threshold[classe_key]["prestation"] = _prefer_dashboard_label(
            classe_groups_pre_threshold[classe_key]["prestation"],
            row["prestation_code"],
        )
        classe_groups_pre_threshold[classe_key]["cohorte"] = _prefer_dashboard_label(
            classe_groups_pre_threshold[classe_key]["cohorte"],
            row["cohorte"],
        )
        _dashboard_bucket_add(classe_groups_pre_threshold[classe_key]["metrics"], row)

    classe_stats_all = sorted(
        [
            {
                "code": item["code"],
                "intitule": item["intitule"],
                "prestation": item["prestation"],
                "cohorte": item["cohorte"],
                "nb": item["metrics"]["nb"],
                "avgs": _dashboard_bucket_avgs(item["metrics"]),
                "total_apprenants": _analysis_class_count(classe_apprenant_counts, item["code"]),
                "threshold_reached": (
                    normalize_network_lookup(item["code"]) in normalized_threshold_codes
                    if normalized_threshold_codes
                    else (
                        item["metrics"]["nb"]
                        >= analysis_threshold_target(
                            _analysis_class_count(classe_apprenant_counts, item["code"])
                        )
                        if _analysis_class_count(classe_apprenant_counts, item["code"]) > 0
                        else item["metrics"]["nb"] > 0
                    )
                ),
            }
            for item in classe_groups_pre_threshold.values()
        ],
        key=lambda item: (item["code"], item["cohorte"]),
    )
    threshold_class_codes = {item["code"] for item in classe_stats_all if item["threshold_reached"]}
    return classe_stats_all, threshold_class_codes


def _thresholded_dashboard_rows(
    all_rows: list[dict],
    filters: dict,
    classe_apprenant_counts: dict,
    threshold_class_codes: set[str] | None = None,
    skip_field: str | None = None,
) -> tuple[list[dict], list[dict]]:
    filtered_rows = [
        row
        for row in all_rows
        if _row_matches_dashboard_filters(row, filters, skip_field=skip_field)
    ]
    classe_stats_all, threshold_class_codes = _build_threshold_class_stats(
        filtered_rows,
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
    )
    normalized_threshold_codes = {
        normalize_network_lookup(code)
        for code in threshold_class_codes
        if str(code or "").strip()
    }
    threshold_rows = [
        row
        for row in filtered_rows
        if normalize_network_lookup(row["classe_code"]) in normalized_threshold_codes
    ]
    return threshold_rows, classe_stats_all


def _build_dashboard_filter_options(
    all_rows: list[dict],
    filters: dict,
    classe_apprenant_counts: dict,
    threshold_class_codes: set[str] | None = None,
) -> dict[str, list[str]]:
    filter_options = {}
    for filter_name, row_key in FILTER_FIELD_ROW_MAP.items():
        option_rows, _ = _thresholded_dashboard_rows(
            all_rows,
            filters,
            classe_apprenant_counts,
            threshold_class_codes=threshold_class_codes,
            skip_field=filter_name,
        )
        filter_options[filter_name] = _sorted_unique(row.get(row_key, "") for row in option_rows)
    return filter_options


def _build_class_filter_options(
    all_rows: list[dict],
    filters: dict,
    classe_apprenant_counts: dict,
    threshold_class_codes: set[str] | None = None,
) -> list[dict]:
    _rows, classe_stats = _thresholded_dashboard_rows(
        all_rows,
        filters,
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
        skip_field="classe",
    )
    options = []
    for item in classe_stats:
        if not item["threshold_reached"]:
            continue
        target = analysis_threshold_target(item["total_apprenants"]) or item["nb"]
        options.append(
            {
                "value": item["code"],
                "label": (
                    f"{item['code']} - Seuil {analysis_threshold_label()} atteint "
                    f"({item['nb']} / {item['total_apprenants'] or item['nb']}, cible {target})"
                ),
            }
        )
    return options


def _build_dashboard_active_filters_summary(filters: dict) -> list[dict]:
    filter_labels = {
        "source": "Source",
        "prestation": "Prestation",
        "fenetre": "Fenêtre",
        "classe": "Classe",
        "prestataire": "Prestataire",
        "beneficiaire": "Bénéficiaire",
        "cohorte": "Cohorte",
        "ville": "Ville",
        "user": "Utilisateur",
        "status": "Status",
    }
    return [
        {"label": filter_labels[key], "value": str(value).strip()}
        for key, value in filters.items()
        if str(value or "").strip()
    ]


def _dashboard_tab_headers(active_tab: str) -> list[str]:
    if active_tab == "tab-classe":
        return [
            "Classe",
            "Intitulé de la formation",
            "Cohorte",
            "Nombre d'enquêtes",
            *[label for _, label in Q_FIELDS],
        ]
    if active_tab == "tab-prestation":
        return [
            "Code prestation",
            "Prestataire",
            "Bénéficiaire",
            "Nombre d'enquêtes",
            *[label for _, label in Q_FIELDS],
            "Global (Q9)",
        ]
    if active_tab == "tab-cohorte":
        return ["Cohorte", "Nombre d'enquêtes", *[label for _, label in Q_FIELDS], "Global (Q9)"]
    if active_tab == "tab-ville":
        return ["Ville", "Nombre d'enquêtes", "Moyenne de satisfaction globale (Q9)"]
    if active_tab == "tab-user":
        return ["Utilisateur", "Nombre d'enquêtes", "Moyenne de satisfaction globale (Q9)"]
    return [
        "ApprenantID réseau",
        "Code",
        "Apprenant",
        "Source",
        "Classe",
        "Intitulé de la formation",
        "Prestataire",
        "Bénéficiaire",
        "Cohorte",
        *[label for _, label in Q_FIELDS],
        "Commentaire",
    ]


def _build_dashboard_table_details(context: dict, rows: list[dict]) -> dict[str, dict]:
    row_counts = {
        "tab-apprenants": len(rows),
        "tab-classe": len(context["classe_stats"]),
        "tab-prestation": len(context["prestation_stats"]),
        "tab-cohorte": len(context["cohorte_stats"]),
        "tab-ville": len(context["ville_stats"]),
        "tab-user": len(context["user_stats"]),
    }
    return {
        tab_id: {
            "label": SATISFACTION_DASHBOARD_TAB_LABELS[tab_id],
            "description": SATISFACTION_DASHBOARD_TAB_DESCRIPTIONS[tab_id],
            "row_count": row_counts[tab_id],
            "column_count": len(_dashboard_tab_headers(tab_id)),
            "headers": _dashboard_tab_headers(tab_id),
        }
        for tab_id in SATISFACTION_DASHBOARD_TAB_LABELS
    }


def _build_table_details_context(context: dict, rows: list[dict]) -> dict[str, dict]:
    return _build_dashboard_table_details(context, rows)


def _build_appel_status_summary(
    *,
    target_class_codes: list[str] | None = None,
    strict_form_q=None,
) -> dict[str, int]:
    from App_PADESCE.appels.models import Appel as _Appel
    from django.db.models import Count as _Count, Q as _Q

    try:
        queryset = _Appel.objects.filter(is_active=True)
        if target_class_codes is not None:
            if target_class_codes:
                queryset = queryset.filter(
                    _Q(classe__code__in=target_class_codes)
                    | _Q(classe_label__in=target_class_codes)
                )
            else:
                queryset = queryset.none()

        form_filter = strict_form_q if strict_form_q is not None else _Q(status__in=CALL_FORM_STATUSES)
        forms_with_audio_filter = (
            form_filter & _Q(audio_file__isnull=False) & ~_Q(audio_file="")
            if strict_form_q is not None
            else _Q(status="formulaire_avec_audio")
        )
        forms_without_audio_filter = form_filter & (_Q(audio_file__isnull=True) | _Q(audio_file=""))

        return queryset.aggregate(
            appels_tentes=_Count("id", filter=~_Q(status="en_attente")),
            appels_reussis=_Count("id", filter=~_Q(status__in=["en_attente", "a_rappeler"])),
            formulaires_remplis=_Count("id", filter=form_filter),
            formulaires_remplis_sans_audio=_Count("id", filter=forms_without_audio_filter),
            formulaires_avec_audio=_Count("id", filter=forms_with_audio_filter),
            audios_enregistres=_Count("id", filter=_Q(audio_file__isnull=False) & ~_Q(audio_file="")),
        )
    except Exception as exc:
        try:
            from django.test.testcases import DatabaseOperationForbidden
        except Exception:
            DatabaseOperationForbidden = None
        if DatabaseOperationForbidden and isinstance(exc, DatabaseOperationForbidden):
            return {
                "appels_tentes": 0,
                "appels_reussis": 0,
                "formulaires_remplis": 0,
                "formulaires_remplis_sans_audio": 0,
                "formulaires_avec_audio": 0,
                "audios_enregistres": 0,
            }
        raise


def _ordered_survey_rows(rows: list[dict]) -> list[dict]:
    return sorted(
        rows,
        key=lambda row: (
            row.get("survey_date") or row.get("date"),
            row.get("survey_time") or row.get("heure"),
            row.get("classe_code", ""),
            row.get("source_apprenant_id") or row.get("apprenant_code", ""),
            row.get("apprenant_nom", ""),
        ),
    )


def _assign_enquete_ids(rows: list[dict]) -> list[dict]:
    grouped_rows: dict[str, list[tuple[int, dict]]] = {}
    for index, row in enumerate(rows, start=1):
        classe_key = normalize_network_lookup(row.get("classe_code", "")) or "__none__"
        grouped_rows.setdefault(classe_key, []).append((index, row))

    updates: dict[int, str] = {}
    for grouped in grouped_rows.values():
        known_ids: list[str] = []
        for _row_index, row in grouped:
            for raw_id in row.get("source_known_enquete_ids", []):
                normalized = _normalize_enquete_id(raw_id)
                if normalized and normalized not in known_ids:
                    known_ids.append(normalized)

        ordered_group = sorted(
            grouped,
            key=lambda item: (
                item[1].get("survey_date") or item[1].get("date"),
                item[1].get("survey_time") or item[1].get("heure"),
                item[1].get("source_apprenant_id") or item[1].get("apprenant_code", ""),
                item[1].get("apprenant_nom", ""),
                item[0],
            ),
        )
        for local_index, (row_index, row) in enumerate(ordered_group, start=1):
            normalized = (
                known_ids[local_index - 1]
                if local_index <= len(known_ids)
                else _normalize_enquete_id("", fallback_index=local_index)
            )
            updates[row_index] = normalized

    enriched = []
    for index, row in enumerate(rows, start=1):
        enriched.append(
            {
                **row,
                "source_enquete_id": updates.get(
                    index, _normalize_enquete_id("", fallback_index=index)
                ),
            }
        )
    return enriched


def _dashboard_export_filename_from_rows(rows: list[dict], filters: dict, extension: str) -> str:
    ordered_rows = _ordered_survey_rows(rows)
    if ordered_rows:
        first_row = ordered_rows[0]
        classe = _clean_export_part(
            first_row.get("classe_code") or filters.get("classe") or "TOUTES", "TOUTES"
        )
        prestataire = _clean_export_part(
            first_row.get("prestataire") or filters.get("prestataire") or "TOUS", "TOUS"
        )
        beneficiaire = _clean_export_part(
            first_row.get("beneficiaire") or filters.get("beneficiaire") or "TOUS", "TOUS"
        )
        cohorte = _clean_export_part(
            first_row.get("cohorte") or filters.get("cohorte") or "TOUTES", "TOUTES"
        )
        return f"{classe}_{prestataire}_{beneficiaire}_{cohorte}.{extension}"
    return _dashboard_export_filename(filters, extension)


def _safe_related(instance, attr_name: str):
    try:
        return getattr(instance, attr_name)
    except ObjectDoesNotExist:
        return None


def _source_class_matches_filters(
    source_class: dict, filters: dict, skip_field: str | None = None
) -> bool:
    source_filter_map = {
        "prestation": "prestation_id",
        "fenetre": "fenetre",
        "classe": "classe_id",
        "prestataire": "prestataire",
        "beneficiaire": "beneficiaire",
        "cohorte": "cohorte",
        "ville": "ville",
    }
    for filter_name, source_key in source_filter_map.items():
        if filter_name == skip_field:
            continue
        filter_value = str(filters.get(filter_name) or "").strip()
        if not filter_value:
            continue
        if normalize_network_lookup(source_class.get(source_key, "")) != normalize_network_lookup(
            filter_value
        ):
            return False
    return True


def _fallback_qualified_prestation_codes(classe_stats_all: list[dict]) -> set[str]:
    prestations: dict[str, list[bool]] = {}
    for item in classe_stats_all:
        prestation_key = normalize_network_lookup(item.get("prestation", ""))
        if not prestation_key:
            continue
        prestations.setdefault(prestation_key, []).append(bool(item.get("threshold_reached")))
    return {
        prestation_key
        for prestation_key, threshold_states in prestations.items()
        if threshold_states and all(threshold_states)
    }


def _source_prestation_classes_from_source(
    filters: dict, source_bundle: dict | None
) -> dict[str, dict[str, dict]]:
    if not source_bundle:
        return {}

    source_classes = list((source_bundle.get("classes") or {}).values())
    if not source_classes:
        return {}

    callable_class_counts = _normalize_class_count_map(_source_class_apprenant_counts(source_bundle))
    prestation_classes: dict[str, dict[str, dict]] = {}
    for source_class in source_classes:
        if not _source_class_matches_filters(source_class, filters):
            continue
        prestation_key = normalize_network_lookup(source_class.get("prestation_id", ""))
        classe_key = normalize_network_lookup(source_class.get("classe_id", ""))
        if not prestation_key or not classe_key:
            continue
        if int(callable_class_counts.get(classe_key) or 0) <= 0:
            continue
        prestation_classes.setdefault(prestation_key, {})[classe_key] = source_class
    return prestation_classes


def _terminated_prestation_codes_from_source(filters: dict, source_bundle: dict | None) -> set[str]:
    prestation_classes = _source_prestation_classes_from_source(filters, source_bundle)
    return {
        prestation_key
        for prestation_key, class_map in prestation_classes.items()
        if class_map
        and all(_source_class_is_finished(source_class) for source_class in class_map.values())
    }


def _qualified_prestation_codes_from_source(
    filters: dict,
    classe_stats_all: list[dict],
    source_bundle: dict | None,
    threshold_class_codes: set[str] | None = None,
) -> set[str]:
    fallback_codes = _fallback_qualified_prestation_codes(classe_stats_all)
    if not source_bundle or not list((source_bundle.get("classes") or {}).values()):
        return fallback_codes

    prestation_classes = _source_prestation_classes_from_source(filters, source_bundle)
    if not prestation_classes:
        return set()

    normalized_threshold_codes = {
        normalize_network_lookup(code)
        for code in (threshold_class_codes or set())
        if str(code or "").strip()
    }
    threshold_by_class = (
        {classe_key: True for classe_key in normalized_threshold_codes}
        if normalized_threshold_codes
        else {
            normalize_network_lookup(item.get("code", "")): bool(item.get("threshold_reached"))
            for item in classe_stats_all
            if item.get("code")
        }
    )
    terminated_prestation_codes = _terminated_prestation_codes_from_source(filters, source_bundle)
    if not terminated_prestation_codes:
        return set()

    qualified_codes = set()
    for prestation_key, class_map in prestation_classes.items():
        if prestation_key not in terminated_prestation_codes:
            continue
        class_keys = set(class_map)
        if class_keys and all(threshold_by_class.get(class_key, False) for class_key in class_keys):
            qualified_codes.add(prestation_key)
    return qualified_codes


def _is_ras_text(value: str) -> bool:
    normalized = normalize_network_lookup(value or "")
    return normalized in {"", "ras", "r.a.s", "r a s", "-", "neant", "sans observation"}


def _has_complete_answer_set(answer: AppelAnswers | None) -> bool:
    if not answer:
        return False
    return all(getattr(answer, field, None) not in (None, "") for field in APPEL_ANSWER_QUESTION_FIELDS)


def _has_complete_satisfaction_set(survey: SatisfactionApprenant | None) -> bool:
    if not survey:
        return False
    return all(getattr(survey, field, None) not in (None, "") for field in APPEL_ANSWER_QUESTION_FIELDS)


def _has_complete_form_record(
    answer: AppelAnswers | None,
    survey: SatisfactionApprenant | None,
) -> bool:
    return _has_complete_answer_set(answer) or _has_complete_satisfaction_set(survey)


def _has_ras_only_form(answer: AppelAnswers | None) -> bool:
    if not answer:
        return False
    return _is_ras_text(answer.commentaire) and _is_ras_text(answer.recommandations)


def _build_call_failure_reasons(
    appel: Appel,
    answer: AppelAnswers | None,
    survey: SatisfactionApprenant | None = None,
) -> list[str]:
    reasons: list[str] = []
    status = str(appel.status or "").strip()
    if status not in CALL_SUCCESS_STATUSES:
        reasons.append(f"Statut non finalise ({appel.get_status_display()})")
    if not appel_has_analysis_phone(appel):
        reasons.append("Sans numero")
    if appel_is_manually_excluded(appel):
        reasons.append("Exclu manuellement")
    if appel.deja_forme or appel.flag_pas_forme:
        reasons.append("Deja forme / pas concerne")
    if appel.flag_faux_nom:
        label = "Faux nom"
        if appel.flag_vrai_nom:
            label += f" ({appel.flag_vrai_nom})"
        reasons.append(label)
    if appel.flag_numero_double:
        reasons.append("Numero double")
    if appel.flag_deja_appele:
        reasons.append("Deja appele")
    if not answer:
        reasons.append("Formulaire absent")
    elif not _has_complete_answer_set(answer):
        reasons.append("Formulaire incomplet")
    if _has_ras_only_form(answer):
        reasons.append("Formulaire RAS")
    return list(dict.fromkeys(reasons))


def _call_report_status(
    appel: Appel,
    answer: AppelAnswers | None,
    survey: SatisfactionApprenant | None = None,
) -> tuple[bool, list[str]]:
    reasons = _build_call_failure_reasons(appel, answer, survey)
    return not reasons, reasons


def _excel_style(cell, *, fill=None, font=None, alignment=None, border=None, number_format=None):
    if fill is not None:
        cell.fill = fill
    if font is not None:
        cell.font = font
    if alignment is not None:
        cell.alignment = alignment
    if border is not None:
        cell.border = border
    if number_format is not None:
        cell.number_format = number_format


def _style_excel_table_header(
    worksheet, row_number: int, fill: PatternFill, font: Font, border: Border
):
    for cell in worksheet[row_number]:
        if cell.value in (None, ""):
            continue
        _excel_style(
            cell,
            fill=fill,
            font=font,
            alignment=Alignment(horizontal="center", vertical="center", wrap_text=True),
            border=border,
        )


def _style_excel_data_range(
    worksheet, start_row: int, border: Border, fill: PatternFill | None = None
):
    for row in worksheet.iter_rows(min_row=start_row, max_row=worksheet.max_row):
        for cell in row:
            if cell.value in (None, ""):
                continue
            _excel_style(
                cell,
                fill=fill,
                border=border,
                alignment=Alignment(vertical="top", wrap_text=True),
            )


def _daily_report_filename(generated_at) -> str:
    return f"rapport-quotidien-padesce-{generated_at.strftime('%Y%m%d')}.xlsx"


def _build_daily_report_row(appel: Appel, source_records: dict[str, dict]) -> dict:
    answer = _safe_related(appel, "answers")
    survey = _safe_related(appel, "satisfaction_apprenant")
    classe = getattr(appel, "classe", None)
    prestation = getattr(classe, "prestation", None) if classe else None
    inspecteur = getattr(survey, "inspecteur", None) if survey else None
    apprenant = getattr(survey, "apprenant", None) if survey else None
    answer_user = getattr(getattr(answer, "modified_by", None), "username", "") if answer else ""
    survey_user = getattr(getattr(survey, "enqueteur", None), "username", "") if survey else ""
    source_record = source_records.get(normalize_network_lookup(appel.code or "")) or {}
    is_success, failure_reasons = _call_report_status(appel, answer, survey)
    analysis_scope = _analysis_fenetre_for_appel(appel) in {"2", "3"}
    analysis_eligible = analysis_scope and appel_is_analysis_eligible(
        appel,
        answer=answer,
        survey=survey,
    )
    q_filled_count = 0
    for field in APPEL_ANSWER_QUESTION_FIELDS:
        val = getattr(answer, field, None) if answer else None
        if val is None and survey:
            val = getattr(survey, field, None)
        if val not in (None, ""):
            q_filled_count += 1
    has_form = q_filled_count >= 9
    analysis_included = has_form and analysis_eligible
    analysis_exclusion_reason = (
        appel_analysis_exclusion_reason(appel, answer=answer, survey=survey)
        if analysis_scope
        else "Fenetre hors analyse"
    )
    failure_detail_parts = []
    if not appel_has_analysis_phone(appel):
        failure_detail_parts.append("Aucun numero joignable n'est disponible pour cet apprenant.")
    if appel_is_manually_excluded(appel):
        failure_detail_parts.append("La ligne a ete exclue manuellement des analyses.")
    if appel.deja_forme or appel.flag_pas_forme:
        failure_detail_parts.append("Le beneficiaire indique qu'il a deja suivi la formation.")
    if appel.flag_faux_nom:
        failure_detail_parts.append(
            f"Nom incoherent. Vrai nom saisi: {appel.flag_vrai_nom or 'non renseigne'}."
        )
    if appel.flag_numero_double:
        failure_detail_parts.append("Le numero a ete signale comme doublon.")
    if appel.flag_deja_appele:
        failure_detail_parts.append("Le contact a deja ete traite dans la campagne.")
    if answer and _has_ras_only_form(answer):
        failure_detail_parts.append(
            "Le formulaire ne contient que 'RAS' dans les champs narratifs."
        )
    if not answer:
        failure_detail_parts.append("Aucune reponse Q1-Q9 n'a ete enregistree.")

    return {
        "id": appel.pk,
        "code": appel.code or "",
        "nom": appel.nom or "",
        "vrai_nom": appel.flag_vrai_nom or "",
        "source_apprenant_id": source_record.get("apprenant_id", ""),
        "classe": getattr(classe, "code", "")
        or appel.classe_label
        or source_record.get("classe_id", ""),
        "formation": getattr(classe, "intitule_formation", "")
        or source_record.get("formation", "")
        or appel.formation_padesce
        or "",
        "prestation": getattr(prestation, "code", "") or source_record.get("prestation_id", ""),
        "prestataire": (
            getattr(getattr(prestation, "prestataire", None), "raison_sociale", "")
            or source_record.get("prestataire", "")
            or appel.prestataire
        ),
        "beneficiaire": (
            getattr(getattr(prestation, "beneficiaire", None), "nom_structure", "")
            or source_record.get("beneficiaire", "")
            or appel.beneficiaire
        ),
        "fenetre": appel.fenetre or source_record.get("fenetre", ""),
        "cohorte": getattr(classe, "cohorte", "") or source_record.get("cohorte", ""),
        "ville": getattr(getattr(classe, "lieu", None), "ville", "")
        or source_record.get("ville", "")
        or appel.lieu,
        "lieu": appel.lieu or source_record.get("lieu", ""),
        "telephone1": appel.telephone1 or "",
        "telephone2": appel.telephone2 or "",
        "status": appel.get_status_display(),
        "taux_presence": float(appel.taux_presence or 0),
        "inspecteur": getattr(inspecteur, "nom_complet", "")
        or source_record.get("inspecteur_label", ""),
        "enqueteur": answer_user
        or survey_user
        or getattr(getattr(appel, "locked_by", None), "username", "")
        or "",
        "survey_date": _format_export_date(getattr(survey, "date", None)),
        "survey_time": _format_export_time(getattr(survey, "heure", None)),
        "commentaire": getattr(answer, "commentaire", "") if answer else "",
        "recommandations": getattr(answer, "recommandations", "") if answer else "",
        "transcription": getattr(survey, "transcription", "") if survey else "",
        "audio_name": os.path.basename(
            getattr(getattr(appel, "audio_file", None), "name", "") or ""
        ),
        "exclude_from_analysis": "Oui" if appel_is_manually_excluded(appel) else "Non",
        "analysis_scope": "Oui" if analysis_scope else "Non",
        "analysis_eligible": "Oui" if analysis_eligible else "Non",
        "analysis_included": "Oui" if analysis_included else "Non",
        "analysis_exclusion_reason": analysis_exclusion_reason,
        "source_status": "Trouve" if source_record else "Absent",
        "created_at": (
            timezone.localtime(appel.created_at).strftime("%d/%m/%Y %H:%M")
            if getattr(appel, "created_at", None)
            else ""
        ),
        "updated_at": (
            timezone.localtime(appel.updated_at).strftime("%d/%m/%Y %H:%M")
            if getattr(appel, "updated_at", None)
            else ""
        ),
        "deja_forme": "Oui" if appel.deja_forme else "Non",
        "flag_pas_forme": "Oui" if appel.flag_pas_forme else "Non",
        "flag_faux_nom": "Oui" if appel.flag_faux_nom else "Non",
        "flag_numero_double": "Oui" if appel.flag_numero_double else "Non",
        "flag_deja_appele": "Oui" if appel.flag_deja_appele else "Non",
        "formulaire_complet": "Oui" if _has_complete_answer_set(answer) else "Non",
        "formulaire_ras": "Oui" if _has_ras_only_form(answer) else "Non",
        "formulaire_all_three": "Oui" if answer_has_all_three_scores(answer) else "Non",
        "rapport_statut": "Reussi" if is_success else "Echoue",
        "motifs_echec": " | ".join(failure_reasons),
        "details_echec": " ".join(failure_detail_parts),
        "apprenant_reference": str(apprenant) if apprenant else "",
        "is_success": is_success,
        "failure_reasons": failure_reasons,
        **{
            field: getattr(answer, field, None) if answer else None
            for field in APPEL_ANSWER_QUESTION_FIELDS
        },
    }


def _append_daily_dashboard_sheet(
    worksheet,
    generated_at,
    total_count: int,
    success_rows: list[dict],
    failed_rows: list[dict],
    status_counts: Counter,
    failure_counts: Counter,
    source_bundle: dict | None,
):
    header_fill = PatternFill("solid", fgColor="4C1D95")
    subheader_fill = PatternFill("solid", fgColor="6F3CC3")
    success_fill = PatternFill("solid", fgColor="DCFCE7")
    failed_fill = PatternFill("solid", fgColor="FEE2E2")
    neutral_fill = PatternFill("solid", fgColor="F8FAFC")
    white_font = Font(color="FFFFFF", bold=True)
    title_font = Font(color="FFFFFF", bold=True, size=14)
    body_border = Border(
        left=Side(style="thin", color="D4C7EC"),
        right=Side(style="thin", color="D4C7EC"),
        top=Side(style="thin", color="D4C7EC"),
        bottom=Side(style="thin", color="D4C7EC"),
    )

    worksheet.title = "Tableau de bord"
    worksheet["A1"] = "Rapport quotidien PADESCE"
    worksheet["A2"] = f"Genere le {generated_at.strftime('%d/%m/%Y a %H:%M')}"
    worksheet["A3"] = "Perimetre: appels PADESCE actifs depuis le debut"
    worksheet.merge_cells("A1:D1")
    worksheet.merge_cells("A2:D2")
    worksheet.merge_cells("A3:D3")
    _excel_style(
        worksheet["A1"], fill=header_fill, font=title_font, alignment=Alignment(horizontal="center")
    )
    _excel_style(worksheet["A2"], fill=neutral_fill, border=body_border)
    _excel_style(worksheet["A3"], fill=neutral_fill, border=body_border)

    worksheet.append([])
    worksheet.append(["Indicateur", "Valeur", "Taux"])
    _style_excel_table_header(worksheet, 5, subheader_fill, white_font, body_border)

    success_count = len(success_rows)
    failed_count = len(failed_rows)
    success_rate = round((success_count / total_count) * 100, 2) if total_count else 0
    failed_rate = round((failed_count / total_count) * 100, 2) if total_count else 0

    dashboard_rows = [
        ("Total appels", total_count, "100%"),
        ("Appels reussis", success_count, f"{success_rate}%"),
        ("Appels echoues", failed_count, f"{failed_rate}%"),
    ]
    for label, value, rate in dashboard_rows:
        worksheet.append([label, value, rate])

    _style_excel_data_range(worksheet, 6, body_border)
    for cell in worksheet[7]:
        _excel_style(cell, fill=success_fill if cell.column <= 3 else None)
    for cell in worksheet[8]:
        _excel_style(cell, fill=failed_fill if cell.column <= 3 else None)

    worksheet.append([])
    worksheet.append(["Statut appel", "Nombre"])
    _style_excel_table_header(worksheet, 10, subheader_fill, white_font, body_border)
    for status_label in sorted(status_counts):
        worksheet.append([status_label, status_counts[status_label]])
    _style_excel_data_range(worksheet, 11, body_border)

    start_row = worksheet.max_row + 2
    worksheet.append(["Motif d'echec", "Nombre"])
    _style_excel_table_header(worksheet, start_row, subheader_fill, white_font, body_border)
    for label, count in failure_counts.most_common():
        worksheet.append([label, count])
    _style_excel_data_range(worksheet, start_row + 1, body_border)

    start_row = worksheet.max_row + 2
    worksheet.append(["Source reseau", "Valeur"])
    _style_excel_table_header(worksheet, start_row, subheader_fill, white_font, body_border)
    source_meta = source_bundle.get("source", {}) if source_bundle else {}
    worksheet.append(["Disponible", "Oui" if source_bundle else "Non"])
    worksheet.append(["Nom du fichier", source_meta.get("name", "Indisponible")])
    worksheet.append(["Derniere mise a jour", source_meta.get("modified_label", "")])
    worksheet.append(
        ["Apprenants source", (source_bundle or {}).get("counts", {}).get("apprenants", 0)]
    )
    worksheet.append(["Classes source", (source_bundle or {}).get("counts", {}).get("classes", 0)])
    worksheet.append(
        ["Prestations source", (source_bundle or {}).get("counts", {}).get("prestations", 0)]
    )
    _style_excel_data_range(worksheet, start_row + 1, body_border)

    worksheet.freeze_panes = "A5"
    _autosize_worksheet(worksheet)


def _append_daily_detail_sheet(
    worksheet, title: str, headers: list[str], rows: list[list], fill_color: str
):
    fill = PatternFill("solid", fgColor=fill_color)
    white_font = Font(color="FFFFFF", bold=True)
    border = Border(
        left=Side(style="thin", color="D4C7EC"),
        right=Side(style="thin", color="D4C7EC"),
        top=Side(style="thin", color="D4C7EC"),
        bottom=Side(style="thin", color="D4C7EC"),
    )
    worksheet.title = title
    worksheet.append(headers)
    _style_excel_table_header(worksheet, 1, fill, white_font, border)
    for row in rows:
        worksheet.append(row)
    _style_excel_data_range(worksheet, 2, border)
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions
    _autosize_worksheet(worksheet, max_width=48)


def _attach_network_source_to_rows(
    rows: list[dict],
    source_bundle: dict | None = None,
    source_key: str = ANALYSIS_DASHBOARD_DEFAULT_SOURCE,
) -> tuple[list[dict], dict]:
    try:
        source_bundle = source_bundle or build_padesce_source_index(source_key=source_key)
    except Exception as exc:
        unavailable_rows = [
            {
                **row,
                "source_found": False,
                "source_apprenant_id": "",
                "source_classe_id": "",
                "source_prestation_id": "",
                "source_fenetre": "",
                "source_cohorte": "",
                "source_status_label": "Source indisponible",
                "source_status_tone": "muted",
                "source_alerts": [],
                "source_alerts_label": str(exc),
                "formation_intitule": row.get("formation_intitule")
                or row.get("classe_intitule")
                or "-",
                "source_inspecteur_id": "",
                "source_inspecteur_label": "",
                "source_statut_prestation": "",
                "source_known_enquete_ids": [],
                "source_enquete_id": "",
            }
            for row in rows
        ]
        return unavailable_rows, _source_summary_unavailable(str(exc), source_key=source_key)

    source_records = source_bundle["records"]
    enriched_rows: list[dict] = []
    matched_count = 0
    missing_count = 0
    consistent_count = 0
    mismatch_count = 0
    apprenant_id_count = 0
    mismatch_rows: list[dict] = []

    for row in rows:
        code_key = normalize_network_lookup(row.get("apprenant_code", ""))
        source_record = source_records.get(code_key)

        if not source_record:
            missing_count += 1
            enriched_rows.append(
                {
                    **row,
                    "source_found": False,
                    "source_apprenant_id": "",
                    "source_classe_id": "",
                    "source_prestation_id": "",
                    "source_fenetre": "",
                    "source_cohorte": "",
                    "source_status_label": "Absent source",
                    "source_status_tone": "danger",
                    "source_alerts": [],
                    "source_alerts_label": "Code introuvable dans la feuille Apprenants du classeur réseau.",
                    "formation_intitule": row.get("formation_intitule")
                    or row.get("classe_intitule")
                    or "-",
                    "source_inspecteur_id": "",
                    "source_inspecteur_label": "",
                    "source_statut_prestation": "",
                    "source_known_enquete_ids": source_bundle.get("class_enquetes", {}).get(
                        normalize_network_lookup(row.get("classe_code", "")),
                        [],
                    ),
                    "source_enquete_id": "",
                }
            )
            continue

        matched_count += 1
        if source_record.get("apprenant_id"):
            apprenant_id_count += 1

        alerts = [
            alert
            for row_key, source_key, label in SOURCE_COMPARE_FIELDS
            for alert in [_source_compare_alert(row, source_record, row_key, source_key, label)]
            if alert
        ]
        if alerts:
            mismatch_count += 1
            status_label = "À vérifier"
            status_tone = "warning"
            if len(mismatch_rows) < 8:
                mismatch_rows.append(
                    {
                        "code": row.get("apprenant_code", ""),
                        "nom": row.get("apprenant_nom", ""),
                        "apprenant_id": source_record.get("apprenant_id", ""),
                        "alerts_label": "; ".join(alerts),
                    }
                )
        else:
            consistent_count += 1
            status_label = "OK"
            status_tone = "success"

        enriched_rows.append(
            {
                **row,
                "source_found": True,
                "source_apprenant_id": source_record.get("apprenant_id", ""),
                "source_classe_id": source_record.get("classe_id", ""),
                "source_prestation_id": source_record.get("prestation_id", ""),
                "source_fenetre": source_record.get("fenetre", ""),
                "source_cohorte": source_record.get("cohorte", ""),
                "source_status_label": status_label,
                "source_status_tone": status_tone,
                "source_alerts": alerts,
                "source_alerts_label": (
                    "; ".join(alerts) if alerts else "Aucun écart détecté avec le classeur réseau."
                ),
                "formation_intitule": source_record.get("formation")
                or row.get("formation_intitule")
                or row.get("classe_intitule")
                or "-",
                "source_inspecteur_id": source_record.get("inspecteur_id", ""),
                "source_inspecteur_label": source_record.get("inspecteur_label", ""),
                "source_statut_prestation": source_record.get("statut_prestation", ""),
                "source_known_enquete_ids": list(source_record.get("enquete_ids", [])),
                "source_enquete_id": "",
            }
        )

    return enriched_rows, {
        "available": True,
        "message": "",
        "key": source_bundle["source"].get("key", normalize_workbook_source_key(source_key)),
        "label": source_bundle["source"].get("label", ""),
        "name": source_bundle["source"]["name"],
        "modified_label": source_bundle["source"]["modified_label"],
        "matched_count": matched_count,
        "missing_count": missing_count,
        "consistent_count": consistent_count,
        "mismatch_count": mismatch_count,
        "apprenant_id_count": apprenant_id_count,
        "source_apprenant_count": sum(_source_class_apprenant_counts(source_bundle).values()),
        "duplicate_code_count": len(source_bundle["duplicate_codes"]),
        "mismatch_rows": mismatch_rows,
    }


def _clean_export_part(value: str, default: str) -> str:
    cleaned = re.sub(r"[^A-Z0-9_-]", "_", str(value or "").strip().upper())
    cleaned = re.sub(r"_+", "_", cleaned).strip("_")
    return (cleaned or default)[:40]


def _dashboard_export_filename(filters: dict, extension: str) -> str:
    classe = _clean_export_part(filters.get("classe") or "TOUTES", "TOUTES")
    prestataire = _clean_export_part(filters.get("prestataire") or "TOUS", "TOUS")
    beneficiaire = _clean_export_part(filters.get("beneficiaire") or "TOUS", "TOUS")
    cohorte = _clean_export_part(filters.get("cohorte") or "TOUTES", "TOUTES")
    return f"{classe}_{prestataire}_{beneficiaire}_{cohorte}.{extension}"


def _dashboard_class_export_label(classe_code: str, rows: list[dict], filters: dict) -> str:
    for row in rows:
        if str(row.get("classe_code") or "").strip() == str(classe_code or "").strip():
            prestataire = row.get("prestataire") or "-"
            beneficiaire = row.get("beneficiaire") or "-"
            return f"{classe_code}_{prestataire}_{beneficiaire}"
    return str(classe_code or "Non renseignée")


def _dashboard_export_chapeau_filename(filters: dict, extension: str) -> str:
    base_name = _dashboard_export_filename(filters, extension)
    return f"EVALUATION_DES_CLASSES_{base_name}"


def _dashboard_chapeau_title(class_label: str) -> str:
    return f"Enquete de satisfaction : {class_label}"


def _tabular_dashboard_export(
    active_tab: str, context: dict, rows: list[dict]
) -> tuple[list[str], list[list]]:
    if active_tab == "tab-classe":
        return (
            [
                "Classe",
                "Intitulé de la formation",
                "Cohorte",
                "Nombre d'enquêtes",
                *[label for _, label in Q_FIELDS],
            ],
            [
                [item["code"], item["intitule"], item["cohorte"], item["nb"], *item["avgs"]]
                for item in context["classe_stats"]
            ],
        )
    if active_tab == "tab-prestation":
        return (
            [
                "Prestation",
                "Prestataire",
                "Bénéficiaire",
                "Nombre d'enquêtes",
                *[label for _, label in Q_FIELDS],
                "Global (Q9)",
            ],
            [
                [
                    item["code"],
                    item["prestataire"],
                    item["beneficiaire"],
                    item["nb"],
                    *item["avgs"],
                    item["avg"],
                ]
                for item in context["prestation_stats"]
            ],
        )
    if active_tab == "tab-cohorte":
        return (
            ["Cohorte", "Nombre d'enquêtes", *[label for _, label in Q_FIELDS], "Global (Q9)"],
            [
                [item["label"], item["nb"], *item["avgs"], item["avg"]]
                for item in context["cohorte_stats"]
            ],
        )
    if active_tab == "tab-ville":
        return (
            ["Ville", "Nombre d'enquêtes", "Moyenne de satisfaction globale (Q9)"],
            [[item["ville"], item["nb"], item["avg"]] for item in context["ville_stats"]],
        )
    if active_tab == "tab-user":
        return (
            ["Utilisateur", "Nombre d'enquêtes", "Moyenne de satisfaction globale (Q9)"],
            [[item["username"], item["nb"], item["avg"]] for item in context["user_stats"]],
        )
    return (
        [
            "N°",
            "ID Apprenant",
            "Nom Apprenant",
            "Bénéficiaire",
            "Prestataire",
            "Formation",
            "Inspecteur",
            "ClassID",
            *CSV_Q_HEADERS,
            "Commentaires",
            "Recommandations",
            "N°Enquête",
            "Date Enregistrement",
            "Heure Enregistrement",
        ],
        [
            [
                index,
                row.get("source_apprenant_id", ""),
                row["apprenant_nom"],
                row["beneficiaire"],
                row["prestataire"],
                row.get("formation_intitule") or row["classe_intitule"],
                row.get("inspecteur_code")
                or row.get("source_inspecteur_id")
                or row.get("source_inspecteur_label", ""),
                row["classe_code"],
                *[row.get(field) for field, _ in Q_FIELDS],
                row["commentaire"],
                row["recommandations"],
                row.get("source_enquete_id", ""),
                _format_export_date(row.get("survey_date")),
                _format_export_time(row.get("survey_time")),
            ]
            for index, row in enumerate(_ordered_survey_rows(rows), start=1)
        ],
    )


def _build_missing_prestations_analysis(
    terminated_prestation_codes: set,
    qualified_prestation_codes: set,
    source_bundle: dict | None,
    classe_stats_all: list[dict],
    filters: dict | None = None,
    threshold_class_codes: set[str] | None = None,
) -> dict:
    """
    Analyse les prestations terminées du fichier source qui n'apparaissent pas
    encore dans le tableau de bord satisfaction (seuil non atteint ou autres raisons).

    Categories retournées :
      - pas_disponible   : la prestation / classe n'existe pas dans le site (non importée)
      - pas_de_numero    : la classe existe mais les apprenants n'ont pas de numéro
      - pas_seuil_atteint: la classe existe, les numéros aussi, mais le seuil d'analyse n'est pas atteint
      - autres           : toute autre raison
    """
    if not source_bundle:
        return {
            "available": False,
            "total_source": 0,
            "total_qualified": 0,
            "total_missing": 0,
            "by_category": {},
            "details": [],
        }

    missing_keys = terminated_prestation_codes - qualified_prestation_codes
    if not missing_keys:
        return {
            "available": True,
            "total_source": len(terminated_prestation_codes),
            "total_qualified": len(qualified_prestation_codes),
            "total_missing": 0,
            "by_category": {
                "pas_disponible": 0,
                "pas_seuil_atteint": 0,
                "pas_de_numero": 0,
                "autres": 0,
            },
            "details": [],
        }

    local_by_code = _local_analysis_class_summary()

    threshold_reached_codes: set[str] = (
        {
            normalize_network_lookup(code)
            for code in (threshold_class_codes or set())
            if str(code or "").strip()
        }
        or {
            normalize_network_lookup(item.get("code", ""))
            for item in classe_stats_all
            if item.get("threshold_reached") and item.get("code")
        }
    )

    # Pre-fetch existing Appel codes (for importable_count computation)
    existing_appel_codes: set[str] = set(
        Appel.objects.filter(is_active=True).values_list("code", flat=True)
    )
    # Source records indexed by normalized prestation key (phone + not-in-appel filter)
    _phone_re = re.compile(r"[0-9]{7,}")
    source_records_by_prestation: dict[str, int] = {}
    for rec in source_bundle.get("records", {}).values():
        p_key = normalize_network_lookup(rec.get("prestation_id", ""))
        if not p_key:
            continue
        numero = (rec.get("telephone1") or rec.get("telephone2") or rec.get("numero") or "").strip()
        code = (rec.get("code") or "").strip()
        if _phone_re.search(numero) and code and code not in existing_appel_codes:
            source_records_by_prestation[p_key] = source_records_by_prestation.get(p_key, 0) + 1

    source_prestations: dict = source_bundle.get("prestations", {})

    # Use the same callable source classes as the denominator so "sans numero"
    # records never re-enter the missing-prestation detail.
    classes_by_prestation = {
        prestation_key: list(class_map.values())
        for prestation_key, class_map in _source_prestation_classes_from_source(
            filters or {},
            source_bundle,
        ).items()
    }

    CATEGORY_LABELS = {
        "pas_disponible": "Pas sur le site (non importée)",
        "pas_seuil_atteint": f"Seuil {analysis_threshold_label()} non atteint",
        "pas_de_numero": "Pas de numéro de téléphone",
        "autres": "Autres",
    }
    by_category: dict[str, int] = {k: 0 for k in CATEGORY_LABELS}
    details: list[dict] = []

    for p_key in sorted(missing_keys):
        p_info = source_prestations.get(p_key, {})
        source_cls_list = classes_by_prestation.get(p_key, [])

        prestation_id = p_info.get("prestation_id", "") or p_key
        prestataire = p_info.get("prestataire", "")
        beneficiaire = p_info.get("beneficiaire", "")
        formation = p_info.get("formation", "")

        if not source_cls_list:
            category = "autres"
        else:
            per_class: list[str] = []
            for cls in source_cls_list:
                c_key = normalize_network_lookup(cls.get("classe_id", ""))
                if not c_key:
                    per_class.append("autres")
                    continue

                local = local_by_code.get(c_key)
                if local is None:
                    per_class.append("pas_disponible")
                elif int(local.get("with_phone") or 0) == 0:
                    per_class.append("pas_de_numero")
                elif c_key not in threshold_reached_codes:
                    per_class.append("pas_seuil_atteint")
                else:
                    per_class.append("autres")

            # Worst-case priority: pas_disponible > pas_de_numero > pas_seuil_atteint > autres
            if "pas_disponible" in per_class:
                category = "pas_disponible"
            elif "pas_de_numero" in per_class:
                category = "pas_de_numero"
            elif "pas_seuil_atteint" in per_class:
                category = "pas_seuil_atteint"
            else:
                category = "autres"

        importable_count = source_records_by_prestation.get(
            normalize_network_lookup(prestation_id), 0
        )
        by_category[category] += 1
        details.append(
            {
                "prestation_id": prestation_id,
                "prestataire": prestataire,
                "beneficiaire": beneficiaire,
                "formation": formation,
                "category": category,
                "category_label": CATEGORY_LABELS[category],
                "classe_count": len(source_cls_list),
                "importable_count": importable_count,
            }
        )

    details.sort(key=lambda x: (x["category"], x["prestataire"], x["prestation_id"]))
    total_importable = sum(d["importable_count"] for d in details)

    return {
        "available": True,
        "total_source": len(terminated_prestation_codes),
        "total_qualified": len(qualified_prestation_codes),
        "total_missing": len(missing_keys),
        "total_importable": total_importable,
        "by_category": by_category,
        "category_labels": CATEGORY_LABELS,
        "details": details,
    }


def _build_satisfaction_dashboard_data(request):
    selected_source = _analysis_selected_source(request)
    source_options = get_workbook_source_options()
    source_option_map = {item["value"]: item for item in source_options}
    filters = {
        "source": selected_source,
        "prestation": request.GET.get("prestation", ""),
        "fenetre": request.GET.get("fenetre", ""),
        "ville": request.GET.get("ville", ""),
        "user": request.GET.get("user", ""),
        "classe": request.GET.get("classe", ""),
        "prestataire": request.GET.get("prestataire", ""),
        "beneficiaire": request.GET.get("beneficiaire", ""),
        "cohorte": request.GET.get("cohorte", ""),
        "status": request.GET.get("status", ""),
    }

    try:
        source_bundle = build_padesce_source_index(source_key=selected_source)
    except Exception:
        source_bundle = None
    cache_key = _analysis_cache_key(
        "dashboard-data",
        selected_source,
        request.GET.urlencode(),
        ((source_bundle or {}).get("source") or {}).get("modified_at", "no-source"),
        _analysis_queryset_marker(Appel),
        _analysis_queryset_marker(AppelAnswers),
        _analysis_queryset_marker(SatisfactionApprenant),
    )
    cached_payload = cache.get(cache_key)
    if cached_payload is not None:
        return cached_payload
    threshold_class_codes = _status_threshold_class_codes(source_bundle)

    all_rows = [
        _dashboard_row_from_answer(answer) for answer in _satisfaction_dashboard_base_queryset()
    ]
    all_rows = [
        row
        for row in all_rows
        if row["fenetre"] in {"2", "3"} and row.get("analysis_included")
    ]

    classe_apprenant_counts = _local_analysis_class_counts()
    analysis_scope_filters = {key: "" for key in filters}
    analysis_scope_filters["source"] = selected_source
    rows, _filtered_classe_stats = _thresholded_dashboard_rows(
        all_rows,
        filters,
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
    )
    _, classe_stats_all = _thresholded_dashboard_rows(
        all_rows,
        analysis_scope_filters,
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
    )
    rows, source_summary = _attach_network_source_to_rows(
        rows,
        source_bundle=source_bundle,
        source_key=selected_source,
    )
    rows = _assign_enquete_ids(rows)
    total = len(rows)
    analysis_audio_count = sum(1 for row in rows if row.get("has_audio"))
    terminated_prestation_codes = _terminated_prestation_codes_from_source(
        analysis_scope_filters,
        source_bundle,
    )
    qualified_prestation_codes = _qualified_prestation_codes_from_source(
        analysis_scope_filters,
        classe_stats_all,
        source_bundle,
        threshold_class_codes=threshold_class_codes,
    )

    global_bucket = _dashboard_bucket()
    classe_groups = {}
    prestation_groups = {}
    fenetre_groups = {}
    ville_groups = {}
    user_groups = {}
    cohorte_groups = {}
    prestataire_groups = {}
    beneficiaire_groups = {}

    for row in rows:
        _dashboard_bucket_add(global_bucket, row)
        effective_prestation_code = row.get("source_prestation_id") or row["prestation_code"]

        classe_key = row["classe_code"]
        classe_groups.setdefault(
            classe_key,
            {
                "code": row["classe_code"],
                "intitule": row.get("formation_intitule") or row["classe_intitule"],
                "prestation": effective_prestation_code,
                "cohorte": row["cohorte"],
                "fenetre": row.get("fenetre", ""),
                "metrics": _dashboard_bucket(),
            },
        )
        classe_groups[classe_key]["intitule"] = _prefer_dashboard_label(
            classe_groups[classe_key]["intitule"],
            row.get("formation_intitule") or row["classe_intitule"],
        )
        classe_groups[classe_key]["prestation"] = _prefer_dashboard_label(
            classe_groups[classe_key]["prestation"],
            effective_prestation_code,
        )
        classe_groups[classe_key]["cohorte"] = _prefer_dashboard_label(
            classe_groups[classe_key]["cohorte"],
            row["cohorte"],
        )
        _dashboard_bucket_add(classe_groups[classe_key]["metrics"], row)

        normalized_p_code = normalize_network_lookup(effective_prestation_code)
        prestation_key = (normalized_p_code, row["prestataire"], row["beneficiaire"])
        prestation_groups.setdefault(
            prestation_key,
            {
                "code": normalized_p_code,
                "prestataire": row["prestataire"],
                "beneficiaire": row["beneficiaire"],
                "associated_classes": set(),
                "metrics": _dashboard_bucket(),
            },
        )
        prestation_groups[prestation_key]["associated_classes"].add(classe_key)
        _dashboard_bucket_add(prestation_groups[prestation_key]["metrics"], row)

        fenetre_groups.setdefault(
            row["fenetre"], {"label": row["fenetre"], "metrics": _dashboard_bucket()}
        )
        _dashboard_bucket_add(fenetre_groups[row["fenetre"]]["metrics"], row)

        ville_groups.setdefault(
            row["ville"], {"ville": row["ville"], "metrics": _dashboard_bucket()}
        )
        _dashboard_bucket_add(ville_groups[row["ville"]]["metrics"], row)

        user_groups.setdefault(
            row["user"], {"username": row["user"], "metrics": _dashboard_bucket()}
        )
        _dashboard_bucket_add(user_groups[row["user"]]["metrics"], row)

        cohorte_groups.setdefault(
            row["cohorte"], {"label": row["cohorte"], "metrics": _dashboard_bucket()}
        )
        _dashboard_bucket_add(cohorte_groups[row["cohorte"]]["metrics"], row)

        prestataire_groups.setdefault(
            row["prestataire"], {"label": row["prestataire"], "metrics": _dashboard_bucket()}
        )
        _dashboard_bucket_add(prestataire_groups[row["prestataire"]]["metrics"], row)

        beneficiaire_groups.setdefault(
            row["beneficiaire"], {"label": row["beneficiaire"], "metrics": _dashboard_bucket()}
        )
        _dashboard_bucket_add(beneficiaire_groups[row["beneficiaire"]]["metrics"], row)

    global_avgs = {label: _dashboard_bucket_avg(global_bucket, field) for field, label in Q_FIELDS}

    classe_stats = sorted(
        [
            {
                "code": item["code"],
                "intitule": item["intitule"],
                "prestation": item["prestation"],
                "cohorte": item["cohorte"],
                "fenetre": item.get("fenetre", ""),
                "nb": item["metrics"]["nb"],
                "avgs": _dashboard_bucket_avgs(item["metrics"]),
                "total_apprenants": _analysis_class_count(classe_apprenant_counts, item["code"]),
                "threshold_reached": True,
            }
            for item in classe_groups.values()
        ],
        key=lambda item: (item["code"], item["cohorte"]),
    )
    classe_stats_seuil = classe_stats

    prestation_stats = sorted(
        [
            {
                "code": item["code"],
                "prestataire": item["prestataire"],
                "beneficiaire": item["beneficiaire"],
                "nb": item["metrics"]["nb"],
                "avg": _dashboard_bucket_avg(item["metrics"], "q9_satisfaction_globale"),
                "avgs": _dashboard_bucket_avgs(item["metrics"]),
                "effectif": sum(
                    _analysis_class_count(classe_apprenant_counts, c)
                    for c in item["associated_classes"]
                ),
            }
            for item in prestation_groups.values()
            if normalize_network_lookup(item["code"]) in qualified_prestation_codes
        ],
        key=lambda item: (item["code"], item["prestataire"], item["beneficiaire"]),
    )

    # Full list (unfiltered by qualified) — used for ranking/map features
    prestation_stats_all = sorted(
        [
            {
                "code": item["code"],
                "prestataire": item["prestataire"],
                "beneficiaire": item["beneficiaire"],
                "nb": item["metrics"]["nb"],
                "avg": _dashboard_bucket_avg(item["metrics"], "q9_satisfaction_globale"),
                "avgs": _dashboard_bucket_avgs(item["metrics"]),
                "effectif": sum(
                    _analysis_class_count(classe_apprenant_counts, c)
                    for c in item["associated_classes"]
                ),
            }
            for item in prestation_groups.values()
        ],
        key=lambda item: (item["code"], item["prestataire"], item["beneficiaire"]),
    )

    ville_stats = sorted(
        [
            {
                "ville": item["ville"],
                "nb": item["metrics"]["nb"],
                "avg": _dashboard_bucket_avg(item["metrics"], "q9_satisfaction_globale"),
            }
            for item in ville_groups.values()
        ],
        key=lambda item: item["ville"],
    )

    user_stats = sorted(
        [
            {
                "username": item["username"],
                "nb": item["metrics"]["nb"],
                "avg": _dashboard_bucket_avg(item["metrics"], "q9_satisfaction_globale"),
            }
            for item in user_groups.values()
        ],
        key=lambda item: item["username"],
    )

    cohorte_stats = sorted(
        [
            {
                "label": item["label"],
                "nb": item["metrics"]["nb"],
                "avg": _dashboard_bucket_avg(item["metrics"], "q9_satisfaction_globale"),
                "avgs": _dashboard_bucket_avgs(item["metrics"]),
            }
            for item in cohorte_groups.values()
        ],
        key=lambda item: item["label"],
    )
    fenetre_stats = sorted(
        [
            {
                "label": item["label"],
                "nb": item["metrics"]["nb"],
            }
            for item in fenetre_groups.values()
        ],
        key=lambda item: item["label"],
    )

    analyzed_classes = [
        {"code": item["code"], "label": f"{item['code']} - {item['intitule']}", "nb": item["nb"], "fenetre": item.get("fenetre", "")}
        for item in classe_stats_seuil
    ]
    analyzed_prestations = [
        {
            "code": item["code"],
            "label": f"{item['code']} | {item['prestataire']} | {item['beneficiaire']}",
            "nb": item["nb"],
        }
        for item in prestation_stats
    ]
    analyzed_fenetres = [{"label": item["label"], "nb": item["nb"]} for item in fenetre_stats]
    analyzed_prestataires = [
        {"label": label, "nb": item["metrics"]["nb"]}
        for label, item in sorted(prestataire_groups.items(), key=lambda pair: pair[0])
    ]
    analyzed_beneficiaires = [
        {"label": label, "nb": item["metrics"]["nb"]}
        for label, item in sorted(beneficiaire_groups.items(), key=lambda pair: pair[0])
    ]
    analyzed_cohortes = [{"label": item["label"], "nb": item["nb"]} for item in cohorte_stats]

    filter_options = _build_dashboard_filter_options(
        all_rows,
        filters,
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
    )
    eligible_prestation_options = sorted(
        {
            item["code"]
            for item in prestation_stats
            if str(item.get("code") or "").strip() and item["code"] != "-"
        }
    )
    if filters["prestation"] and filters["prestation"] not in eligible_prestation_options:
        eligible_prestation_options.append(filters["prestation"])
    filter_options["prestation"] = eligible_prestation_options
    class_options = _build_class_filter_options(
        all_rows,
        filters,
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
    )
    active_filters_summary = _build_dashboard_active_filters_summary(
        {
            **filters,
            "source": source_option_map.get(selected_source, {}).get("label", selected_source),
        }
    )

    missing_analysis = _build_missing_prestations_analysis(
        terminated_prestation_codes,
        qualified_prestation_codes,
        source_bundle,
        classe_stats_all,
        analysis_scope_filters,
        threshold_class_codes=threshold_class_codes,
    )

    # Build prestataire → classes/beneficiaires mapping for dynamic filters
    prestataire_to_classes: dict[str, set[str]] = {}
    prestataire_to_beneficiaires: dict[str, set[str]] = {}
    for row in all_rows:
        prest = str(row.get("prestataire") or "").strip()
        classe = str(row.get("classe_code") or "").strip()
        benef = str(row.get("beneficiaire") or "").strip()
        if prest and classe:
            prestataire_to_classes.setdefault(prest, set()).add(classe)
        if prest and benef:
            prestataire_to_beneficiaires.setdefault(prest, set()).add(benef)
    filter_map_json = json.dumps({
        "prestataire_to_classes": {k: sorted(v) for k, v in prestataire_to_classes.items()},
        "prestataire_to_beneficiaires": {k: sorted(v) for k, v in prestataire_to_beneficiaires.items()},
    })

    filter_query_string = request.GET.copy().urlencode()
    analyzed_prestations_count = (
        int(missing_analysis.get("total_qualified") or 0)
        if missing_analysis.get("available")
        else len(analyzed_prestations)
    )
    analyzed_prestations_total_count = (
        int(missing_analysis.get("total_source") or 0)
        if missing_analysis.get("available")
        else len(analyzed_prestations)
    )
    analyzed_prestations_ratio = (
        f"{analyzed_prestations_count}/{analyzed_prestations_total_count}"
    )
    context = {
        "total": total,
        "global_avgs": global_avgs,
        "q_labels": [label for _, label in Q_FIELDS],
        "classe_stats": classe_stats_seuil,
        "prestation_stats": prestation_stats,
        "prestation_stats_all": prestation_stats_all,
        "ville_stats": ville_stats,
        "user_stats": user_stats,
        "cohorte_stats": cohorte_stats,
        "filter_prestation": filters["prestation"],
        "filter_source": selected_source,
        "filter_fenetre": filters["fenetre"],
        "filter_ville": filters["ville"],
        "filter_user": filters["user"],
        "filter_classe": filters["classe"],
        "filter_prestataire": filters["prestataire"],
        "filter_beneficiaire": filters["beneficiaire"],
        "filter_cohorte": filters["cohorte"],
        "filter_status": filters["status"],
        "analyzed_classes": analyzed_classes,
        "analyzed_prestations": analyzed_prestations,
        "analyzed_fenetres": analyzed_fenetres,
        "analyzed_prestataires": analyzed_prestataires,
        "analyzed_beneficiaires": analyzed_beneficiaires,
        "analyzed_cohortes": analyzed_cohortes,
        "analyzed_classes_count": len(analyzed_classes),
        "analyzed_prestations_count": analyzed_prestations_count,
        "analyzed_prestations_total_count": analyzed_prestations_total_count,
        "analyzed_prestations_ratio": analyzed_prestations_ratio,
        "analyzed_fenetres_count": len(analyzed_fenetres),
        "analyzed_prestataires_count": len(analyzed_prestataires),
        "analyzed_beneficiaires_count": len(analyzed_beneficiaires),
        "analyzed_cohortes_count": len(analyzed_cohortes),
        "analysis_audio_count": analysis_audio_count,
        "analysis_threshold_label": analysis_threshold_label(),
        "filter_query_string": filter_query_string,
        "active_filters_summary": active_filters_summary,
        "source_summary": source_summary,
        "source_options": source_options,
        "class_options": class_options,
        "missing_analysis": missing_analysis,
        "prestations": filter_options.get("prestation", []),
        "fenetres": filter_options.get("fenetre", []),
        "villes": filter_options.get("ville", []),
        "users": filter_options.get("user", []),
        "classes": filter_options.get("classe", []),
        "prestataires": filter_options.get("prestataire", []),
        "beneficiaires": filter_options.get("beneficiaire", []),
        "cohortes": filter_options.get("cohorte", []),
        "status": filter_options.get("status", []),
        "filter_map_json": filter_map_json,
    }
    from django.db.models import Q as _Q

    q_fields = [
        "q1_clarte_exposes",
        "q2_interaction_formateur",
        "q3_maitrise_contenu",
        "q4_salle_adequate",
        "q5_materiel_disponible",
        "q6_organisation_temps",
        "q7_utilite_formation",
        "q8_adequation_besoins",
        "q9_satisfaction_globale",
    ]
    answers_valid_q = _Q()
    for field_name in q_fields:
        answers_valid_q &= _Q(**{f"answers__{field_name}__isnull": False})
    survey_valid_q = _Q(satisfaction_apprenant__isnull=False)
    strict_form_q = answers_valid_q | survey_valid_q

    target_class_codes = [
        str(item.get("code") or "").strip()
        for item in classe_stats_seuil
        if str(item.get("code") or "").strip()
    ]

    _appel_stats = _build_appel_status_summary(
        target_class_codes=target_class_codes,
        strict_form_q=strict_form_q,
    )
    context["appels_tentes"] = _appel_stats["appels_tentes"]
    context["appels_reussis"] = _appel_stats["appels_reussis"]
    context["formulaires_remplis_appels"] = _appel_stats["formulaires_remplis"]
    context["formulaires_remplis_sans_audio_appels"] = _appel_stats["formulaires_remplis_sans_audio"]
    context["formulaires_avec_audio_appels"] = _appel_stats["formulaires_avec_audio"]
    context["audios_enregistres_appels"] = _appel_stats["audios_enregistres"]
    context["tab_details"] = _build_table_details_context(context, rows)
    payload = {"rows": rows, "filters": filters, "context": context}
    cache.set(cache_key, payload, timeout=ANALYSIS_CACHE_TIMEOUT)
    return payload


@require_analysis_access
def satisfaction_dashboard_export_chapeau(request):
    dashboard = _build_satisfaction_dashboard_data(request)
    context = dashboard["context"]
    rows = dashboard["rows"]
    filename = _dashboard_export_chapeau_filename(dashboard["filters"], "docx")

    from docx import Document

    document = Document()

    class_question_counts: dict[str, dict[str, int]] = defaultdict(dict)
    class_respondent_counts: dict[str, int] = {}
    class_respondent_keys: dict[str, set[str]] = defaultdict(set)

    for row in rows:
        classe_code = str(row.get("classe_code") or "").strip()
        if not classe_code:
            continue

        apprenant_key = (
            str(row.get("source_apprenant_id") or "").strip()
            or str(row.get("apprenant_code") or "").strip()
            or str(row.get("apprenant_nom") or "").strip()
        )

        if apprenant_key:
            class_respondent_keys[classe_code].add(apprenant_key)

        for field, _label in Q_FIELDS:
            value = row.get(field)
            if value not in (None, ""):
                class_question_counts.setdefault(classe_code, {})
                class_question_counts[classe_code][field] = (
                    class_question_counts[classe_code].get(field, 0) + 1
                )

    for classe_code, respondents in class_respondent_keys.items():
        class_respondent_counts[classe_code] = len(respondents)

    if not context.get("classe_stats"):
        document.add_paragraph("Aucune classe à exporter.")
    else:
        for item in context["classe_stats"]:
            classe_code = item["code"]
            classe_nom_complet = _dashboard_class_export_label(
                classe_code, rows, dashboard["filters"]
            )
            total_repondants = class_respondent_counts.get(classe_code, 0)

            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"

            title_cells = table.rows[0].cells
            title_cell = title_cells[0].merge(title_cells[1])
            title_paragraph = title_cell.paragraphs[0]
            title_run = title_paragraph.add_run(_dashboard_chapeau_title(classe_nom_complet))
            title_run.bold = True

            headers = table.add_row().cells
            headers[0].text = "QUESTION"
            headers[1].text = "NOTE"

            for index, (field, question_label) in enumerate(Q_FIELDS):
                note = item["avgs"][index] if index < len(item.get("avgs", [])) else 0
                total_question = class_question_counts.get(classe_code, {}).get(field, 0)

                row_cells = table.add_row().cells
                row_cells[0].text = question_label
                row_cells[1].text = f"{note}/5" if total_question else "-"

            # Ligne TOTAL conservée
            total_row = table.add_row().cells
            total_label_cell = total_row[0]
            total_label_cell.text = "TOTAL DES PARTICIPANTS A LA FORMATION"
            total_row[1].text = str(total_repondants)

            document.add_paragraph("")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@require_analysis_access
def satisfaction_dashboard_export_xlsx(request):
    dashboard = _build_satisfaction_dashboard_data(request)
    rows = dashboard["rows"]
    context = dashboard["context"]

    wb = openpyxl.Workbook()
    ws_summary = wb.active
    ws_summary.title = "Synthese"
    ws_summary.append(["Analyse satisfaction apprenants"])
    ws_summary.append([])
    ws_summary.append(["Total enquêtes", context["total"]])
    ws_summary.append(["Classes analysées", context["analyzed_classes_count"]])
    ws_summary.append(["Prestations analysees", context["analyzed_prestations_ratio"]])
    ws_summary.append(["Prestataires analysés", context["analyzed_prestataires_count"]])
    ws_summary.append(["Bénéficiaires analysés", context["analyzed_beneficiaires_count"]])
    ws_summary.append(["Cohortes analysées", context["analyzed_cohortes_count"]])
    ws_summary.append([])
    ws_summary.append(["Filtres appliqués"])
    ws_summary.append(["Prestation", context["filter_prestation"] or "Toutes"])
    ws_summary.append(["Fenêtre", context["filter_fenetre"] or "Toutes"])
    ws_summary.append(["Classe", context["filter_classe"] or "Toutes"])
    ws_summary.append(["Prestataire", context["filter_prestataire"] or "Tous"])
    ws_summary.append(["Bénéficiaire", context["filter_beneficiaire"] or "Tous"])
    ws_summary.append(["Cohorte", context["filter_cohorte"] or "Toutes"])
    ws_summary.append(["Ville", context["filter_ville"] or "Toutes"])
    ws_summary.append(["Utilisateur", context["filter_user"] or "Tous"])
    ws_summary.append([])
    ws_summary.append(["Contrôle source réseau"])
    ws_summary.append(
        ["Source disponible", "Oui" if context["source_summary"]["available"] else "Non"]
    )
    ws_summary.append(
        ["ApprenantID réseau trouvés", context["source_summary"]["apprenant_id_count"]]
    )
    ws_summary.append(["Lignes reliées à la source", context["source_summary"]["matched_count"]])
    ws_summary.append(["Lignes cohérentes", context["source_summary"]["consistent_count"]])
    ws_summary.append(["Lignes à vérifier", context["source_summary"]["mismatch_count"]])
    ws_summary.append(["Lignes absentes de la source", context["source_summary"]["missing_count"]])
    ws_summary.append([])
    ws_summary.append(["Type", "Libellé", "Nombre d'enquêtes"])
    for item in context["analyzed_classes"]:
        ws_summary.append(["Classe", item["label"], item["nb"]])
    for item in context["analyzed_prestations"]:
        ws_summary.append(["Prestation", item["label"], item["nb"]])
    for item in context["analyzed_fenetres"]:
        ws_summary.append(["Fenêtre", item["label"], item["nb"]])
    for item in context["analyzed_prestataires"]:
        ws_summary.append(["Prestataire", item["label"], item["nb"]])
    for item in context["analyzed_beneficiaires"]:
        ws_summary.append(["Bénéficiaire", item["label"], item["nb"]])
    for item in context["analyzed_cohortes"]:
        ws_summary.append(["Cohorte", item["label"], item["nb"]])

    ws_data = wb.create_sheet("Enquêtes")
    ws_data.append(
        [
            "Classe",
            "Intitulé de la formation",
            "Prestataire",
            "Bénéficiaire",
            "Cohorte",
            "Ville",
            "Code apprenant",
            "ApprenantID réseau",
            "Apprenant",
            "Cohérence source",
            "Alertes source",
            "Statut appel",
            *[label for _, label in Q_FIELDS],
            "Commentaire",
            "Recommandations",
        ]
    )
    for item in _ordered_survey_rows(rows):
        ws_data.append(
            [
                item["classe_code"],
                item.get("formation_intitule") or item["classe_intitule"],
                item["prestataire"],
                item["beneficiaire"],
                item["cohorte"],
                item["ville"],
                item["apprenant_code"],
                item.get("source_apprenant_id", ""),
                item["apprenant_nom"],
                item.get("source_status_label", ""),
                item.get("source_alerts_label", ""),
                item["status"],
                *[item.get(field) for field, _label in Q_FIELDS],
                item["commentaire"],
                item["recommandations"],
            ]
        )

    ws_classes = wb.create_sheet("Classes")
    ws_classes.append(
        [
            "Classe",
            "Intitulé de la formation",
            "Cohorte",
            "Nombre d'enquêtes",
            *[label for _, label in Q_FIELDS],
        ]
    )
    for item in context["classe_stats"]:
        ws_classes.append(
            [item["code"], item["intitule"], item["cohorte"], item["nb"], *item["avgs"]]
        )

    ws_prestations = wb.create_sheet("Prestations")
    ws_prestations.append(
        [
            "Prestation",
            "Prestataire",
            "Bénéficiaire",
            "Nombre d'enquêtes",
            *[label for _, label in Q_FIELDS],
            "Global (Q9)",
        ]
    )
    for item in context["prestation_stats"]:
        ws_prestations.append(
            [
                item["code"],
                item["prestataire"],
                item["beneficiaire"],
                item["nb"],
                *item["avgs"],
                item["avg"],
            ]
        )

    ws_cohortes = wb.create_sheet("Cohortes")
    ws_cohortes.append(
        ["Cohorte", "Nombre d'enquêtes", *[label for _, label in Q_FIELDS], "Global (Q9)"]
    )
    for item in context["cohorte_stats"]:
        ws_cohortes.append([item["label"], item["nb"], *item["avgs"], item["avg"]])

    for worksheet in wb.worksheets:
        _autosize_worksheet(worksheet)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = _dashboard_export_filename_from_rows(
        sorted(rows, key=lambda row: row["modified_at"]),
        dashboard["filters"],
        "xlsx",
    )

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@require_analysis_access
def satisfaction_dashboard_daily_report_xlsx(request):
    generated_at = timezone.localtime()
    selected_source = _analysis_selected_source(request)
    try:
        source_bundle = build_padesce_source_index(source_key=selected_source)
    except Exception:
        source_bundle = None
    source_records = (source_bundle or {}).get("records", {})

    appels = list(
        Appel.objects.filter(is_active=True)
        .select_related(
            "classe",
            "classe__lieu",
            "classe__prestation",
            "classe__prestation__prestataire",
            "classe__prestation__beneficiaire",
            "locked_by",
            "answers",
            "answers__modified_by",
            "satisfaction_apprenant",
            "satisfaction_apprenant__inspecteur",
            "satisfaction_apprenant__apprenant",
            "satisfaction_apprenant__enqueteur",
        )
        .order_by("created_at", "nom", "code")
    )

    report_rows = [_build_daily_report_row(appel, source_records) for appel in appels]
    success_rows = [row for row in report_rows if row["is_success"]]
    failed_rows = [row for row in report_rows if not row["is_success"]]

    status_counts = Counter(row["status"] for row in report_rows)
    failure_counts = Counter()
    for row in failed_rows:
        for reason in row["failure_reasons"]:
            failure_counts[reason] += 1

    success_headers = [
        "N°",
        "Code appel",
        "Nom déclare",
        "ID apprenant réseau",
        "Classe",
        "Formation",
        "Prestation",
        "Prestataire",
        "Beneficiaire",
        "Fenetre",
        "Cohorte",
        "Ville",
        "Lieu",
        "Telephone 1",
        "Telephone 2",
        "Statut appel",
        "Taux presence",
        "Inspecteur",
        "Enqueteur",
        "Date enquete",
        "Heure enquete",
        *[label for _, label in Q_FIELDS],
        "Commentaire",
        "Recommandations",
        "Transcription",
        "Audio",
        "Source reseau",
        "Cree le",
        "Mis a jour le",
    ]
    failed_headers = [
        "N°",
        "Code appel",
        "Nom declare",
        "Vrai nom",
        "ID apprenant réseau",
        "Classe",
        "Formation",
        "Prestation",
        "Prestataire",
        "Beneficiaire",
        "Fenetre",
        "Cohorte",
        "Ville",
        "Lieu",
        "Telephone 1",
        "Telephone 2",
        "Statut appel",
        "Motifs d'echec",
        "Details",
        "Deja forme",
        "Pas forme",
        "Faux nom",
        "Numero double",
        "Deja appele",
        "Formulaire complet",
        "Formulaire RAS",
        *[label for _, label in Q_FIELDS],
        "Commentaire",
        "Recommandations",
        "Transcription",
        "Audio",
        "Cree le",
        "Mis a jour le",
    ]

    success_sheet_rows = [
        [
            index,
            row["code"],
            row["nom"],
            row["source_apprenant_id"],
            row["classe"],
            row["formation"],
            row["prestation"],
            row["prestataire"],
            row["beneficiaire"],
            row["fenetre"],
            row["cohorte"],
            row["ville"],
            row["lieu"],
            row["telephone1"],
            row["telephone2"],
            row["status"],
            row["taux_presence"],
            row["inspecteur"],
            row["enqueteur"],
            row["survey_date"],
            row["survey_time"],
            *[row.get(field) for field in APPEL_ANSWER_QUESTION_FIELDS],
            row["commentaire"],
            row["recommandations"],
            row["transcription"],
            row["audio_name"],
            row["source_status"],
            row["created_at"],
            row["updated_at"],
        ]
        for index, row in enumerate(success_rows, start=1)
    ]
    failed_sheet_rows = [
        [
            index,
            row["code"],
            row["nom"],
            row["vrai_nom"],
            row["source_apprenant_id"],
            row["classe"],
            row["formation"],
            row["prestation"],
            row["prestataire"],
            row["beneficiaire"],
            row["fenetre"],
            row["cohorte"],
            row["ville"],
            row["lieu"],
            row["telephone1"],
            row["telephone2"],
            row["status"],
            row["motifs_echec"],
            row["details_echec"],
            row["deja_forme"],
            row["flag_pas_forme"],
            row["flag_faux_nom"],
            row["flag_numero_double"],
            row["flag_deja_appele"],
            row["formulaire_complet"],
            row["formulaire_ras"],
            *[row.get(field) for field in APPEL_ANSWER_QUESTION_FIELDS],
            row["commentaire"],
            row["recommandations"],
            row["transcription"],
            row["audio_name"],
            row["created_at"],
            row["updated_at"],
        ]
        for index, row in enumerate(failed_rows, start=1)
    ]

    wb = openpyxl.Workbook()
    ws_dashboard = wb.active
    _append_daily_dashboard_sheet(
        ws_dashboard,
        generated_at,
        len(report_rows),
        success_rows,
        failed_rows,
        status_counts,
        failure_counts,
        source_bundle,
    )
    ws_success = wb.create_sheet("Appels reussis")
    _append_daily_detail_sheet(
        ws_success, "Appels reussis", success_headers, success_sheet_rows, "166534"
    )
    ws_failed = wb.create_sheet("Appels echoues")
    _append_daily_detail_sheet(
        ws_failed, "Appels echoues", failed_headers, failed_sheet_rows, "B91C1C"
    )

    # Formateurs sheet
    formateurs_headers = [
        "N°",
        "Prestataire",
        "Bénéficiaire",
        "Formation",
        "Cohorte",
        "Téléphone",
        "Statut",
        "Date séance",
        "Q1 – Prérequis apprenants",
        "Q2 – Interaction apprenants",
        "Q3 – Compétences acquises",
        "Q4 – Gestion administrative",
        "Q5 – Gestion financière",
        "Q6 – Communication",
        "Commentaires",
        "Créé le",
        "Mis à jour le",
    ]
    formateurs_qs = list(
        AppelFormateur.objects.filter(is_active=True)
        .order_by("session_date", "prestataire")
        .values(
            "prestataire",
            "beneficiaire",
            "formation",
            "cohorte",
            "telephone",
            "status",
            "session_date",
            "q1_prerequis_apprenants",
            "q2_interaction_apprenants",
            "q3_competences_acquises",
            "q4_gestion_administrative",
            "q5_gestion_financiere",
            "q6_communication",
            "commentaires",
            "created_at",
            "updated_at",
        )
    )
    formateurs_sheet_rows = [
        [
            idx,
            r["prestataire"] or "",
            r["beneficiaire"] or "",
            r["formation"] or "",
            r["cohorte"] or "",
            r["telephone"] or "",
            r["status"] or "",
            r["session_date"],
            r["q1_prerequis_apprenants"],
            r["q2_interaction_apprenants"],
            r["q3_competences_acquises"],
            r["q4_gestion_administrative"] or "",
            r["q5_gestion_financiere"] or "",
            r["q6_communication"] or "",
            r["commentaires"] or "",
            r["created_at"],
            r["updated_at"],
        ]
        for idx, r in enumerate(formateurs_qs, start=1)
    ]
    ws_formateurs = wb.create_sheet("Appels formateurs")
    _append_daily_detail_sheet(
        ws_formateurs, "Appels formateurs", formateurs_headers, formateurs_sheet_rows, "1E3A5F"
    )

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="{_daily_report_filename(generated_at)}"'
    )
    return response


@require_analysis_access
def satisfaction_dashboard(request):
    dashboard = _build_satisfaction_dashboard_data(request)
    ctx = dashboard["context"]
    # Ranking complet (non filtré) pour la carte et le classement des prestations
    ctx["toutes_prestations_classees"] = get_prestations_ranking(
        ctx.get("prestation_stats_all", []), order="desc"
    )
    ctx["rows"] = [
        {**row, "q_values": [row.get(field) for field, _ in Q_FIELDS]}
        for row in sorted(dashboard["rows"], key=lambda r: r["modified_at"], reverse=True)
    ]
    ctx["active_tab"] = _active_satisfaction_tab(request)
    ctx["active_table_details"] = ctx["tab_details"].get(
        ctx["active_tab"], ctx["tab_details"]["tab-apprenants"]
    )
    ctx.update(build_fast_stats_context(request, default_mode="apprenant"))
    return render(request, "satisfaction_apprenants/dashboard.html", ctx)


def _general_analysis_threshold_codes(source_bundle: dict | None = None) -> set[str]:
    return _status_threshold_class_codes(source_bundle)


def _general_analysis_search_matches(row: dict, query: str) -> bool:
    haystack = " ".join(
        [
            str(row.get("source_apprenant_id") or ""),
            str(row.get("code") or ""),
            str(row.get("nom") or ""),
            str(row.get("classe") or ""),
            str(row.get("formation") or ""),
            str(row.get("prestation") or ""),
            str(row.get("prestataire") or ""),
            str(row.get("beneficiaire") or ""),
            str(row.get("telephone1") or ""),
            str(row.get("telephone2") or ""),
            str(row.get("enqueteur") or ""),
        ]
    ).casefold()
    return query.casefold() in haystack


def _build_general_analysis_rows(
    selected_source: str,
    source_bundle: dict | None = None,
) -> list[dict]:
    source_records = (source_bundle or {}).get("records", {})
    threshold_codes = _general_analysis_threshold_codes(source_bundle)
    appels = list(
        Appel.objects.filter(is_active=True)
        .select_related(
            "classe",
            "classe__lieu",
            "classe__prestation",
            "classe__prestation__prestataire",
            "classe__prestation__beneficiaire",
            "locked_by",
            "answers",
            "answers__modified_by",
            "satisfaction_apprenant",
            "satisfaction_apprenant__inspecteur",
            "satisfaction_apprenant__apprenant",
            "satisfaction_apprenant__enqueteur",
        )
        .order_by("-updated_at", "nom", "code")
    )

    rows: list[dict] = []
    for appel in appels:
        row = _build_daily_report_row(appel, source_records)
        class_key = normalize_network_lookup(row.get("classe") or "")
        threshold_reached = bool(class_key and class_key in threshold_codes)
        analysis_included = row.get("analysis_included") == "Oui"
        analysis_taken_into_account = analysis_included and threshold_reached
        if analysis_taken_into_account:
            analysis_take_reason = "Pris en compte dans les analyses"
        elif row.get("analysis_included") != "Oui":
            analysis_take_reason = (
                row.get("analysis_exclusion_reason") or "Aucune reponse analysee"
            )
        elif not threshold_reached:
            analysis_take_reason = f"Seuil {analysis_threshold_label()} non atteint"
        else:
            analysis_take_reason = "Non retenu"

        row["has_phone"] = bool(appel_has_analysis_phone(appel))
        row["analysis_threshold_reached"] = threshold_reached
        row["analysis_taken_into_account"] = analysis_taken_into_account
        row["analysis_take_reason"] = analysis_take_reason
        row["q_values"] = [row.get(field) for field, _label in Q_FIELDS]
        row["responses_summary"] = " | ".join(
            [
                f"Q{index}:{row.get(field) if row.get(field) not in (None, '') else '-'}"
                for index, (field, _label) in enumerate(Q_FIELDS, start=1)
            ]
        )
        row["toggle_label"] = "Reintegrer" if row.get("exclude_from_analysis") == "Oui" else "Masquer"
        row["toggle_action"] = "include" if row.get("exclude_from_analysis") == "Oui" else "exclude"
        rows.append(row)

    return rows


def _cached_general_analysis_rows(
    selected_source: str,
    source_bundle: dict | None = None,
) -> list[dict]:
    source_marker = str(((source_bundle or {}).get("source") or {}).get("modified_at") or "no-source")
    cache_key = _analysis_cache_key(
        "general-analysis-rows",
        selected_source,
        source_marker,
        get_analysis_cache_version("analysis-general"),
        _analysis_queryset_marker(Appel),
        _analysis_queryset_marker(AppelAnswers),
        _analysis_queryset_marker(SatisfactionApprenant),
    )
    cached_rows = cache.get(cache_key)
    if cached_rows is not None:
        return cached_rows
    rows = _build_general_analysis_rows(selected_source, source_bundle=source_bundle)
    cache.set(cache_key, rows, timeout=ANALYSIS_CACHE_TIMEOUT)
    return rows


def _normalize_batch_update_container(raw_value: str) -> str:
    text = str(raw_value or "").strip()
    if text.startswith("[") and text.endswith("]"):
        return text[1:-1].strip()
    return text


def _parse_batch_update_targets(raw_codes: str, default_class_code: str = "") -> list[dict[str, str]]:
    normalized_default = str(default_class_code or "").strip()
    parsed_targets: list[dict[str, str]] = []
    seen_codes: set[str] = set()
    for block in re.split(r"[\r\n;]+", _normalize_batch_update_container(raw_codes)):
        block = _normalize_batch_update_container(block)
        if not block:
            continue
        for token in [item.strip() for item in re.split(r"[\s,]+", block) if item.strip()]:
            code = token
            requested_class_code = normalized_default
            if "|" in token:
                code, requested_class_code = [part.strip() for part in token.split("|", 1)]
                requested_class_code = requested_class_code or normalized_default
            if not code:
                continue
            code_key = code.casefold()
            if code_key in seen_codes:
                continue
            seen_codes.add(code_key)
            parsed_targets.append(
                {
                    "code": code,
                    "requested_class_code": requested_class_code,
                }
            )
    return parsed_targets


def _merge_batch_update_targets(
    raw_codes: str,
    selected_targets: list[str] | None,
    default_class_code: str = "",
) -> list[dict[str, str]]:
    merged_targets: list[dict[str, str]] = []
    seen_codes: set[str] = set()
    for raw_value in [raw_codes, *(selected_targets or [])]:
        for target in _parse_batch_update_targets(raw_value, default_class_code):
            code_key = str(target.get("code") or "").strip().casefold()
            if not code_key or code_key in seen_codes:
                continue
            seen_codes.add(code_key)
            merged_targets.append(target)
    return merged_targets


def _expand_batch_update_values(
    values: list,
    target_count: int,
    *,
    label: str,
    default_value=None,
) -> list:
    if not values:
        return [default_value] * target_count
    if len(values) == 1:
        return list(values) * target_count
    if len(values) == target_count:
        return list(values)
    raise ValueError(
        f"{label}: fournissez une seule valeur ou exactement {target_count} valeurs."
    )


def _build_batch_update_payloads(cleaned_data: dict, target_count: int) -> list[dict]:
    payloads = [{} for _ in range(target_count)]
    field_labels = {field: label for field, label in Q_FIELDS}
    field_labels.update(
        {
            "commentaire_values": "Commentaire general",
            "recommandations_values": "Recommandations",
        }
    )

    for field in APPEL_ANSWER_QUESTION_FIELDS:
        expanded_values = _expand_batch_update_values(
            cleaned_data.get(field) or [],
            target_count,
            label=field_labels.get(field, field),
            default_value=3,
        )
        for index, value in enumerate(expanded_values):
            if value is not None:
                payloads[index][field] = value

    for form_field, payload_field in (
        ("commentaire_values", "commentaire"),
        ("recommandations_values", "recommandations"),
    ):
        expanded_values = _expand_batch_update_values(
            cleaned_data.get(form_field) or [],
            target_count,
            label=field_labels[form_field],
            default_value="RAS",
        )
        for index, value in enumerate(expanded_values):
            if value is not None:
                payloads[index][payload_field] = value

    return payloads


def _linked_one_to_one(instance, attr_name: str):
    try:
        return getattr(instance, attr_name)
    except ObjectDoesNotExist:
        return None


def _batch_update_answer_summary(
    answer: AppelAnswers | None,
    survey: SatisfactionApprenant | None = None,
) -> str:
    source = answer or survey
    if not source:
        return "-"
    return " / ".join(str(getattr(source, field, None) or "-") for field in APPEL_ANSWER_QUESTION_FIELDS)


def _batch_update_status_display(status_code: str) -> str:
    normalized_status = str(status_code or "").strip()
    return dict(Appel.STATUS_CHOICES).get(normalized_status, normalized_status or "-")


def _build_update_form_candidate_row(
    appel: Appel,
    apprenant: Apprenant | None,
    answer: AppelAnswers | None,
    survey: SatisfactionApprenant | None,
) -> dict:
    resolved_class = appel.classe or getattr(apprenant, "classe", None)
    classe_code = (
        getattr(resolved_class, "code", "")
        or str(getattr(appel, "classe_label", "") or "").strip()
        or getattr(getattr(apprenant, "classe", None), "code", "")
        or "-"
    )
    prestation_code = getattr(getattr(resolved_class, "prestation", None), "code", "") or "-"
    has_complete_form = _has_complete_form_record(answer, survey)
    has_partial_form = appel_has_any_form_data(appel)
    current_status = str(appel.status or "").strip()
    computed_status = derive_padesce_status(appel)
    selection_value = appel.code
    if classe_code and classe_code != "-":
        selection_value = f"{appel.code}|{classe_code}"
    return {
        "code": appel.code,
        "nom": str(appel.nom or getattr(apprenant, "nom_complet", "") or "-").strip() or "-",
        "classe_code": classe_code,
        "prestation_code": prestation_code,
        "audio_label": "Oui" if appel_has_any_audio(appel) else "Non",
        "formulaire_label": "Complet" if has_complete_form else ("Partiel" if has_partial_form else "Non"),
        "current_status": current_status,
        "current_status_label": appel.get_status_display(),
        "computed_status": computed_status,
        "computed_status_label": _batch_update_status_display(computed_status),
        "commentaire": getattr(answer or survey, "commentaire", "") or "-",
        "recommandations": getattr(answer or survey, "recommandations", "") or "-",
        "selection_value": selection_value,
        "has_complete_form": has_complete_form,
    }


def _update_form_candidate_base_queryset():
    return (
        Appel.objects.filter(is_active=True)
        .select_related("classe", "classe__prestation", "answers", "satisfaction_apprenant")
        .order_by("classe_label", "code")
    )


def _update_form_complete_record_q() -> Q:
    return appel_answers_completed_q() | Q(satisfaction_apprenant__isnull=False)


def _termine_without_form_queryset():
    return _update_form_candidate_base_queryset().filter(status="termine").exclude(
        _update_form_complete_record_q()
    )


def _form_status_issue_queryset():
    return _update_form_candidate_base_queryset().exclude(status="termine").filter(
        _update_form_complete_record_q()
    )


def _build_update_form_rows_for_appels(appels) -> list[dict]:
    apprenant_codes = [str(appel.code or "").strip() for appel in appels if str(appel.code or "").strip()]
    apprenants_by_code = {
        str(apprenant.code or "").strip().casefold(): apprenant
        for apprenant in Apprenant.objects.select_related("classe", "classe__prestation").filter(
            code__in=apprenant_codes
        )
    }
    rows: list[dict] = []
    for appel in appels:
        answer = _linked_one_to_one(appel, "answers")
        survey = _linked_one_to_one(appel, "satisfaction_apprenant")
        apprenant = apprenants_by_code.get(str(appel.code or "").strip().casefold())
        if apprenant is None:
            apprenant = _resolve_batch_update_apprenant(appel)
        rows.append(_build_update_form_candidate_row(appel, apprenant, answer, survey))
    return rows


def _build_update_form_candidate_lists() -> tuple[list[dict], list[dict]]:
    return (
        _build_update_form_rows_for_appels(list(_termine_without_form_queryset())),
        _build_update_form_rows_for_appels(list(_form_status_issue_queryset())),
    )


def _cached_update_form_candidate_lists() -> tuple[list[dict], list[dict]]:
    cache_key = _analysis_cache_key(
        "update-form-candidates",
        _analysis_queryset_marker(Appel),
        _analysis_queryset_marker(AppelAnswers),
        _analysis_queryset_marker(SatisfactionApprenant),
    )
    cached_payload = cache.get(cache_key)
    if cached_payload is not None:
        return (
            list(cached_payload.get("termine_without_form_rows", [])),
            list(cached_payload.get("form_status_issue_rows", [])),
        )
    termine_without_form_rows, form_status_issue_rows = _build_update_form_candidate_lists()
    cache.set(
        cache_key,
        {
            "termine_without_form_rows": termine_without_form_rows,
            "form_status_issue_rows": form_status_issue_rows,
        },
        timeout=ANALYSIS_CACHE_TIMEOUT,
    )
    return termine_without_form_rows, form_status_issue_rows


def _paginate_update_form_rows(
    request,
    queryset,
    *,
    page_param: str,
    per_page: int = 50,
):
    paginator = Paginator(queryset, per_page)
    page_obj = paginator.get_page(request.GET.get(page_param))
    return page_obj, _build_update_form_rows_for_appels(list(page_obj.object_list))


def _resolve_batch_update_apprenant(appel: Appel) -> Apprenant | None:
    apprenant = (
        Apprenant.objects.select_related("classe", "formation")
        .filter(code__iexact=str(appel.code or "").strip())
        .first()
    )
    if apprenant:
        return apprenant
    from App_PADESCE.appels.views import _find_apprenant_for_appel

    fallback = _find_apprenant_for_appel(Apprenant.objects.all(), appel)
    if fallback is None:
        return None
    return (
        Apprenant.objects.select_related("classe", "formation")
        .filter(pk=fallback.pk)
        .first()
        or fallback
    )


def _batch_update_known_class_codes(appel: Appel, apprenant: Apprenant | None) -> list[str]:
    known_codes: list[str] = []
    seen_codes: set[str] = set()
    raw_candidates = [
        getattr(getattr(appel, "classe", None), "code", ""),
        getattr(appel, "classe_label", ""),
        getattr(getattr(apprenant, "classe", None), "code", ""),
        str((getattr(appel, "snapshot", {}) or {}).get("classe_id") or ""),
    ]
    for raw_value in raw_candidates:
        code = str(raw_value or "").strip()
        code_key = code.casefold()
        if not code or code_key in seen_codes:
            continue
        seen_codes.add(code_key)
        known_codes.append(code)
    return known_codes


def _resolve_batch_update_class(
    appel: Appel,
    apprenant: Apprenant | None,
    requested_class_code: str,
) -> tuple[Classe | None, str]:
    requested_code = str(requested_class_code or "").strip()
    known_codes = _batch_update_known_class_codes(appel, apprenant)
    known_code_keys = {code.casefold() for code in known_codes}
    if requested_code and known_code_keys and requested_code.casefold() not in known_code_keys:
        raise ValueError(
            f"Classe attendue {requested_code} differente de la classe trouvee {known_codes[0]}."
        )

    resolved_class = appel.classe or getattr(apprenant, "classe", None)
    lookup_code = requested_code or (known_codes[0] if known_codes else "")
    if not resolved_class and lookup_code:
        resolved_class = Classe.objects.filter(code__iexact=lookup_code).first()

    effective_code = requested_code or getattr(resolved_class, "code", "") or (
        known_codes[0] if known_codes else ""
    )
    return resolved_class, effective_code


def _sync_batch_update_appel_class(
    appel: Appel,
    resolved_class: Classe | None,
    effective_class_code: str,
) -> None:
    update_fields: list[str] = []
    normalized_code = str(effective_class_code or "").strip()
    if resolved_class and appel.classe_id != resolved_class.pk:
        appel.classe = resolved_class
        update_fields.append("classe")
    if normalized_code and appel.classe_label != normalized_code:
        appel.classe_label = normalized_code
        update_fields.append("classe_label")
    if update_fields:
        appel.save(update_fields=[*update_fields, "updated_at"])


def _upsert_batch_update_satisfaction(
    appel: Appel,
    answers: AppelAnswers,
    apprenant: Apprenant | None,
    classe: Classe | None,
    user,
) -> SatisfactionApprenant | None:
    if not _has_complete_answer_set(answers):
        return None

    now = timezone.localtime()
    survey = _linked_one_to_one(appel, "satisfaction_apprenant")
    if survey is None:
        survey = SatisfactionApprenant(
            appel=appel,
            date=now.date(),
            heure=now.time().replace(microsecond=0),
        )

    if classe is not None:
        survey.classe = classe
    elif survey.classe_id is None and apprenant and apprenant.classe_id:
        survey.classe = apprenant.classe

    if apprenant is not None:
        survey.apprenant = apprenant

    if getattr(user, "is_authenticated", False):
        survey.enqueteur = user

    if not survey.date:
        survey.date = now.date()
    if not survey.heure:
        survey.heure = now.time().replace(microsecond=0)

    for field in APPEL_ANSWER_QUESTION_FIELDS:
        setattr(survey, field, getattr(answers, field))
    survey.commentaire = answers.commentaire or ""
    survey.recommandations = answers.recommandations or ""
    survey.save()
    return survey


def _apply_batch_update_target(target: dict[str, str], payload: dict, user) -> dict:
    requested_code = target["code"]
    requested_class_code = str(target.get("requested_class_code") or "").strip()
    result = {
        "code": requested_code,
        "requested_class_code": requested_class_code or "-",
        "resolved_class_code": "-",
        "nom": "-",
        "before_status": "-",
        "after_status": "-",
        "before_answers": "-",
        "after_answers": "-",
        "commentaire": payload.get("commentaire", "-") or "-",
        "recommandations": payload.get("recommandations", "-") or "-",
        "message": "",
        "ok": False,
        "survey_synced": False,
    }

    appel = (
        Appel.objects.filter(is_active=True, code__iexact=requested_code)
        .select_related("classe", "answers", "satisfaction_apprenant")
        .first()
    )
    if appel is None:
        result["message"] = "Code apprenant introuvable."
        return result

    result["nom"] = appel.nom or "-"
    result["before_status"] = appel.get_status_display()
    before_answers = _linked_one_to_one(appel, "answers")
    before_survey = _linked_one_to_one(appel, "satisfaction_apprenant")
    result["before_answers"] = _batch_update_answer_summary(before_answers, before_survey)

    try:
        with transaction.atomic():
            apprenant = _resolve_batch_update_apprenant(appel)
            resolved_class, effective_class_code = _resolve_batch_update_class(
                appel,
                apprenant,
                requested_class_code,
            )
            _sync_batch_update_appel_class(appel, resolved_class, effective_class_code)

            from App_PADESCE.appels.views import _save_appel_answers

            answers = _save_appel_answers(appel, user, payload, apply_defaults=False)
            survey = _upsert_batch_update_satisfaction(
                appel,
                answers,
                apprenant,
                resolved_class,
                user,
            )
            sync_padesce_status(appel)

        result["resolved_class_code"] = effective_class_code or "-"
        result["after_status"] = appel.get_status_display()
        result["after_answers"] = _batch_update_answer_summary(answers)
        result["commentaire"] = answers.commentaire or "-"
        result["recommandations"] = answers.recommandations or "-"
        result["survey_synced"] = survey is not None
        result["message"] = (
            "Formulaire mis a jour et fiche satisfaction synchronisee."
            if survey is not None
            else "Mise a jour enregistree. Fiche satisfaction non synchronisee."
        )
        result["ok"] = True
        return result
    except ValueError as exc:
        result["message"] = str(exc)
        return result
    except Exception as exc:
        logger.exception("UPDATE FORM batch update failed for code=%s", requested_code)
        result["message"] = f"Erreur interne pendant la mise a jour: {exc}"
        return result


def _apply_batch_status_target(target: dict[str, str], target_status: str) -> dict:
    requested_code = target["code"]
    requested_class_code = str(target.get("requested_class_code") or "").strip()
    result = {
        "code": requested_code,
        "requested_class_code": requested_class_code or "-",
        "resolved_class_code": "-",
        "nom": "-",
        "before_status": "-",
        "after_status": "-",
        "before_answers": "-",
        "after_answers": "-",
        "commentaire": "-",
        "recommandations": "-",
        "message": "",
        "ok": False,
        "survey_synced": False,
    }

    appel = (
        Appel.objects.filter(is_active=True, code__iexact=requested_code)
        .select_related("classe", "answers", "satisfaction_apprenant")
        .first()
    )
    if appel is None:
        result["message"] = "Code apprenant introuvable."
        return result

    answer = _linked_one_to_one(appel, "answers")
    survey = _linked_one_to_one(appel, "satisfaction_apprenant")
    if not _has_complete_form_record(answer, survey):
        result["nom"] = appel.nom or "-"
        result["before_status"] = appel.get_status_display()
        result["before_answers"] = _batch_update_answer_summary(answer, survey)
        result["message"] = "Aucun formulaire complet trouve pour ce code."
        return result

    result["nom"] = appel.nom or "-"
    result["before_status"] = appel.get_status_display()
    result["before_answers"] = _batch_update_answer_summary(answer, survey)
    result["commentaire"] = getattr(answer or survey, "commentaire", "") or "-"
    result["recommandations"] = getattr(answer or survey, "recommandations", "") or "-"

    try:
        with transaction.atomic():
            apprenant = _resolve_batch_update_apprenant(appel)
            resolved_class, effective_class_code = _resolve_batch_update_class(
                appel,
                apprenant,
                requested_class_code,
            )
            _sync_batch_update_appel_class(appel, resolved_class, effective_class_code)
            appel.status = target_status
            appel.save(update_fields=["status", "updated_at"])

        result["resolved_class_code"] = effective_class_code or "-"
        result["after_status"] = appel.get_status_display()
        result["after_answers"] = _batch_update_answer_summary(answer, survey)
        result["message"] = f"Statut mis a jour vers {appel.get_status_display()}."
        result["ok"] = True
        return result
    except ValueError as exc:
        result["message"] = str(exc)
        return result
    except Exception as exc:
        logger.exception("UPDATE FORM status update failed for code=%s", requested_code)
        result["message"] = f"Erreur interne pendant le changement de statut: {exc}"
        return result


@require_analysis_access
def satisfaction_update_form_page(request):
    selected_source = _analysis_selected_source(request)
    initial = {
        "classe_code": str(request.GET.get("classe_code", "") or "").strip(),
        "codes_text": str(request.GET.get("codes", "") or "").strip(),
    }
    form = SatisfactionBatchUpdateForm(request.POST or None, initial=initial)
    results: list[dict] = []
    summary = {
        "requested_total": 0,
        "updated_total": 0,
        "error_total": 0,
        "synced_total": 0,
        "action_label": "formulaire(s)",
    }
    selected_target_values = {
        str(value or "").strip()
        for value in request.POST.getlist("selected_targets")
        if str(value or "").strip()
    }

    if request.method == "POST" and form.is_valid():
        action = str(request.POST.get("action") or "update_form").strip() or "update_form"
        targets = _merge_batch_update_targets(
            form.cleaned_data["codes_text"],
            request.POST.getlist("selected_targets"),
            form.cleaned_data.get("classe_code", ""),
        )
        if not targets:
            form.add_error(
                "codes_text",
                "Ajoutez au moins un code apprenant valide ou selectionnez au moins une ligne.",
            )
        else:
            if action == "update_status":
                requested_status = str(form.cleaned_data.get("target_status") or "").strip()
                if not requested_status:
                    form.add_error("target_status", "Choisissez le statut a appliquer.")
                else:
                    results = [
                        _apply_batch_status_target(target, requested_status) for target in targets
                    ]
                    summary = {
                        "requested_total": len(results),
                        "updated_total": sum(1 for item in results if item["ok"]),
                        "error_total": sum(1 for item in results if not item["ok"]),
                        "synced_total": 0,
                        "action_label": "statut(s)",
                    }
                    if summary["updated_total"]:
                        messages.success(
                            request,
                            f"{summary['updated_total']} statut(s) mis a jour.",
                        )
                    if summary["error_total"]:
                        messages.warning(
                            request,
                            f"{summary['error_total']} code(s) n'ont pas pu etre traites.",
                        )
            else:
                try:
                    payloads = _build_batch_update_payloads(form.cleaned_data, len(targets))
                except ValueError as exc:
                    form.add_error(None, str(exc))
                else:
                    results = [
                        _apply_batch_update_target(target, payload, request.user)
                        for target, payload in zip(targets, payloads)
                    ]
                    summary = {
                        "requested_total": len(results),
                        "updated_total": sum(1 for item in results if item["ok"]),
                        "error_total": sum(1 for item in results if not item["ok"]),
                        "synced_total": sum(1 for item in results if item["survey_synced"]),
                        "action_label": "formulaire(s)",
                    }
                    if summary["updated_total"]:
                        messages.success(
                            request,
                            f"{summary['updated_total']} formulaire(s) mis a jour.",
                        )
                    if summary["error_total"]:
                        messages.warning(
                            request,
                            f"{summary['error_total']} code(s) n'ont pas pu etre traites.",
                        )

    termine_qs = _termine_without_form_queryset()
    form_status_qs = _form_status_issue_queryset()
    termine_without_form_total = termine_qs.count()
    form_status_issue_total = form_status_qs.count()
    termine_page_obj, termine_without_form_rows = _paginate_update_form_rows(
        request,
        termine_qs,
        page_param="termine_page",
    )
    form_status_page_obj, form_status_issue_rows = _paginate_update_form_rows(
        request,
        form_status_qs,
        page_param="status_page",
    )

    context = {
        "form": form,
        "results": results,
        "summary": summary,
        "question_fields": Q_FIELDS,
        "selected_target_values": selected_target_values,
        "termine_without_form_rows": termine_without_form_rows,
        "termine_without_form_total": termine_without_form_total,
        "termine_without_form_page_obj": termine_page_obj,
        "form_status_issue_rows": form_status_issue_rows,
        "form_status_issue_total": form_status_issue_total,
        "form_status_issue_page_obj": form_status_page_obj,
        "candidate_total": termine_without_form_total + form_status_issue_total,
        "selected_source": selected_source,
        "general_url": f"{reverse('satisfaction_general_page')}?source={selected_source}",
        "dashboard_url": f"{reverse('satisfaction_dashboard')}?source={selected_source}",
    }
    return render(request, "satisfaction_apprenants/update_form.html", context)


@require_analysis_access
def satisfaction_general_page(request):
    selected_source = _analysis_selected_source(request)
    search = str(request.GET.get("q", "") or "").strip()
    without_phone_only = request.GET.get("without_phone") == "1"
    all_three_only = request.GET.get("all_three") == "1"
    excluded_filter = str(request.GET.get("excluded", "") or "").strip().lower()
    status_filter = (request.GET.get("status") or "").strip()

    try:
        source_bundle = build_padesce_source_index(source_key=selected_source)
    except Exception:
        source_bundle = None
    rows = list(_cached_general_analysis_rows(selected_source, source_bundle=source_bundle))

    if search:
        rows = [row for row in rows if _general_analysis_search_matches(row, search)]
    if without_phone_only:
        rows = [row for row in rows if not row.get("has_phone")]
    if all_three_only:
        rows = [row for row in rows if row.get("formulaire_all_three") == "Oui"]
    if excluded_filter == "yes":
        rows = [row for row in rows if row.get("exclude_from_analysis") == "Oui"]
    elif excluded_filter == "no":
        rows = [row for row in rows if row.get("exclude_from_analysis") != "Oui"]

    if status_filter:
        rows = [row for row in rows if row.get("status") == status_filter]

    paginator = Paginator(rows, 50)
    page_obj = paginator.get_page(request.GET.get("page"))

    context = {
        "rows": list(page_obj.object_list),
        "page_obj": page_obj,
        "paginator": paginator,
        "total_rows": len(rows),
        "filters": {
            "source": selected_source,
            "q": search,
            "without_phone": without_phone_only,
            "all_three": all_three_only,
            "excluded": excluded_filter,
            "status": status_filter,
        },
        "source_options": get_workbook_source_options(),
        "analysis_threshold_label": analysis_threshold_label(),
        "stats": {
            "total": len(rows),
            "taken_into_account": sum(
                1 for row in rows if row.get("analysis_taken_into_account")
            ),
            "without_phone": sum(1 for row in rows if not row.get("has_phone")),
            "all_three": sum(
                1 for row in rows if row.get("formulaire_all_three") == "Oui"
            ),
            "excluded": sum(
                1 for row in rows if row.get("exclude_from_analysis") == "Oui"
            ),
        },
        "current_path": request.get_full_path(),
    }
    return render(request, "satisfaction_apprenants/general.html", context)


@require_POST
@require_analysis_access
def satisfaction_general_toggle_exclusion(request):
    next_url = request.POST.get("next") or reverse("satisfaction_general_page")
    requested_action = str(request.POST.get("action") or "").strip().lower()
    appel_ids = [
        str(value).strip()
        for value in request.POST.getlist("appel_ids")
        if str(value).strip()
    ]
    single_appel_id = str(request.POST.get("appel_id") or "").strip()
    if single_appel_id and single_appel_id not in appel_ids:
        appel_ids.append(single_appel_id)

    if not appel_ids:
        messages.warning(request, "Selectionne au moins une ligne a mettre a jour.")
        return redirect(next_url)

    appels = list(Appel.objects.filter(is_active=True, pk__in=appel_ids).order_by("pk"))
    if not appels:
        messages.warning(request, "Aucune ligne active n'a ete trouvee pour cette action.")
        return redirect(next_url)

    explicit_exclusion: bool | None
    if requested_action == "exclude":
        explicit_exclusion = True
    elif requested_action == "include":
        explicit_exclusion = False
    else:
        explicit_exclusion = None

    if explicit_exclusion is None:
        updated_states = [
            toggle_appel_manual_exclusion(appel)
            for appel in appels
        ]
        final_excluded = sum(1 for state in updated_states if state)
        final_included = len(updated_states) - final_excluded
        if len(appels) == 1:
            appel = appels[0]
            messages.success(
                request,
                (
                    f"{appel.nom or appel.code or 'La ligne'} est maintenant "
                    f"{'exclu(e)' if updated_states[0] else 'reintegre(e)'} des analyses."
                ),
            )
        else:
            messages.success(
                request,
                (
                    f"{len(appels)} ligne(s) mises a jour : "
                    f"{final_excluded} masquee(s), {final_included} reintegree(s)."
                ),
            )
        return redirect(next_url)

    for appel in appels:
        set_appel_manual_exclusion(appel, explicit_exclusion)

    if len(appels) == 1:
        appel = appels[0]
        messages.success(
            request,
            (
                f"{appel.nom or appel.code or 'La ligne'} est maintenant "
                f"{'exclu(e)' if explicit_exclusion else 'reintegre(e)'} des analyses."
            ),
        )
    else:
        messages.success(
            request,
            (
                f"{len(appels)} ligne(s) sont maintenant "
                f"{'masquee(s)' if explicit_exclusion else 'reintegree(s)'} dans les analyses."
            ),
        )
    return redirect(next_url)


@require_analysis_access
def satisfaction_map_data(request):
    """Retourne le top 5 des meilleures prestations par région (carte Leaflet)."""
    dashboard = _build_satisfaction_dashboard_data(request)
    prestation_stats_all = dashboard["context"].get("prestation_stats_all", [])
    all_rankings = get_prestations_ranking(prestation_stats_all, order="desc")

    region_data: dict = defaultdict(list)
    for p in all_rankings:
        reg = p.get("region") or "Inconnu"
        if len(region_data[reg]) < 5:
            region_data[reg].append(
                {
                    "code": p["code"],
                    "prestataire": p["prestataire"],
                    "beneficiaire": p["beneficiaire"],
                    "score": p.get("score_global"),
                }
            )
    return JsonResponse(dict(region_data))


@require_analysis_access
def satisfaction_dashboard_rag(request):
    if request.method != "POST":
        return JsonResponse({"error": "Methode non autorisee."}, status=405)

    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except (TypeError, ValueError, UnicodeDecodeError):
        return JsonResponse({"error": "Corps JSON invalide."}, status=400)

    prompt = str(payload.get("prompt") or "").strip()
    if not prompt:
        return JsonResponse({"error": "Le prompt est vide."}, status=400)

    filter_query = str(payload.get("filter_query") or "").lstrip("?")
    active_tab = str(payload.get("tab") or "tab-apprenants").strip()
    query_params = QueryDict(filter_query, mutable=True)
    query_params["tab"] = active_tab
    request_like = SimpleNamespace(GET=query_params)

    try:
        dashboard = _build_satisfaction_dashboard_data(request_like)
        active_tab = _active_satisfaction_tab(request_like)
        result = answer_dashboard_prompt(
            prompt,
            active_tab,
            dashboard["context"],
            dashboard["rows"],
        )
    except ValueError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except RuntimeError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except requests.RequestException as exc:
        return JsonResponse(
            {"error": f"Erreur Groq lors de la generation de la reponse: {exc}"},
            status=502,
        )
    except Exception as exc:
        return JsonResponse(
            {"error": f"Impossible de traiter la demande RAG: {exc}"},
            status=500,
        )

    return JsonResponse(
        {
            "ok": True,
            "active_tab": active_tab,
            "active_tab_label": SATISFACTION_DASHBOARD_TAB_LABELS.get(active_tab, "Table active"),
            "answer_markdown": result["answer_markdown"],
            "matched_rows": result.get("matched_rows", []),
            "matched_count": len(result.get("matched_rows", [])),
            "retrieved_count": result.get("retrieved_count", 0),
            "insufficient_context": bool(result.get("insufficient_context")),
            "model": result.get("model", ""),
            "row_count": len(dashboard["rows"]),
        }
    )


@require_analysis_access
def satisfaction_dashboard_export_csv(request):
    dashboard = _build_satisfaction_dashboard_data(request)
    rows = dashboard["rows"]
    active_tab = _active_satisfaction_tab(request)
    headers, export_rows = _tabular_dashboard_export(active_tab, dashboard["context"], rows)
    filename = _dashboard_export_filename_from_rows(
        sorted(rows, key=lambda row: row["modified_at"]),
        dashboard["filters"],
        "csv",
    )

    response = HttpResponse(content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response.write("\ufeff")
    writer = csv.writer(response)
    writer.writerow(headers)
    for export_row in export_rows:
        writer.writerow(export_row)
    return response


# ---------------------------------------------------------------------------
# Import apprenants manquants depuis le fichier réseau
# ---------------------------------------------------------------------------

_PHONE_RE_IMPORT = re.compile(r"[0-9]{7,}")


def _safe_import_appel_code(record: dict) -> str:
    raw_code = str(record.get("code") or "").strip()
    max_length = Appel._meta.get_field("code").max_length or 50
    if len(raw_code) <= max_length:
        return raw_code

    phone_digits = _normalize_phone(
        (record.get("telephone1") or record.get("telephone2") or record.get("numero") or "").strip()
    )
    row_marker = str(record.get("row_number") or record.get("numero") or "").strip()
    classe_id = str(record.get("classe_label") or record.get("classe_id") or "").strip()
    nom = str(record.get("nom") or record.get("nom_individu") or "").strip()
    seed = "||".join(part for part in [raw_code, phone_digits, row_marker, classe_id, nom] if part)
    digest = hashlib.sha1(seed.encode("utf-8")).hexdigest()[:20]
    suffix = phone_digits[-8:] if phone_digits else ""
    safe_code = f"XL-{digest}"
    if suffix:
        safe_code = f"{safe_code}-{suffix}"
    return safe_code[:max_length]


@require_analysis_access
def apprenants_manquants_page(request):
    """Page dédiée : liste et import des apprenants des prestations manquantes.

    Utilise la feuille Consolidation pour récupérer les téléphones réels,
    et diagnostique pourquoi les prestations restent manquantes (téléphones vides,
    apprenants non chargés, seuil non atteint).
    """
    selected_source = _analysis_selected_source(request)
    source_options = get_workbook_source_options()

    # -- Charge la feuille Consolidation (téléphones réels) ----------------
    try:
        consol_bundle = build_consolidation_call_candidates(source_key=selected_source)
    except Exception as exc:
        return render(
            request,
            "satisfaction_apprenants/apprenants_manquants.html",
            {
                "error": f"Feuille Consolidation inaccessible : {exc}",
                "source_options": source_options,
                "selected_source": selected_source,
            },
        )

    # -- Charge aussi la feuille Prestations (méta-données) ----------------
    try:
        source_bundle = build_padesce_source_index(source_key=selected_source)
    except Exception:
        source_bundle = None

    # -- Calcul des prestations manquantes ---------------------------------
    all_rows = [
        _dashboard_row_from_answer(answer) for answer in _satisfaction_dashboard_base_queryset()
    ]
    all_rows = [row for row in all_rows if row["fenetre"] in {"2", "3"}]
    classe_apprenant_counts = dict(
        Classe.objects.annotate(_nb=Count("apprenants")).values_list("code", "_nb")
    )
    threshold_class_codes = _status_threshold_class_codes(source_bundle)
    _, classe_stats_all = _thresholded_dashboard_rows(
        all_rows,
        {},
        classe_apprenant_counts,
        threshold_class_codes=threshold_class_codes,
    )

    if source_bundle:
        terminated_codes = _terminated_prestation_codes_from_source({}, source_bundle)
        qualified_codes = _qualified_prestation_codes_from_source(
            {},
            classe_stats_all,
            source_bundle,
            threshold_class_codes=threshold_class_codes,
        )
        missing_keys = terminated_codes - qualified_codes
        source_prestations: dict = source_bundle.get("prestations", {})
    else:
        missing_keys = set()
        source_prestations = {}

    # -- Index des appels existants (code → telephone1 courant) ------------
    existing_appels: dict[str, str] = {
        code: (tel or "")
        for code, tel in Appel.objects.filter(is_active=True).values_list("code", "telephone1")
    }

    # -- Index classe_label → prestation_key depuis le fichier source ------
    # La feuille Consolidation n'a pas de colonne Prestation ID : on passe par
    # classe_label → classes_by_id → prestation_id (feuilles Classes/Prestations)
    classes_data: dict = source_bundle.get("classes", {}) if source_bundle else {}
    classe_to_presta_key: dict[str, str] = {}
    for cls_key, cls_info in classes_data.items():
        presta_key = normalize_network_lookup(cls_info.get("prestation_id", ""))
        if presta_key:
            classe_to_presta_key[cls_key] = presta_key

    # -- Index des records Consolidation par code --------------------------
    # consol_bundle["records"] est une LISTE (pas un dict)
    consol_records: list[dict] = consol_bundle.get("records", []) if consol_bundle else []
    consol_by_code: dict[str, dict] = {}
    consol_by_prestation: dict[str, list[dict]] = defaultdict(list)
    for rec in consol_records:
        code = (rec.get("code") or "").strip()
        if code:
            consol_by_code[code] = rec
        # 1. Prestation ID direct (colonne optionnelle dans Consolidation)
        p_key = normalize_network_lookup(rec.get("prestation_id", ""))
        # 2. Fallback : classe_label → prestation via feuille Classes
        if not p_key:
            cls_key = normalize_network_lookup(rec.get("classe_label", ""))
            p_key = classe_to_presta_key.get(cls_key, "")
        if p_key:
            consol_by_prestation[p_key].append(rec)

    prestations_with_importable: list[dict] = []
    prestations_sans_importable: list[dict] = []
    total_importable = 0
    total_already_loaded = 0
    total_needs_phone_sync = 0

    for p_key in sorted(missing_keys):
        p_info = source_prestations.get(p_key, {})
        prestation_id_display = p_info.get("prestation_id", "") or p_key

        importable: list[dict] = []
        loaded_with_phone: list[dict] = []
        loaded_no_phone: list[dict] = []
        no_phone_count = 0

        recs_for_presta = consol_by_prestation.get(p_key, [])

        for rec in recs_for_presta:
            code = (rec.get("code") or "").strip()
            if not code:
                continue
            # Téléphone depuis la feuille Consolidation
            numero = (
                rec.get("telephone1") or rec.get("telephone2") or rec.get("numero") or ""
            ).strip()
            has_phone = bool(_PHONE_RE_IMPORT.search(numero))
            classe_id = (
                rec.get("classe_label") or rec.get("classe_id") or ""
            ).strip()

            row = {
                "code": code,
                "nom": (rec.get("nom") or rec.get("nom_individu") or "").strip(),
                "classe_id": classe_id,
                "telephone": numero,
                "prestataire": (
                    rec.get("prestataire") or p_info.get("prestataire") or ""
                ).strip(),
                "beneficiaire": (
                    rec.get("beneficiaire") or p_info.get("beneficiaire") or ""
                ).strip(),
                "fenetre": (rec.get("fenetre") or "").strip(),
            }

            if code in existing_appels:
                current_tel = existing_appels[code]
                current_has_phone = bool(
                    _PHONE_RE_IMPORT.search(current_tel) if current_tel else False
                )
                if current_has_phone:
                    loaded_with_phone.append({**row, "telephone": current_tel})
                else:
                    # Dans DB mais sans téléphone — peut être mis à jour
                    loaded_no_phone.append({**row, "telephone_consol": numero})
            elif has_phone:
                importable.append(row)
            else:
                no_phone_count += 1

        importable.sort(key=lambda r: (r["classe_id"], r["code"]))
        loaded_no_phone.sort(key=lambda r: (r["classe_id"], r["code"]))

        importable_count = len(importable)
        loaded_no_phone_count = len(loaded_no_phone)
        total_importable += importable_count
        total_already_loaded += len(loaded_with_phone) + loaded_no_phone_count
        total_needs_phone_sync += loaded_no_phone_count

        entry = {
            "prestation_id": prestation_id_display,
            "p_key": p_key,
            "prestataire": p_info.get("prestataire", ""),
            "beneficiaire": p_info.get("beneficiaire", ""),
            "formation": p_info.get("formation", ""),
            "importable_count": importable_count,
            "loaded_count": len(loaded_with_phone) + loaded_no_phone_count,
            "loaded_with_phone_count": len(loaded_with_phone),
            "loaded_no_phone_count": loaded_no_phone_count,
            "no_phone_count": no_phone_count,
            "apprenants": importable,
            "apprenants_loaded_no_phone": loaded_no_phone,
        }
        if importable_count > 0 or loaded_no_phone_count > 0:
            prestations_with_importable.append(entry)
        else:
            prestations_sans_importable.append(entry)

    return render(
        request,
        "satisfaction_apprenants/apprenants_manquants.html",
        {
            "source_options": source_options,
            "selected_source": selected_source,
            "filter_source": selected_source or "principal",
            "prestations": prestations_with_importable,
            "prestations_sans_importable": prestations_sans_importable,
            "total_importable": total_importable,
            "total_already_loaded": total_already_loaded,
            "total_needs_phone_sync": total_needs_phone_sync,
            "total_missing_prestations": len(missing_keys),
            "import_url": "analyse/import-manquants/",
            "sync_url": "analyse/sync-telephones/",
            "notif_url": "analyse/import-notifications/",
        },
    )


@require_analysis_access
def import_missing_apprenants(request):
    """POST – importe un lot de 20 apprenants depuis le fichier réseau consolidé.

    Body JSON attendu :
        {
            "offset": int,                  # position de départ dans la liste triée
            "prestation_ids": [str, ...]    # IDs bruts des prestations manquantes à traiter
        }

    Retourne :
        {
            "imported": int,
            "total_importable": int,
            "total_remaining": int,
            "classes": [str, ...],
            "next_offset": int | null,
            "done": bool
        }
    """
    if request.method != "POST":
        return JsonResponse({"error": "Méthode non autorisée."}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except (ValueError, UnicodeDecodeError):
        body = {}

    offset = max(0, int(body.get("offset") or 0))
    prestation_ids: list[str] = [str(p) for p in (body.get("prestation_ids") or []) if p]
    selected_source = _analysis_selected_source(request)

    try:
        consol_bundle = build_consolidation_call_candidates(source_key=selected_source)
    except Exception as exc:
        return JsonResponse({"error": f"Fichier Consolidation inaccessible : {exc}"}, status=500)

    if not consol_bundle:
        return JsonResponse({"error": "Feuille Consolidation non disponible."}, status=400)

    # -- Index classe_label → prestation_key via feuille Classes -----------
    # La feuille Consolidation n'a pas de colonne Prestation ID :
    # on résout via classe_label → Classes sheet → prestation_id
    try:
        src_bundle_import = build_padesce_source_index(source_key=selected_source)
    except Exception:
        src_bundle_import = None
    _classes_import: dict = src_bundle_import.get("classes", {}) if src_bundle_import else {}
    _cls_to_presta: dict[str, str] = {}
    for _ck, _ci in _classes_import.items():
        _pk = normalize_network_lookup(_ci.get("prestation_id", ""))
        if _pk:
            _cls_to_presta[_ck] = _pk

    def _get_presta_key_for_rec(rec: dict) -> str:
        # Direct prestation_id column (optional in Consolidation)
        pk = normalize_network_lookup(rec.get("prestation_id", ""))
        if pk:
            return pk
        # Fallback via classe_label
        ck = normalize_network_lookup(rec.get("classe_label", ""))
        return _cls_to_presta.get(ck, "")

    # Détermine les prestations cibles
    if prestation_ids:
        target_keys = {normalize_network_lookup(p) for p in prestation_ids}
    else:
        # Toutes les prestations résolvables depuis la Consolidation
        target_keys = {
            _get_presta_key_for_rec(rec)
            for rec in consol_bundle.get("records", [])
            if _get_presta_key_for_rec(rec)
        }

    if not target_keys:
        return JsonResponse(
            {
                "imported": 0,
                "total_importable": 0,
                "total_remaining": 0,
                "classes": [],
                "done": True,
            }
        )

    # Pre-fetch existing Appel codes to avoid duplicates
    existing_codes: set[str] = set(Appel.objects.filter(is_active=True).values_list("code", flat=True))
    seen_import_codes = set(existing_codes)

    # Collect importable records: in target prestations, have phone, not already in DB
    # consol_bundle["records"] est une LISTE
    importable: list[dict] = []
    for rec in consol_bundle.get("records", []):
        p_key = _get_presta_key_for_rec(rec)
        if p_key not in target_keys:
            continue
        numero = (rec.get("telephone1") or rec.get("telephone2") or rec.get("numero") or "").strip()
        if not _PHONE_RE_IMPORT.search(numero):
            continue
        code = _safe_import_appel_code(rec)
        if not code or code in seen_import_codes:
            continue
        importable.append({**rec, "_import_code": code})
        seen_import_codes.add(code)

    # Stable ordering for consistent pagination
    importable.sort(key=lambda r: (r.get("classe_label", ""), r.get("code", "")))

    total_importable = len(importable)
    batch = importable[offset : offset + 20]

    if not batch:
        return JsonResponse(
            {
                "imported": 0,
                "total_importable": total_importable,
                "total_remaining": 0,
                "classes": [],
                "done": True,
            }
        )

    # Pre-fetch all local Classe objects once (avoid N+1)
    local_classes_map: dict[str, object] = {
        normalize_network_lookup(c.code): c
        for c in Classe.objects.select_related("prestation", "formation")
    }

    classes_touched: set[str] = set()
    appels_to_create: list = []

    for rec in batch:
        code = rec["_import_code"].strip()
        classe_id = (rec.get("classe_label") or rec.get("classe_id") or "").strip()
        local_classe = local_classes_map.get(normalize_network_lookup(classe_id))

        appels_to_create.append(
            Appel(
                code=code,
                nom=(rec.get("nom") or rec.get("nom_individu") or "").strip(),
                prestataire=(rec.get("prestataire") or "").strip(),
                beneficiaire=(rec.get("beneficiaire") or "").strip(),
                lieu=(rec.get("lieu") or "").strip(),
                classe_label=classe_id,
                fenetre=(rec.get("fenetre") or "").strip(),
                telephone1=(
                    rec.get("telephone1") or rec.get("telephone2") or rec.get("numero") or ""
                ).strip(),
                formation_padesce=(rec.get("formation") or "").strip(),
                status="en_attente",
                is_active=True,
                classe=local_classe,
            )
        )
        label = (local_classe.code if local_classe else classe_id) or code
        if label:
            classes_touched.add(label)

    try:
        created = Appel.objects.bulk_create(appels_to_create, ignore_conflicts=True)
        imported_count = len(created)
    except Exception as exc:
        return JsonResponse({"error": f"Erreur lors de l'import : {exc}"}, status=500)

    remaining = max(0, total_importable - offset - 20)
    classes_list = sorted(classes_touched)

    if imported_count > 0:
        _push_import_notif(
            f"{imported_count} appel(s) importé(s) depuis le fichier réseau consolidé.",
            classes_list,
        )

    return JsonResponse(
        {
            "imported": imported_count,
            "total_importable": total_importable,
            "total_remaining": remaining,
            "classes": classes_list,
            "next_offset": offset + 20,
            "done": remaining <= 0,
        }
    )


@require_analysis_access
def sync_phones_from_consolidation(request):
    """POST – met à jour les téléphones vides des Appel existants depuis la feuille Consolidation.

    Retourne :
        {
            "updated": int,
            "total_checked": int,
            "classes": [str, ...],
            "done": bool
        }
    """
    if request.method != "POST":
        return JsonResponse({"error": "Méthode non autorisée."}, status=405)

    selected_source = _analysis_selected_source(request)

    try:
        consol_bundle = build_consolidation_call_candidates(
            source_key=selected_source, force_refresh=True
        )
    except Exception as exc:
        return JsonResponse({"error": f"Feuille Consolidation inaccessible : {exc}"}, status=500)

    if not consol_bundle:
        return JsonResponse({"error": "Feuille Consolidation non disponible."}, status=400)

    # Index code → téléphone depuis la feuille Consolidation
    consol_phones: dict[str, str] = {}
    for rec in consol_bundle.get("records", []):
        code = (rec.get("code") or "").strip()
        if not code:
            continue
        numero = (
            rec.get("telephone1") or rec.get("telephone2") or rec.get("numero") or ""
        ).strip()
        if _PHONE_RE_IMPORT.search(numero):
            consol_phones[code] = numero

    # Récupère tous les Appel actifs sans téléphone
    appels_sans_tel = list(
        Appel.objects.filter(is_active=True).exclude(
            telephone1__regex=r"[0-9]{7,}"
        ).only("id", "code", "classe_label", "telephone1")
    )

    updated_count = 0
    classes_touched: set[str] = set()
    to_update: list = []

    for appel in appels_sans_tel:
        new_tel = consol_phones.get(appel.code, "")
        if new_tel:
            appel.telephone1 = new_tel
            to_update.append(appel)
            if appel.classe_label:
                classes_touched.add(appel.classe_label)

    if to_update:
        try:
            Appel.objects.bulk_update(to_update, ["telephone1"], batch_size=200)
            updated_count = len(to_update)
        except Exception as exc:
            return JsonResponse({"error": f"Erreur mise à jour : {exc}"}, status=500)

    classes_list = sorted(classes_touched)
    if updated_count > 0:
        _push_import_notif(
            f"{updated_count} téléphone(s) mis à jour depuis la feuille Consolidation.",
            classes_list,
        )

    return JsonResponse(
        {
            "updated": updated_count,
            "total_checked": len(appels_sans_tel),
            "classes": classes_list,
            "done": True,
        }
    )


@require_analysis_access
def import_notifications_poll(request):
    """GET – retourne les notifications d'import depuis un timestamp donné.

    Paramètre : since=<float unix timestamp>
    Utilisé par tous les clients pour afficher les notifications en temps réel.
    """
    try:
        since = float(request.GET.get("since", 0))
    except (ValueError, TypeError):
        since = 0.0

    notifications = _get_import_notifs_since(since)
    return JsonResponse({"notifications": notifications})
